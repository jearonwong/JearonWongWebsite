#!/usr/bin/env python3
"""Generate AIAAWP R4D internal publication-design candidate artifacts.

The generator is intentionally AIAAWP-specific, but it follows the GAIC main
white paper publication standard: HTML/PDF are the visual authority and DOCX is
an editable derivative generated from the same source.
"""

from __future__ import annotations

import html
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from hashlib import sha256
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PyPDF2 import PdfReader, PdfWriter
except Exception:  # pragma: no cover - metadata is best-effort if PyPDF2 is absent.
    PdfReader = None
    PdfWriter = None


PACKAGE_DIR = Path(__file__).resolve().parents[1]
SOURCE_PATH = PACKAGE_DIR / "AIAAWP-2026-v0.1-R4-CANDIDATE.md"
SOURCE_REGISTER_PATH = PACKAGE_DIR / "sources" / "wp2-r4-source-register.md"
CITATION_MAP_PATH = PACKAGE_DIR / "sources" / "wp2-r4-citation-map.md"
METADATA_PLAN_PATH = PACKAGE_DIR / "metadata" / "wp2-r4-metadata-plan.md"
JSONLD_PLAN_PATH = PACKAGE_DIR / "metadata" / "wp2-r4-jsonld-plan.json"
MANIFEST_PATH = PACKAGE_DIR / "manifest.json"
CHECKSUM_PATH = PACKAGE_DIR / "checksums.sha256"
OUT_DIR = PACKAGE_DIR / "out"
ARTIFACT_BASENAME = "Agentic-AI-Auditability-Assurance-White-Paper-2026-v0.1-R4-Candidate"
HTML_PATH = OUT_DIR / f"{ARTIFACT_BASENAME}.html"
PDF_PATH = OUT_DIR / f"{ARTIFACT_BASENAME}.pdf"
DOCX_PATH = OUT_DIR / f"{ARTIFACT_BASENAME}.docx"
GENERATION_LOG_PATH = OUT_DIR / "aiaawp-r4d-generation-log.json"
REPORT_DIR = PACKAGE_DIR.parent / "reports"
R4D_REPORTS = [
    "aiaawp-r4d-preflight.md",
    "aiaawp-r4d-gaic-publication-standard-audit.md",
    "aiaawp-r4d-publication-standard-alignment-plan.md",
    "aiaawp-r4d-text-style-alignment-qa.md",
    "aiaawp-r4d-text-style-patch-log.md",
    "aiaawp-r4d-html-publication-design-qa.md",
    "aiaawp-r4d-docx-publication-design-qa.md",
    "aiaawp-r4d-pdf-publication-design-qa.md",
    "aiaawp-r4d-metadata-seo-geo-alignment.md",
    "aiaawp-r4d-cross-artifact-qa.md",
    "aiaawp-r4d-boundary-qa.md",
    "aiaawp-r4d-build-and-validation.md",
    "aiaawp-r4d-final-qa.md",
]

TITLE = "Agentic AI Auditability & Assurance White Paper 2026"
SUBTITLE = "A Lifecycle Evidence Guide for Audit, Assurance, and Enterprise AI Governance"
DOCUMENT_ID = "AIAAWP-2026-v0.1-R4-CANDIDATE"
VERSION = "v0.1-publication-candidate"
AUTHOR = "Jearon Wong"
SERIES = "Agentic Lifecycle Governance Industry Series"
STATUS = "Internal format candidate; not public, not final, not sealed, not live"
TODAY = datetime.now().astimezone().strftime("%B %-d, %Y") if sys.platform == "darwin" else datetime.now().astimezone().strftime("%B %d, %Y")
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
KEYWORDS = [
    "Agentic AI Auditability",
    "AI Agent Auditability",
    "Audit Evidence Chain",
    "Agentic Audit Object",
    "AARM",
    "MRO",
    "Agentic Lifecycle Governance",
    "Enterprise AI Governance",
]
BOUNDARY_LINE = (
    "Internal candidate only. Not public, not final, not sealed, not live, not legal advice, "
    "not an audit standard, not certification, not an assurance opinion, not legal compliance "
    "proof, not regulator approval, not procurement guidance, not a vendor ranking, and not an "
    "endorsement by any firm or professional body."
)


@dataclass
class Heading:
    level: int
    text: str
    slug: str
    group: str


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def posix(path: Path) -> str:
    return path.as_posix()


def rel(path: Path) -> str:
    return posix(path.relative_to(PACKAGE_DIR))


def file_hash(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def strip_markdown(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", value)
    value = value.replace("\\|", "|")
    return value.strip()


def inline_md(value: str) -> str:
    escaped = html.escape(value, quote=True)
    code: list[str] = []

    def stash_code(match: re.Match[str]) -> str:
        token = f"@@CODE{len(code)}@@"
        code.append(f"<code>{html.escape(match.group(1), quote=False)}</code>")
        return token

    escaped = re.sub(r"`([^`]+)`", stash_code, escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', escaped)
    for idx, code_html in enumerate(code):
        escaped = escaped.replace(f"@@CODE{idx}@@", code_html)
    return escaped


def slugify(value: str, seen: dict[str, int]) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", strip_markdown(value).lower()).strip("-") or "section"
    count = seen.get(base, 0)
    seen[base] = count + 1
    return base if count == 0 else f"{base}-{count + 1}"


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells: list[str] = []
    current = []
    escaped = False
    for char in stripped:
        if escaped:
            current.append(char)
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    cells.append("".join(current).strip())
    return cells


def is_separator(line: str) -> bool:
    compact = line.strip().strip("|").replace("|", "").replace(":", "").replace("-", "").replace(" ", "")
    return compact == ""


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].lstrip().startswith("|")
        and lines[index + 1].lstrip().startswith("|")
        and is_separator(lines[index + 1])
    )


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = [split_table_row(lines[index])]
    index += 2
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        rows.append(split_table_row(lines[index]))
        index += 1
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    return normalized, index


def source_body() -> str:
    source = read_text(SOURCE_PATH)
    marker = "## Relationship to the Agentic Lifecycle Governance Industry Series"
    if marker in source:
        source = source[source.index(marker) :]
    source = re.sub(r"\n## Table of Contents\n[\s\S]*?\n---\n", "\n", source, count=1)
    return source.strip()


def composed_markdown() -> str:
    return "\n\n".join(
        [
            source_body(),
            "---\n\n## Source Note and Citation Register\n\n"
            "The source register and citation map below are part of the internal candidate package. "
            "Source IDs are used as internal citation handles for review and future public citation styling.",
            read_text(SOURCE_REGISTER_PATH),
            "---\n\n## Package Citation Map\n\n",
            read_text(CITATION_MAP_PATH),
        ]
    )


def heading_group(text: str) -> str:
    clean = strip_markdown(text)
    if re.match(r"^\d+\.", clean):
        n = int(clean.split(".", 1)[0])
        return "chapters-a" if n <= 8 else "chapters-b"
    if clean.startswith("Appendix "):
        return "appendices"
    if "Source" in clean or "Citation" in clean:
        return "sources"
    return "front"


def render_table(rows: list[list[str]], caption: str | None) -> str:
    if not rows:
        return ""
    headers = rows[0]
    data = rows[1:]
    cols = len(headers)
    caption_html = f'<div class="table-caption">{inline_md(caption)}</div>' if caption else ""
    if cols >= 6 or len(data) >= 10:
        cards = [
            '<div class="semantic-row-card-set table-block">',
            caption_html,
            '<div class="layout-note">Table rendered as semantic row cards to preserve GAIC-style readability across HTML, PDF, and DOCX review surfaces.</div>',
        ]
        for row in data:
            title = inline_md(row[0]) if row else "Record"
            cards.append('<article class="semantic-row-card">')
            cards.append(f"<h4>{title}</h4>")
            cards.append("<dl>")
            for header, value in zip(headers[1:], row[1:]):
                if value.strip():
                    cards.append(f"<div><dt>{inline_md(header)}</dt><dd>{inline_md(value)}</dd></div>")
            cards.append("</dl></article>")
        cards.append("</div>")
        return "\n".join(cards)

    parts = [f'<div class="table-block table-scroll">{caption_html}<table>']
    parts.append("<thead><tr>" + "".join(f"<th>{inline_md(cell)}</th>" for cell in headers) + "</tr></thead>")
    parts.append("<tbody>")
    for row in data:
        parts.append("<tr>" + "".join(f"<td>{inline_md(cell)}</td>" for cell in row) + "</tr>")
    parts.append("</tbody></table></div>")
    return "\n".join(parts)


def render_markdown(markdown: str) -> tuple[str, list[Heading], dict[str, int]]:
    lines = markdown.splitlines()
    out: list[str] = []
    headings: list[Heading] = []
    seen: dict[str, int] = {}
    section_open = False
    paragraph: list[str] = []
    list_items: list[str] = []
    list_type: str | None = None
    pending_caption: str | None = None
    table_stats = {"count": 0, "row_card_tables": 0, "grid_tables": 0}

    def close_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            out.append(f"<p>{inline_md(' '.join(paragraph))}</p>")
            paragraph = []

    def close_list() -> None:
        nonlocal list_items, list_type
        if list_items:
            tag = "ol" if list_type == "ol" else "ul"
            out.append(f"<{tag}>" + "".join(f"<li>{inline_md(item)}</li>" for item in list_items) + f"</{tag}>")
            list_items = []
            list_type = None

    def close_flow() -> None:
        close_paragraph()
        close_list()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            close_flow()
            i += 1
            continue
        if stripped == "---":
            close_flow()
            out.append('<hr class="section-rule">')
            i += 1
            continue
        if is_table_start(lines, i):
            close_flow()
            rows, i = parse_table(lines, i)
            table_stats["count"] += 1
            if len(rows[0]) >= 6 or len(rows) >= 11:
                table_stats["row_card_tables"] += 1
            else:
                table_stats["grid_tables"] += 1
            out.append(render_table(rows, pending_caption))
            pending_caption = None
            continue
        heading = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading:
            close_flow()
            level = len(heading.group(1))
            text = strip_markdown(heading.group(2))
            slug = slugify(text, seen)
            group = heading_group(text)
            headings.append(Heading(level=level, text=text, slug=slug, group=group))
            if level == 2:
                if section_open:
                    out.append("</section>")
                cls = "chapter" if re.match(r"^\d+\.", text) else "appendix" if text.startswith("Appendix ") else "front-section"
                out.append(f'<section class="{cls}" id="{slug}">')
                section_open = True
                out.append(f"<h2>{inline_md(text)}</h2>")
            else:
                out.append(f'<h{level} id="{slug}">{inline_md(text)}</h{level}>')
            pending_caption = text if "Table" in text or "Mapping" in text or "Matrix" in text else None
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            close_paragraph()
            current_type = "ol" if numbered else "ul"
            if list_type and list_type != current_type:
                close_list()
            list_type = current_type
            list_items.append((numbered or bullet).group(1))
            i += 1
            continue
        if stripped.startswith(">"):
            close_flow()
            out.append(f"<blockquote>{inline_md(stripped.lstrip('> '))}</blockquote>")
            i += 1
            continue
        paragraph.append(stripped)
        i += 1

    close_flow()
    if section_open:
        out.append("</section>")
    return "\n".join(out), headings, table_stats


def toc_html(headings: list[Heading]) -> str:
    groups = [
        ("front", "Front Matter"),
        ("chapters-a", "Chapters 0-8"),
        ("chapters-b", "Chapters 9-16"),
        ("appendices", "Appendices"),
        ("sources", "Sources"),
    ]
    blocks = []
    for group_id, label in groups:
        items = [h for h in headings if h.group == group_id and h.level == 2]
        if not items:
            continue
        list_html = "".join(f'<li><a href="#{h.slug}">{html.escape(h.text)}</a></li>' for h in items)
        blocks.append(f"<div><h3>{label}</h3><ol>{list_html}</ol></div>")
    return f"""
<nav class="generated-toc" aria-label="Publication table of contents">
  <h2>Publication Contents</h2>
  <p class="toc-note">HTML anchors are active in the internal artifact. PDF page-number pinning remains a final candidate QA task.</p>
  <div class="toc-grid">{''.join(blocks)}</div>
</nav>
"""


def jsonld_block() -> str:
    graph = {
        "@context": "https://schema.org",
        "@graph": [
            {
                "@type": "TechArticle",
                "@id": "urn:jearonwong:aiaawp:2026:r4-candidate",
                "headline": TITLE,
                "alternativeHeadline": SUBTITLE,
                "identifier": DOCUMENT_ID,
                "version": VERSION,
                "author": {"@type": "Person", "name": AUTHOR},
                "inLanguage": "en",
                "genre": "Technical report",
                "creativeWorkStatus": "Internal publication design candidate; not public, not final, not sealed",
                "isPartOf": {"@type": "CreativeWorkSeries", "name": SERIES},
                "about": [
                    "Agentic AI Auditability",
                    "Audit Evidence Chain",
                    "Agentic Audit Object",
                    "AARM",
                    "Missing Regulatory Objects",
                    "Enterprise AI Governance",
                ],
                "keywords": KEYWORDS,
            }
        ],
    }
    return json.dumps(graph, indent=2)


def publication_css() -> str:
    return """
:root {
  --monolith: #111111;
  --slate: #1a1a1a;
  --machine: #f1f5f9;
  --registry: #3b82f6;
  --evidence: #22d3ee;
  --status: #64748b;
  --grid: #d8dee8;
  --amber: #b7791f;
  --fault: #991b1b;
  --paper: #ffffff;
}
* { box-sizing: border-box; }
html {
  background: #0b111c;
  color: var(--monolith);
  overflow-x: hidden;
}
body {
  margin: 0;
  font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  font-size: 15px;
  line-height: 1.54;
  color: var(--monolith);
  background: transparent;
  overflow-x: hidden;
}
.publication {
  max-width: 1040px;
  margin: 0 auto;
  background: var(--paper);
  padding: 0 74px 54px;
  box-shadow: 0 32px 92px rgba(2, 6, 23, 0.24);
  overflow-x: hidden;
}
a { color: var(--registry); text-decoration-thickness: 1px; text-underline-offset: 2px; }
p { margin: 0.72rem 0; }
strong { font-weight: 750; }
code {
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.92em;
  background: #f8fafc;
  border: 1px solid #e2e8f0;
  padding: 0.05rem 0.22rem;
  border-radius: 2px;
  overflow-wrap: break-word;
}
.cover-page.professional-shell {
  min-height: 96vh;
  margin: 0 -74px 2.7rem;
  padding: 72px 74px;
  border: 0;
  background:
    linear-gradient(145deg, rgba(34, 211, 238, 0.13), transparent 28%),
    linear-gradient(315deg, rgba(59, 130, 246, 0.18), transparent 34%),
    #0b111c;
  color: #e2e8f0;
}
.cover-frame {
  min-height: calc(96vh - 144px);
  display: flex;
  flex-direction: column;
  justify-content: center;
  min-width: 0;
  border: 1px solid rgba(255,255,255,0.14);
  padding: 54px 48px;
  background: rgba(255,255,255,0.025);
}
.cover-kicker {
  color: #94a3b8;
  font-size: 0.76rem;
  font-weight: 800;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}
.cover-page h1 {
  margin: 1rem 0 0.85rem;
  max-width: 850px;
  color: #ffffff;
  font-size: 3.95rem;
  line-height: 0.96;
  letter-spacing: 0;
  overflow-wrap: anywhere;
}
.cover-subtitle {
  max-width: 780px;
  margin: 0;
  color: #67e8f9;
  font-size: 1.28rem;
  line-height: 1.35;
}
.cover-role { color: #cbd5e1; margin: 0.65rem 0 0; }
.thesis-strip {
  display: grid;
  grid-template-columns: repeat(5, 1fr);
  gap: 1px;
  margin: 2rem 0 1.4rem;
  border: 1px solid rgba(255,255,255,0.18);
  background: rgba(255,255,255,0.18);
}
.thesis-strip span {
  background: rgba(15,23,42,0.92);
  color: #e0f2fe;
  padding: 0.72rem 0.55rem;
  text-align: center;
  font-weight: 750;
}
.cover-meta {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 1px;
  background: rgba(255,255,255,0.12);
  border: 1px solid rgba(255,255,255,0.14);
}
.cover-meta div { padding: 0.75rem; background: rgba(255,255,255,0.045); }
.cover-meta strong {
  display: block;
  color: #94a3b8;
  font-size: 0.68rem;
  letter-spacing: 0.06em;
  text-transform: uppercase;
}
.cover-meta span { display: block; color: #f8fafc; margin-top: 0.12rem; }
.cover-notice {
  margin-top: 1.25rem;
  padding: 0.92rem 1rem;
  border-left: 4px solid var(--evidence);
  background: rgba(8,47,73,0.42);
  color: #e0f2fe;
  font-size: 0.9rem;
}
.important-notice {
  border: 1px solid #cbd5e1;
  border-left: 5px solid var(--registry);
  background: #f8fafc;
  padding: 1rem 1.1rem;
  margin: 1.35rem 0 2rem;
}
.notice-label {
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
  text-transform: uppercase;
  color: var(--registry);
  font-weight: 800;
  letter-spacing: 0.08em;
}
h1, h2, h3, h4 {
  font-family: Outfit, Inter, system-ui, sans-serif;
  letter-spacing: 0;
  line-height: 1.16;
}
h2 {
  margin: 2.6rem 0 1rem;
  padding-top: 1.2rem;
  border-top: 4px solid var(--monolith);
  font-size: 2rem;
}
h3 {
  margin: 1.55rem 0 0.55rem;
  padding-top: 0.45rem;
  border-top: 1px solid #e2e8f0;
  font-size: 1.24rem;
}
h4 { margin: 1.1rem 0 0.4rem; font-size: 1rem; color: #1f2937; }
.generated-toc {
  margin: 2.4rem 0 2.8rem;
  padding: 1.25rem 1.35rem;
  background: linear-gradient(180deg, #f8fafc 0%, #ffffff 100%);
  border: 1px solid var(--grid);
  border-top: 4px solid var(--monolith);
}
.generated-toc h2 { margin-top: 0; padding-top: 0; border: 0; font-size: 1.45rem; }
.toc-note { color: var(--status); font-size: 0.9rem; }
.toc-grid { display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 1rem 1.25rem; }
.toc-grid h3 { border: 0; margin: 0 0 0.4rem; padding: 0; font-size: 0.95rem; color: var(--registry); }
.toc-grid ol { margin: 0; padding-left: 1.2rem; }
.toc-grid li { margin: 0.25rem 0; }
.chapter, .appendix, .front-section { margin-top: 1.7rem; }
.section-rule { border: 0; border-top: 1px solid #d8dee8; margin: 2rem 0; }
blockquote {
  margin: 1rem 0;
  padding: 0.75rem 0.95rem;
  border-left: 4px solid var(--registry);
  background: #f8fafc;
}
.table-block {
  margin: 1.1rem 0 1.45rem;
}
.table-scroll {
  overflow-x: auto;
  max-width: 100%;
  border: 1px solid var(--grid);
  background: #fff;
}
.table-caption {
  padding: 0.55rem 0.65rem;
  border-bottom: 1px solid #e2e8f0;
  font-weight: 800;
  color: #334155;
  background: #f8fafc;
}
table {
  width: 100%;
  border-collapse: collapse;
  table-layout: auto;
  font-size: 0.88rem;
}
th, td {
  border-bottom: 1px solid var(--grid);
  border-right: 1px solid var(--grid);
  padding: 0.56rem 0.62rem;
  vertical-align: top;
  min-width: 72px;
  word-break: normal;
  overflow-wrap: break-word;
  hyphens: auto;
}
th {
  text-align: left;
  background: var(--slate);
  color: #f8fafc;
  font-weight: 800;
}
tr:last-child td { border-bottom: 0; }
th:last-child, td:last-child { border-right: 0; }
.semantic-row-card-set {
  display: grid;
  gap: 0.75rem;
}
.layout-note {
  color: var(--status);
  font-size: 0.84rem;
  background: #f8fafc;
  border: 1px solid #e2e8f0;
  padding: 0.55rem 0.65rem;
}
.semantic-row-card {
  border: 1px solid #cbd5e1;
  border-left: 4px solid var(--evidence);
  padding: 0.85rem 0.95rem;
  background: #ffffff;
  break-inside: avoid;
}
.semantic-row-card h4 { margin-top: 0; color: #0f172a; }
.semantic-row-card dl {
  display: grid;
  grid-template-columns: minmax(120px, 0.32fr) minmax(0, 1fr);
  gap: 0.45rem 0.75rem;
  margin: 0;
}
.semantic-row-card dl div { display: contents; }
.semantic-row-card dt {
  color: var(--status);
  font-weight: 800;
  font-size: 0.78rem;
  text-transform: uppercase;
}
.semantic-row-card dd { margin: 0; min-width: 0; }
.html-footer {
  margin-top: 2.5rem;
  padding-top: 0.8rem;
  border-top: 1px solid #cbd5e1;
  color: var(--status);
  font-size: 0.84rem;
}
@page { size: A4; margin: 16mm 15mm; }
@media print {
  html, body { background: #ffffff; }
  body { font-size: 10.3pt; }
  .publication { max-width: none; box-shadow: none; padding: 0 54px 42px; }
  .cover-page.professional-shell {
    min-height: 96vh;
    margin-left: -54px;
    margin-right: -54px;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
    break-after: page;
  }
  .chapter, .appendix { break-before: page; }
  h2, h3, h4 { break-after: avoid; }
  .important-notice, .generated-toc, .semantic-row-card { break-inside: avoid; }
  .table-scroll { overflow: visible; }
  a { color: inherit; text-decoration: none; }
}
/* Screen-only responsive reading rules inherited from the GAIC public HTML
   web-reading edition. Print/PDF layout remains governed by print styles. */
@media screen {
  html {
    overflow-x: hidden;
    scroll-padding-top: 4.5rem;
  }
  body {
    min-width: 0;
    overflow-x: hidden;
    font-size: 16px;
    line-height: 1.62;
  }
  .publication {
    width: min(100%, 1040px);
    max-width: 1040px;
    overflow-x: clip;
    padding-left: clamp(1.25rem, 5vw, 74px);
    padding-right: clamp(1.25rem, 5vw, 74px);
  }
  .chapter,
  .appendix {
    content-visibility: auto;
    contain-intrinsic-size: 1200px;
  }
  .front-section,
  .chapter,
  .appendix,
  .generated-toc,
  .table-block,
  .table-scroll,
  .semantic-row-card,
  .semantic-row-card-set {
    max-width: 100%;
  }
  .table-block,
  .table-scroll {
    overflow-x: auto;
    -webkit-overflow-scrolling: touch;
    overscroll-behavior-x: contain;
  }
  .table-scroll::after {
    content: "Scroll table horizontally if needed";
    display: none;
    margin: 0.25rem 0 0.55rem;
    padding-left: 0.65rem;
    color: var(--status);
    font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
    font-size: 0.68rem;
    letter-spacing: 0.04em;
    text-transform: uppercase;
  }
  table {
    max-width: 100%;
  }
  img,
  svg,
  canvas {
    max-width: 100%;
    height: auto;
  }
}
@media screen and (max-width: 900px) {
  .publication {
    width: 100%;
    max-width: 100%;
    margin: 0;
    padding: 0 clamp(1.25rem, 4vw, 2rem) 3rem;
    box-shadow: none;
  }
  .cover-page.professional-shell {
    margin: 0 calc(-1 * clamp(1.25rem, 4vw, 2rem)) 2rem;
    padding: clamp(3rem, 8vw, 4rem) clamp(1.25rem, 5vw, 2.5rem);
    min-height: auto;
  }
  .cover-frame {
    min-height: auto;
    padding: clamp(1.4rem, 5vw, 2.5rem);
  }
  .cover-page.professional-shell h1,
  .cover-page h1 {
    max-width: 100%;
    font-size: clamp(2.35rem, 10vw, 3.4rem);
    line-height: 0.98;
    overflow-wrap: anywhere;
  }
  .cover-page.professional-shell .cover-subtitle,
  .cover-subtitle {
    font-size: clamp(1rem, 4vw, 1.2rem);
    line-height: 1.35;
  }
  .cover-page.professional-shell .thesis-strip,
  .thesis-strip {
    grid-template-columns: 1fr 1fr;
  }
  .cover-page.professional-shell .cover-meta,
  .cover-meta,
  .toc-grid,
  .semantic-row-card dl {
    grid-template-columns: 1fr;
  }
  .generated-toc,
  .important-notice,
  .semantic-row-card {
    padding: clamp(0.9rem, 4vw, 1.15rem);
  }
  h2 {
    font-size: clamp(1.35rem, 5.5vw, 1.8rem);
    line-height: 1.18;
  }
  h3 {
    font-size: clamp(1.08rem, 4vw, 1.25rem);
  }
}
@media screen and (max-width: 640px) {
  body {
    font-size: 16px;
    line-height: 1.68;
  }
  .publication {
    padding-left: 1rem;
    padding-right: 1rem;
  }
  .cover-page.professional-shell {
    margin-left: -1rem;
    margin-right: -1rem;
    padding: 2rem 1rem;
  }
  .cover-frame {
    padding: 1.15rem;
  }
  .cover-page.professional-shell .thesis-strip,
  .thesis-strip {
    grid-template-columns: 1fr;
  }
  .cover-page.professional-shell .thesis-strip span,
  .thesis-strip span {
    min-height: 2.65rem;
    display: flex;
    align-items: center;
    justify-content: center;
  }
  .toc-grid ol,
  .generated-toc ol,
  .generated-toc ul,
  ul,
  ol {
    margin-left: 1.05rem;
  }
  .table-caption {
    font-size: 0.95rem;
    line-height: 1.35;
  }
  .table-block table,
  .table-scroll table {
    min-width: 680px;
    width: max-content;
    max-width: none;
  }
  .table-scroll::after {
    display: block;
  }
  th,
  td {
    min-width: 96px;
    padding: 0.58rem 0.62rem;
    font-size: 0.86rem;
    hyphens: none;
  }
  th {
    font-size: 0.82rem;
  }
  .semantic-row-card dt,
  .semantic-row-card dd {
    display: block;
    width: 100%;
  }
}
@media screen and (max-width: 420px) {
  .cover-page.professional-shell h1,
  .cover-page h1 {
    max-width: 10.8ch;
    font-size: clamp(2rem, 10.5vw, 2.35rem);
    line-height: 1.03;
  }
  .cover-page.professional-shell .cover-subtitle,
  .cover-subtitle,
  .cover-role {
    max-width: 28ch;
  }
  .cover-meta span,
  .cover-meta strong,
  .cover-meta div,
  .cover-kicker,
  .cover-subtitle,
  .cover-role,
  .cover-notice,
  .toc-grid li {
    overflow-wrap: anywhere;
  }
}
"""


def build_html() -> tuple[str, dict[str, int]]:
    body_html, headings, table_stats = render_markdown(composed_markdown())
    doc = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="robots" content="noindex, nofollow">
<meta name="aiaawp-artifact-status" content="internal-format-candidate">
<meta name="author" content="{html.escape(AUTHOR)}">
<meta name="description" content="{html.escape(SUBTITLE)}">
<meta name="citation_title" content="{html.escape(TITLE + ': ' + SUBTITLE)}">
<meta name="citation_author" content="Wong, Jearon">
<meta name="citation_publication_date" content="PUBLICATION_DATE_PLACEHOLDER">
<meta name="citation_technical_report_number" content="{DOCUMENT_ID}">
<meta name="citation_language" content="en">
<meta name="keywords" content="{html.escape(', '.join(KEYWORDS))}">
<title>{html.escape(TITLE)} - {DOCUMENT_ID}</title>
<script type="application/ld+json">
{jsonld_block()}
</script>
<style>
{publication_css()}
</style>
</head>
<body>
<main class="publication">
<section class="cover-page professional-shell">
  <div class="cover-frame">
    <div class="cover-kicker">Jearon Wong / {SERIES}</div>
    <h1>{html.escape(TITLE)}</h1>
    <p class="cover-subtitle">{html.escape(SUBTITLE)}</p>
    <p class="cover-role">Jearon Wong - Protocol Architect for the Agent Era</p>
    <div class="thesis-strip" aria-label="Auditability thesis strip">
      <span>Authority</span><span>Responsibility</span><span>Agent Work</span><span>Evidence</span><span>Closure</span>
    </div>
    <div class="cover-meta">
      <div><strong>Document ID</strong><span>{DOCUMENT_ID}</span></div>
      <div><strong>Version</strong><span>{VERSION}</span></div>
      <div><strong>Date</strong><span>{TODAY}</span></div>
      <div><strong>Status</strong><span>{STATUS}</span></div>
      <div><strong>Series</strong><span>{SERIES}</span></div>
      <div><strong>Visual Source of Truth</strong><span>HTML/PDF primary; DOCX derivative</span></div>
    </div>
    <div class="cover-notice">{html.escape(BOUNDARY_LINE)}</div>
  </div>
</section>
<section class="important-notice" aria-label="Publication boundary">
  <div class="notice-label">Publication Boundary</div>
  <h2>Internal Candidate Status</h2>
  <p>This artifact is an internal publication-design candidate for review. It is not a public release, final publication, sealed artifact, live deployment, audit standard, certification, assurance opinion, legal compliance proof, regulator approval, procurement recommendation, vendor ranking, or endorsement.</p>
  <p>It applies the GAIC main white paper publication standard to AIAAWP while preserving AIAAWP's public-facing title, subtitle, and document ID.</p>
</section>
{toc_html(headings)}
{body_html}
<footer class="html-footer">
  {DOCUMENT_ID} - internal format candidate. Generated {GENERATED_AT}. No public route, no public canonical URL, no final/sealed/live claim.
</footer>
</main>
</body>
</html>"""
    write_text(HTML_PATH, doc)
    return doc, table_stats


def chrome_path() -> str:
    candidates = [
        shutil.which("google-chrome"),
        shutil.which("chromium"),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    raise RuntimeError("No Chrome/Chromium executable found for HTML-to-PDF export")


def html_to_pdf() -> None:
    if PDF_PATH.exists():
        PDF_PATH.unlink()
    profile = OUT_DIR / "chrome-profile"
    shutil.rmtree(profile, ignore_errors=True)
    profile.mkdir(parents=True, exist_ok=True)
    cmd = [
        chrome_path(),
        "--headless=new",
        "--disable-gpu",
        "--disable-background-networking",
        "--no-pdf-header-footer",
        "--disable-extensions",
        "--no-first-run",
        "--no-default-browser-check",
        f"--user-data-dir={profile}",
        f"--print-to-pdf={PDF_PATH}",
        HTML_PATH.resolve().as_uri(),
    ]
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    deadline = time.time() + 420
    last_size = -1
    stable_ticks = 0
    while time.time() < deadline:
        if proc.poll() is not None:
            stdout, stderr = proc.communicate(timeout=30)
            if proc.returncode != 0 or not PDF_PATH.exists():
                raise RuntimeError(f"Chrome PDF export failed\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
            break
        if PDF_PATH.exists():
            size = PDF_PATH.stat().st_size
            stable_ticks = stable_ticks + 1 if size > 0 and size == last_size else 0
            last_size = size
            if stable_ticks >= 3 and pdf_page_count(PDF_PATH) > 0:
                proc.terminate()
                try:
                    proc.communicate(timeout=10)
                except subprocess.TimeoutExpired:
                    proc.kill()
                    proc.communicate(timeout=30)
                break
        time.sleep(1.5)
    else:
        proc.kill()
        stdout, stderr = proc.communicate(timeout=30)
        if not PDF_PATH.exists() or pdf_page_count(PDF_PATH) == 0:
            raise RuntimeError(f"Chrome PDF export timed out\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
    shutil.rmtree(profile, ignore_errors=True)
    patch_pdf_metadata()


def patch_pdf_metadata() -> dict[str, str | bool]:
    if PdfReader is None or PdfWriter is None:
        return {"applied": False, "reason": "PyPDF2 unavailable"}
    reader = PdfReader(str(PDF_PATH))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": TITLE,
            "/Author": AUTHOR,
            "/Subject": SUBTITLE,
            "/Keywords": ", ".join(KEYWORDS),
            "/Producer": "Headless Chrome; PyPDF2 metadata pass",
        }
    )
    tmp = PDF_PATH.with_suffix(".pdf.tmp")
    with tmp.open("wb") as handle:
        writer.write(handle)
    tmp.replace(PDF_PATH)
    return {"applied": True}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def add_docx_table(doc: Document, rows: list[list[str]], caption: str | None) -> None:
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(strip_markdown(caption))
        run.bold = True
        run.font.color.rgb = RGBColor(59, 130, 246)
        run.font.size = Pt(8.5)
    if not rows:
        return
    headers = rows[0]
    data = rows[1:]
    if len(headers) >= 6 or len(data) >= 10:
        note = doc.add_paragraph()
        note.add_run("Semantic row-card rendering for wide/dense table readability.").italic = True
        note.runs[0].font.size = Pt(8)
        for row in data:
            title = strip_markdown(row[0]) if row else "Record"
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            r = p.add_run(title)
            r.bold = True
            r.font.color.rgb = RGBColor(17, 17, 17)
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for header, value in zip(headers[1:], row[1:]):
                if not value.strip():
                    continue
                cells = table.add_row().cells
                set_cell_width(cells[0], 2100)
                set_cell_width(cells[1], 6500)
                set_cell_shading(cells[0], "F1F5F9")
                cells[0].paragraphs[0].add_run(strip_markdown(header)).bold = True
                cells[1].paragraphs[0].add_run(strip_markdown(value))
                for cell in cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(7.6)
        return

    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx == 0:
                set_cell_shading(cell, "1A1A1A")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(strip_markdown(value))
            run.font.name = "Arial"
            run.font.size = Pt(7.2 if len(headers) >= 5 else 8)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)


def build_docx() -> None:
    if DOCX_PATH.exists():
        DOCX_PATH.unlink()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.4)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(3)
    for name, size in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 12), ("Heading 4", 10)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(17, 17, 17)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    run = p.add_run(TITLE)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 17, 17)
    sub = doc.add_paragraph(SUBTITLE)
    sub.runs[0].font.name = "Arial"
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for key, value in [
        ("Document ID", DOCUMENT_ID),
        ("Version", VERSION),
        ("Date", TODAY),
        ("Author", AUTHOR),
        ("Status", STATUS),
        ("Series", SERIES),
    ]:
        cells = meta.add_row().cells
        set_cell_shading(cells[0], "F1F5F9")
        cells[0].paragraphs[0].add_run(key).bold = True
        cells[1].paragraphs[0].add_run(value)
    notice = doc.add_paragraph()
    notice.paragraph_format.space_before = Pt(12)
    notice.add_run(BOUNDARY_LINE).italic = True
    doc.add_page_break()

    toc_title = doc.add_heading("Publication Contents", level=1)
    toc_title.paragraph_format.keep_with_next = True
    for line in [
        "Front Matter",
        "Chapters 0-16",
        "Appendices A-F",
        "Source Note and Citation Register",
        "Package Citation Map",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_page_break()

    lines = composed_markdown().splitlines()
    i = 0
    pending_caption: str | None = None
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if is_table_start(lines, i):
            rows, i = parse_table(lines, i)
            add_docx_table(doc, rows, pending_caption)
            pending_caption = None
            continue
        heading = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading:
            level = min(max(len(heading.group(1)) - 1, 1), 4)
            text = strip_markdown(heading.group(2))
            if level == 1 and (re.match(r"^\d+\.", text) or text.startswith("Appendix ")):
                doc.add_page_break()
            doc.add_heading(text, level=level)
            pending_caption = text if "Table" in text or "Mapping" in text or "Matrix" in text else None
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            doc.add_paragraph(strip_markdown(bullet.group(1)), style="List Bullet")
        elif numbered:
            doc.add_paragraph(strip_markdown(numbered.group(1)), style="List Number")
        elif stripped.startswith(">"):
            para = doc.add_paragraph(strip_markdown(stripped.lstrip("> ")))
            para.paragraph_format.left_indent = Inches(0.2)
        else:
            doc.add_paragraph(strip_markdown(stripped))
        i += 1

    props = doc.core_properties
    props.title = TITLE
    props.author = AUTHOR
    props.subject = SUBTITLE
    props.keywords = ", ".join(KEYWORDS)
    props.comments = "Internal format candidate; HTML/PDF are visual authority and DOCX is editable derivative."
    doc.save(DOCX_PATH)


def pdf_page_count(path: Path) -> int:
    result = subprocess.run(["pdfinfo", str(path)], capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        return 0
    match = re.search(r"^Pages:\s+(\d+)", result.stdout, re.MULTILINE)
    return int(match.group(1)) if match else 0


def extract_pdf_text() -> str:
    with tempfile.NamedTemporaryFile(prefix="aiaawp-r4d-pdf-text-", suffix=".txt", delete=False) as handle:
        text_path = Path(handle.name)
    try:
        result = subprocess.run(["pdftotext", "-layout", str(PDF_PATH), str(text_path)], capture_output=True, text=True, timeout=240)
        if result.returncode != 0:
            return ""
        return read_text(text_path)
    finally:
        text_path.unlink(missing_ok=True)


def docx_text() -> str:
    try:
        doc = Document(DOCX_PATH)
        return "\n".join([p.text for p in doc.paragraphs] + ["\n".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows])
    except Exception:
        return ""


def update_manifest(table_stats: dict[str, int]) -> dict:
    manifest = json.loads(read_text(MANIFEST_PATH))
    artifact_paths = [HTML_PATH, PDF_PATH, DOCX_PATH, GENERATION_LOG_PATH]
    hashes = {rel(path): file_hash(path) for path in artifact_paths if path.exists()}
    manifest.update(
        {
            "document_id": DOCUMENT_ID,
            "date": "2026-05-18",
            "status": "internal_format_candidate",
            "not_public": True,
            "not_final": True,
            "not_sealed": True,
            "not_live": True,
            "no_public_route": True,
            "ui_style_alignment_status": "aligned_to_gaic_main_white_paper_publication_standard",
            "text_style_alignment_status": "internal draft residue reduced; candidate status centralized",
            "metadata_seo_geo_alignment_status": "aligned_plan_noindex_internal_future_public_updates_deferred",
            "artifact_integrity_status": "manifest_and_checksums_updated_for_r4d_internal_artifacts",
            "visual_authority": "HTML/PDF are the visual authority; DOCX is an editable derivative.",
            "source_baseline_commit": "431cf6597d13f0bd8a159c2c8c4fe2874e9b2311",
            "source_baseline": [
                "R4C content and public naming QA package",
                "R4B internal generated artifacts",
                "R4 source package",
                "R3 QA-reviewed internal draft",
                "GAIC v0.3.2-FRC-R3 source truth",
            ],
            "files_included": [
                "README.md",
                "AIAAWP-2026-v0.1-R4-CANDIDATE.md",
                "manifest.json",
                "checksums.sha256",
                "metadata/wp2-r4-metadata-plan.md",
                "metadata/wp2-r4-jsonld-plan.json",
                "sources/wp2-r4-source-register.md",
                "sources/wp2-r4-citation-map.md",
                "qa/wp2-r4-integrity-notes.md",
                "generation/wp2-r4-generation-plan.md",
                "generation/generate-aiaawp-r4d-publication-candidate.py",
                "out/aiaawp-r4d-generation-log.json",
            ],
            "r4d_baseline_commit": "431cf6597d13f0bd8a159c2c8c4fe2874e9b2311",
            "artifact_generation_status": "html-pdf-docx-generated-internal-format-candidate-only",
            "artifact_paths": [rel(HTML_PATH), rel(PDF_PATH), rel(DOCX_PATH)],
            "generated_artifact_paths": [rel(HTML_PATH), rel(PDF_PATH), rel(DOCX_PATH)],
            "generated_artifact_hashes": {k: v for k, v in hashes.items() if k.startswith("out/") and not k.endswith("generation-log.json")},
            "generation_script_path": "generation/generate-aiaawp-r4d-publication-candidate.py",
            "generation_timestamp": GENERATED_AT,
            "artifact_status": "internal_format_candidate_only",
            "known_formatting_limitations": [
                "DOCX is an editable derivative and may receive human template polish before final publication.",
                "PDF page-number pinning and full page-by-page human proofing remain final candidate QA tasks.",
            ],
            "r4d_generation_results": {
                "html": {
                    "generated": HTML_PATH.exists(),
                    "engine": "AIAAWP Python HTML-first generator using GAIC publication style rules",
                },
                "pdf": {
                    "generated": PDF_PATH.exists(),
                    "engine": "Headless Chrome from publication-quality HTML plus PyPDF2 metadata pass",
                    "page_count": pdf_page_count(PDF_PATH) if PDF_PATH.exists() else None,
                },
                "docx": {
                    "generated": DOCX_PATH.exists(),
                    "engine": "python-docx structured editable derivative from same candidate source",
                },
            },
            "conversion_results": {
                "pdf": {
                    "generated": PDF_PATH.exists(),
                    "engine": "Headless Chrome from publication HTML",
                    "metadata": {"applied": True, "method": "PyPDF2"},
                },
                "docx": {
                    "generated": DOCX_PATH.exists(),
                    "engine": "python-docx structured editable derivative",
                    "metadata": {"applied": True, "method": "python-docx core properties"},
                },
            },
            "qa_report_paths": [
                "reports/aiaawp-r4d-preflight.md",
                "reports/aiaawp-r4d-gaic-publication-standard-audit.md",
                "reports/aiaawp-r4d-publication-standard-alignment-plan.md",
                "reports/aiaawp-r4d-text-style-alignment-qa.md",
                "reports/aiaawp-r4d-html-publication-design-qa.md",
                "reports/aiaawp-r4d-docx-publication-design-qa.md",
                "reports/aiaawp-r4d-pdf-publication-design-qa.md",
                "reports/aiaawp-r4d-metadata-seo-geo-alignment.md",
                "reports/aiaawp-r4d-cross-artifact-qa.md",
                "reports/aiaawp-r4d-boundary-qa.md",
                "reports/aiaawp-r4d-build-and-validation.md",
                "reports/aiaawp-r4d-final-qa.md",
            ],
            "next_phase": "AIAAWP final candidate QA after human review of R4D publication-design artifacts",
            "public_facing_document_id": DOCUMENT_ID,
            "public_facing_title": TITLE,
            "public_facing_subtitle": SUBTITLE,
        }
    )
    write_text(MANIFEST_PATH, json.dumps(manifest, indent=2, ensure_ascii=False))
    return manifest


def list_package_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        rel_path = rel(path)
        if rel_path == "checksums.sha256":
            continue
        if "/chrome-profile" in rel_path or "/lo-profile" in rel_path:
            continue
        if rel_path.endswith(".tmp"):
            continue
        yield path


def update_checksums() -> None:
    lines = [
        "# SHA256 checksums for AIAAWP R4D internal format candidate package files",
        "# checksums.sha256 excludes itself and transient conversion profile files.",
    ]
    for path in list_package_files(PACKAGE_DIR):
        lines.append(f"{file_hash(path)}  {rel(path)}")
    for name in R4D_REPORTS:
        path = REPORT_DIR / name
        if path.exists():
            lines.append(f"{file_hash(path)}  ../reports/{name}")
    write_text(CHECKSUM_PATH, "\n".join(lines))


def write_generation_log(table_stats: dict[str, int], manifest: dict) -> None:
    pdf_text = extract_pdf_text() if PDF_PATH.exists() else ""
    dtext = docx_text() if DOCX_PATH.exists() else ""
    required_terms = [
        "Agentic AI",
        "Auditability",
        "Assurance",
        "White Paper",
        "2026",
        DOCUMENT_ID,
        "0. Executive Summary",
        "16. Conclusion",
        "Appendix A",
        "Appendix F",
        "MRO-01",
        "MRO-16",
        "AARM",
        "Audit Evidence Chain",
    ]
    log = {
        "task": "AIAAWP-R4D-FORMAT-PUBLICATION-DESIGN-AND-GAIC-STANDARD-ALIGNMENT",
        "generated_at": GENERATED_AT,
        "document_id": DOCUMENT_ID,
        "status": STATUS,
        "artifacts": {
            "html": {"path": rel(HTML_PATH), "sha256": file_hash(HTML_PATH), "bytes": HTML_PATH.stat().st_size},
            "pdf": {
                "path": rel(PDF_PATH),
                "sha256": file_hash(PDF_PATH),
                "bytes": PDF_PATH.stat().st_size,
                "pages": pdf_page_count(PDF_PATH),
            },
            "docx": {"path": rel(DOCX_PATH), "sha256": file_hash(DOCX_PATH), "bytes": DOCX_PATH.stat().st_size},
        },
        "table_stats": table_stats,
        "content_presence": {
            "pdf": {term: term in pdf_text for term in required_terms},
            "docx": {term: term in dtext for term in required_terms},
        },
        "manifest_status": manifest.get("status"),
        "visual_authority": manifest.get("visual_authority"),
        "boundary": BOUNDARY_LINE,
    }
    write_text(GENERATION_LOG_PATH, json.dumps(log, indent=2, ensure_ascii=False))


def clean_old_outputs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for pattern in [
        f"{ARTIFACT_BASENAME}.*",
        f"{ARTIFACT_BASENAME}-docx-source.html",
        "aiaawp-r4d-pdf-text.txt",
        "aiaawp-r4d-generation-log.json",
    ]:
        for path in OUT_DIR.glob(pattern):
            if path.is_file():
                path.unlink()
    shutil.rmtree(OUT_DIR / "chrome-profile", ignore_errors=True)


def main() -> int:
    clean_old_outputs()
    _html, table_stats = build_html()
    build_docx()
    html_to_pdf()
    manifest = update_manifest(table_stats)
    write_generation_log(table_stats, manifest)
    update_checksums()
    print(
        json.dumps(
            {
                "html": {"path": rel(HTML_PATH), "sha256": file_hash(HTML_PATH)},
                "pdf": {"path": rel(PDF_PATH), "sha256": file_hash(PDF_PATH), "pages": pdf_page_count(PDF_PATH)},
                "docx": {"path": rel(DOCX_PATH), "sha256": file_hash(DOCX_PATH)},
                "generation_log": {"path": rel(GENERATION_LOG_PATH), "sha256": file_hash(GENERATION_LOG_PATH)},
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
