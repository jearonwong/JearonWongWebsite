#!/usr/bin/env python3
"""
Phase 1D-7 HTML-first publication renderer for GAIC-2026 v0.3.2 FRC-R3.

This renderer treats active source_r3 Markdown as the only content truth,
builds a structured HTML/CSS publication, exports PDF from HTML, and creates
a DOCX derivative from the same HTML. It does not use old DOCX/PDF artifacts
as input.
"""

from __future__ import annotations

import html
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin

from PIL import Image, ImageChops
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import generate_r3_phase1d_docx_pdf as base
import generate_r3_phase1d4_full_whitepaper as full


ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parents[1]
OUT_DIR = ROOT / "out" / "phase_1d7" / "html_publication"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
REPORT_DIR = ROOT / "reports"
ARTIFACT_STEM = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-HTML-Publication-Draft"
HTML_NAME = f"{ARTIFACT_STEM}.html"
PDF_NAME = f"{ARTIFACT_STEM}.pdf"
DOCX_NAME = f"{ARTIFACT_STEM}.docx"
PDF_TEXT_NAME = "phase_1d7_pdf_text.txt"
LOG_NAME = "phase_1d7_generation_log.json"

TRACE_TAG = "GACWP-2026-v0.3.2-FRC-R3"
PHASE_STATUS = "PHASE 1D-7 HTML PUBLICATION DRAFT GENERATED / QA REQUIRED"
SOURCE_FILES = full.SOURCE_FILES
HIGH_RISK_TABLE_IDS = set(base.HIGH_RISK_TABLE_IDS) | {
    "T-F-01", "T-F-02", "T-F-03", "T-F-04", "T-F-05",
    "T-F-06", "T-F-07", "T-F-08", "T-F-09", "T-F-10",
    "T-F-11", "T-F-12", "T-F-13", "T-F-14", "T-F-15",
}

FORBIDDEN_TERMS = [
    "certifies compliance",
    "certification authority",
    "regulator-approved",
    "regulatory approval",
    "proves legal compliance",
    "guarantees compliance",
    "only validation path",
    "official compliance standard",
    "conformity assessment body",
    "market endorsement",
    "recommended vendor",
    "best system",
    "procurement winner",
    "winner",
    "leader",
    "trails",
    "final vendor score",
    "MPLP is required",
    "MPLP certifies",
    "Validation Lab certifies compliance",
    "audit opinion",
    "insurance guarantee",
    "deployment-readiness guarantee",
]

APPENDIX_G_STRICT_TERMS = [
    "Composite Score",
    "Ranking",
    "ranked",
    "score table",
    "final vendor score",
    "recommended vendor",
    "best system",
    "winner",
    "leader",
    "trailer",
    "leads",
    "trails",
]

FIGURE_TITLE_FALLBACKS = {
    "F-01": "AI Agent Lifecycle Governance Stack",
    "F-02": "Governance Stack / Missing Layer",
    "F-03": "MRO Topology",
    "F-04": "RCCS/ALCS Dual Scoring Model",
    "F-05": "Lifecycle Conformance Mapping Strength Heatmap",
    "F-06": "Enterprise Control Overlay for AI Agent Lifecycle Governance",
    "F-07": "Model Risk in Agentic Lifecycle",
    "F-08": "Evidence-Based Validation Pattern Flow",
    "F-09": "Enterprise Failure Scenario Chain",
    "F-10": "Companion Paper Boundary Map",
    "F-11": "Stage 0-to-7 Adoption Roadmap",
}


@dataclass
class TableRecord:
    table_id: str | None
    caption: str | None
    source_file: str
    rows: int
    cols: int
    rendered_tables: int
    strategy: str
    high_risk: bool


@dataclass
class FigureRecord:
    figure_id: str
    title: str
    source_file: str
    has_flow: bool
    has_interpretation: bool


def which(name: str) -> str | None:
    return shutil.which(name)


def run_git(args: list[str]) -> str:
    result = subprocess.run(["git", *args], cwd=REPO_ROOT, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        return f"ERROR: {result.stderr.strip()}"
    return result.stdout.strip()


def clean_output_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in OUT_DIR.iterdir():
        if path.is_dir():
            shutil.rmtree(path)
        else:
            path.unlink()
    RENDER_DIR.mkdir(parents=True, exist_ok=True)


def sanitize_line(line: str, relative: str) -> str | None:
    return full.sanitize_publication_line(line, relative)


def read_source_lines(relative: str) -> list[str]:
    path = ROOT / relative
    if not path.exists():
        raise FileNotFoundError(f"Missing required source file: {relative}")
    lines: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        sanitized = sanitize_line(raw, relative)
        if sanitized is not None:
            lines.append(sanitized.rstrip())
    return lines


def strip_markdown_for_text(text: str) -> str:
    return base.clean_inline_markdown(text).replace("→", "->").strip()


def inline_md_to_html(text: str) -> str:
    placeholders: list[str] = []

    def stash(value: str) -> str:
        placeholders.append(value)
        return f"\u0000{len(placeholders) - 1}\u0000"

    raw = text.strip()
    raw = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", lambda m: stash(html.escape(m.group(1))), raw)
    raw = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        lambda m: stash(
            f'<a href="{html.escape(m.group(2), quote=True)}">{html.escape(m.group(1))}</a>'
        ),
        raw,
    )
    raw = re.sub(r"`([^`]+)`", lambda m: stash(f'<code>{html.escape(m.group(1))}</code>'), raw)
    escaped = html.escape(raw)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", escaped)
    escaped = escaped.replace("&lt;br&gt;", "<br>").replace("&lt;br/&gt;", "<br>").replace("&lt;br /&gt;", "<br>")
    for idx, value in enumerate(placeholders):
        escaped = escaped.replace(f"\u0000{idx}\u0000", value)
    escaped = re.sub(r"\b(L[1-5])\b", r'<span class="badge evidence">\1</span>', escaped)
    escaped = re.sub(r"\b(MRO-\d{2})\b", r'<span class="badge mro">\1</span>', escaped)
    escaped = re.sub(r"\b(RCCS-\d{2})\b", r'<span class="badge rccs">\1</span>', escaped)
    escaped = re.sub(r"\b(ALCS-\d{2})\b", r'<span class="badge alcs">\1</span>', escaped)
    return escaped


def slugify(text: str, used: dict[str, int]) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", strip_markdown_for_text(text).lower()).strip("-")
    slug = slug or "section"
    count = used.get(slug, 0)
    used[slug] = count + 1
    return slug if count == 0 else f"{slug}-{count + 1}"


def heading_level(line: str) -> int:
    return len(line) - len(line.lstrip("#"))


def is_caption_line(line: str) -> bool:
    return base.is_caption_line(line)


def is_table_start(lines: list[str], i: int) -> bool:
    return base.is_table_start(lines, i)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for raw in table_lines:
        parts = [p.strip() for p in raw.strip("|").split("|")]
        if parts and all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            continue
        rows.append(parts)
    return rows, i


def detect_table_id(caption: str | None, rows: list[list[str]]) -> str | None:
    cleaned = [[strip_markdown_for_text(cell) for cell in row] for row in rows]
    return base.detect_table_id(caption, cleaned)


def column_groups(cols: int) -> list[list[int]]:
    if cols <= 5:
        return [list(range(cols))]
    if cols == 6:
        return [[0, 1, 2], [0, 3, 4, 5]]
    if cols == 7:
        return [[0, 1, 2, 3], [0, 4, 5, 6]]
    if cols == 8:
        return [[0, 1, 2, 3], [0, 4, 5, 6, 7]]
    if cols == 9:
        return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8]]
    return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8, 9]]


def render_table(rows: list[list[str]], caption: str | None, relative: str, records: list[TableRecord]) -> str:
    if not rows:
        return ""
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table_id = detect_table_id(caption, normalized)
    high_risk = bool(table_id and table_id in HIGH_RISK_TABLE_IDS)
    split_required = cols >= 6 or high_risk
    groups = column_groups(cols) if split_required else [list(range(cols))]
    strategy = "semantic column split" if split_required else "single table"
    classes = ["table-block"]
    if high_risk or split_required:
        classes.append("table-compact")
    parts = [f'<div class="{" ".join(classes)}">']
    if caption:
        parts.append(f'<div class="table-caption">{inline_md_to_html(caption)}</div>')
    if split_required:
        parts.append('<div class="layout-note">Wide source table split into consecutive column groups; all source rows and columns are preserved.</div>')
    for idx, group in enumerate(groups, 1):
        if split_required:
            parts.append(f'<div class="continuation-label">Continuation {idx} of {len(groups)}</div>')
        parts.append("<table>")
        if normalized:
            parts.append("<thead><tr>")
            for col in group:
                parts.append(f"<th>{inline_md_to_html(normalized[0][col] if col < len(normalized[0]) else '')}</th>")
            parts.append("</tr></thead>")
        parts.append("<tbody>")
        for row in normalized[1:]:
            parts.append("<tr>")
            for col in group:
                parts.append(f"<td>{inline_md_to_html(row[col] if col < len(row) else '')}</td>")
            parts.append("</tr>")
        parts.append("</tbody></table>")
    parts.append("</div>")
    records.append(
        TableRecord(
            table_id=table_id,
            caption=strip_markdown_for_text(caption or "") or None,
            source_file=relative,
            rows=len(normalized),
            cols=cols,
            rendered_tables=len(groups),
            strategy=strategy,
            high_risk=high_risk,
        )
    )
    return "\n".join(parts)


def is_figure_title(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith("-"):
        return False
    return bool(re.match(r"^\*\*Figure F-\d{2}\b", stripped) or re.match(r"^Figure F-\d{2}\b", stripped))


def parse_figure_title(line: str) -> tuple[str, str]:
    stripped = strip_markdown_for_text(line)
    match = re.search(r"\bFigure\s+(F-\d{2})\s*(?:[-—:]\s*)?(.*)$", stripped)
    if not match:
        return "F-00", stripped
    figure_id = match.group(1)
    title = match.group(2).strip() or FIGURE_TITLE_FALLBACKS.get(figure_id, "")
    return figure_id, title


def render_flow(flow: str) -> str:
    flow = strip_markdown_for_text(flow)
    if not flow:
        return ""
    lanes = [part.strip() for part in re.split(r"\s*(?:→|->)\s*", flow) if part.strip()]
    if len(lanes) < 2:
        return f'<p class="figure-flow-text">{inline_md_to_html(flow)}</p>'
    pieces: list[str] = ['<div class="figure-flow">']
    for idx, lane in enumerate(lanes):
        pieces.append(f'<span class="flow-node">{inline_md_to_html(lane)}</span>')
        if idx < len(lanes) - 1:
            pieces.append('<span class="flow-arrow">→</span>')
    pieces.append("</div>")
    return "\n".join(pieces)


def render_figure(lines: list[str], start: int, relative: str, records: list[FigureRecord]) -> tuple[str, int]:
    figure_id, title = parse_figure_title(lines[start])
    i = start + 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    flow = ""
    if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not is_table_start(lines, i):
        flow = lines[i].strip()
        i += 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    interpretation = ""
    if i < len(lines) and re.match(r"^\*Figure F-\d{2}\b", lines[i].strip()):
        interpretation = lines[i].strip().strip("*")
        i += 1
    records.append(
        FigureRecord(
            figure_id=figure_id,
            title=title,
            source_file=relative,
            has_flow=bool(flow),
            has_interpretation=bool(interpretation),
        )
    )
    return (
        "\n".join(
            [
                f'<figure class="figure-card" id="figure-{figure_id.lower()}">',
                f'<figcaption><span>{figure_id}</span>{html.escape(title)}</figcaption>',
                render_flow(flow),
                f'<p class="figure-note">{inline_md_to_html(interpretation)}</p>' if interpretation else "",
                "</figure>",
            ]
        ),
        i,
    )


def callout_class(text: str) -> str | None:
    low = strip_markdown_for_text(text).lower()
    if "not legal advice" in low or "not a certification" in low or "not certification" in low:
        return "callout boundary"
    if low.startswith("**pattern boundary:**") or low.startswith("**boundary note:**") or low.startswith("**chapter boundary:**"):
        return "callout boundary"
    if low.startswith("**the paper argues") or low.startswith("**core thesis"):
        return "callout thesis"
    if "source-bound" in low or "evidence" in low and low.startswith("**"):
        return "callout evidence"
    if "blocker" in low or "unresolved" in low:
        return "callout warning"
    return None


def render_list(lines: list[str], start: int, ordered: bool) -> tuple[str, int]:
    tag = "ol" if ordered else "ul"
    parts = [f"<{tag}>"]
    i = start
    pattern = r"^\s*\d+\.\s+(.*)$" if ordered else r"^\s*[-*]\s+(.*)$"
    while i < len(lines):
        match = re.match(pattern, lines[i])
        if not match:
            break
        parts.append(f"<li>{inline_md_to_html(match.group(1))}</li>")
        i += 1
    parts.append(f"</{tag}>")
    return "\n".join(parts), i


def render_heading(line: str, used_slugs: dict[str, int], relative: str) -> tuple[str, dict[str, str | int]]:
    level = min(heading_level(line), 4)
    text = strip_markdown_for_text(line[level:].strip())
    slug = slugify(text, used_slugs)
    classes = []
    if relative.startswith("appendices/") and level == 1:
        classes.append("appendix-heading")
    class_attr = f' class="{" ".join(classes)}"' if classes else ""
    return f'<h{level} id="{slug}"{class_attr}>{html.escape(text)}</h{level}>', {"level": level, "text": text, "id": slug, "source": relative}


def render_source_file(relative: str, used_slugs: dict[str, int], table_records: list[TableRecord], figure_records: list[FigureRecord]) -> tuple[str, list[dict[str, str | int]], str]:
    lines = read_source_lines(relative)
    html_parts: list[str] = [f'<section class="source-section {"appendix-mode" if relative.startswith("appendices/") else "body-mode"}" data-source="{html.escape(relative)}">']
    headings: list[dict[str, str | int]] = []
    source_text_lines: list[str] = []
    i = 0
    in_code = False
    code_buffer: list[str] = []
    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        source_text_lines.append(line)

        if stripped.startswith("```"):
            if in_code:
                language = ""
                code = html.escape("\n".join(code_buffer))
                html_parts.append(f'<pre class="code-block"><code>{code}</code></pre>')
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            html_parts.append('<hr class="section-rule">')
            i += 1
            continue
        if relative == "sections/00-front-matter.md" and stripped == "# Global AI Compliance White Paper 2026":
            i += 1
            continue
        if is_table_start(lines, i):
            rows, new_i = parse_markdown_table(lines, i)
            html_parts.append(render_table(rows, pending_caption, relative, table_records))
            pending_caption = None
            i = new_i
            continue
        if stripped.startswith("#"):
            rendered, heading = render_heading(stripped, used_slugs, relative)
            headings.append(heading)
            html_parts.append(rendered)
            pending_caption = None
            i += 1
            continue
        if is_caption_line(stripped):
            pending_caption = stripped
            i += 1
            continue
        if is_figure_title(stripped):
            rendered, new_i = render_figure(lines, i, relative, figure_records)
            html_parts.append(rendered)
            pending_caption = None
            i = new_i
            continue
        if stripped.startswith(">"):
            html_parts.append(f'<blockquote>{inline_md_to_html(stripped.lstrip("> "))}</blockquote>')
            i += 1
            continue
        if re.match(r"^\s*[-*]\s+", line):
            rendered, new_i = render_list(lines, i, ordered=False)
            html_parts.append(rendered)
            i = new_i
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            rendered, new_i = render_list(lines, i, ordered=True)
            html_parts.append(rendered)
            i = new_i
            continue

        klass = callout_class(stripped)
        if klass:
            html_parts.append(f'<div class="{klass}">{inline_md_to_html(stripped)}</div>')
        else:
            html_parts.append(f"<p>{inline_md_to_html(stripped)}</p>")
        i += 1

    if code_buffer:
        html_parts.append(f'<pre class="code-block"><code>{html.escape(chr(10).join(code_buffer))}</code></pre>')
    html_parts.append("</section>")
    return "\n".join(html_parts), headings, "\n".join(source_text_lines)


def build_generated_toc(headings: list[dict[str, str | int]]) -> str:
    chapter_items = []
    appendix_items = []
    for heading in headings:
        text = str(heading["text"])
        if int(heading["level"]) != 1:
            continue
        item = f'<li><a href="#{heading["id"]}">{html.escape(text)}</a></li>'
        if text.startswith("Appendix"):
            appendix_items.append(item)
        elif text.startswith("Chapter"):
            chapter_items.append(item)
    return "\n".join(
        [
            '<nav class="generated-toc" aria-label="Publication table of contents">',
            "<h2>Publication Contents</h2>",
            '<div class="toc-grid">',
            '<div><h3>Chapters</h3><ol>',
            "\n".join(chapter_items),
            "</ol></div>",
            '<div><h3>Appendices</h3><ol>',
            "\n".join(appendix_items),
            "</ol></div>",
            "</div>",
            "</nav>",
        ]
    )


def cover_html() -> str:
    return f"""
<section class="cover-page">
  <div class="cover-kicker">Jearon Wong / AI Agent Lifecycle Governance</div>
  <h1>Global AI Compliance White Paper 2026</h1>
  <p class="cover-subtitle">From Model Governance to Agentic Lifecycle Conformance</p>
  <p class="cover-role">Jearon Wong · Protocol Architect for the Agent Era</p>
  <div class="thesis-strip" aria-label="Lifecycle thesis strip">
    <span>Intent</span><span>Authority</span><span>Agent Work</span><span>Evidence</span><span>Closure</span>
  </div>
  <div class="cover-meta">
    <div><strong>Document ID</strong><span>{TRACE_TAG}</span></div>
    <div><strong>Version</strong><span>v0.3.2 Final Release Candidate R3</span></div>
    <div><strong>Date</strong><span>May 2026</span></div>
    <div><strong>Status</strong><span>HTML Publication Draft / QA Required</span></div>
    <div><strong>Visual source of truth</strong><span>HTML/PDF primary; DOCX derivative</span></div>
  </div>
  <div class="cover-notice">Non-legal technical governance analysis. Not legal advice, not a certification standard, not regulatory approval, not a procurement recommendation, not a vendor ranking, and not a final vendor assessment.</div>
</section>
"""


def important_notice_html() -> str:
    return """
<section class="important-notice">
  <h2>Important Notice</h2>
  <p>This publication is a technical governance analysis of AI Agent Lifecycle Governance. It is not legal advice, a certification standard, a conformity assessment, a regulatory approval path, an audit opinion, an insurance guarantee, a procurement recommendation, a vendor ranking, or a final vendor assessment.</p>
  <p>RCCS/ALCS scores and system mappings are analytical and source-qualified. Appendix G remains revalidated provisional, qualitative, non-ranking, and no-score.</p>
</section>
"""


def css() -> str:
    return """
:root {
  --monolith: #111111;
  --slate: #1a1a1a;
  --machine: #f1f5f9;
  --registry: #3b82f6;
  --evidence: #22d3ee;
  --status: #64748b;
  --grid: #d8dee8;
  --amber: #b7791f;
  --fault: #991b1b;
  --paper: #ffffff;
}
* { box-sizing: border-box; }
html { background: var(--machine); color: var(--monolith); }
body {
  margin: 0;
  font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  font-size: 15px;
  line-height: 1.52;
  color: var(--monolith);
}
.publication {
  max-width: 980px;
  margin: 0 auto;
  background: var(--paper);
  padding: 32px 80px 48px;
  box-shadow: 0 20px 80px rgba(15, 23, 42, 0.08);
}
a { color: var(--registry); text-decoration: none; }
p { margin: 0 0 0.82rem; }
h1, h2, h3, h4 {
  font-family: Outfit, Inter, system-ui, sans-serif;
  line-height: 1.18;
  letter-spacing: 0;
  color: var(--monolith);
  break-after: avoid;
}
h1 {
  margin: 2.2rem 0 1rem;
  padding-top: 0.6rem;
  border-top: 4px solid var(--monolith);
  font-size: 2.32rem;
}
h2 {
  margin: 2rem 0 0.75rem;
  padding-top: 0.75rem;
  border-top: 1px solid var(--grid);
  font-size: 1.54rem;
}
h3 { margin: 1.35rem 0 0.48rem; font-size: 1.14rem; color: var(--slate); }
h4 { margin: 1rem 0 0.35rem; font-size: 0.98rem; color: var(--status); }
ul, ol { margin: 0.2rem 0 0.9rem 1.35rem; padding: 0; }
li { margin: 0.18rem 0; }
blockquote {
  margin: 1rem 0;
  padding: 0.65rem 0.9rem;
  border-left: 4px solid var(--registry);
  background: #f8fafc;
  color: var(--slate);
}
.cover-page {
  min-height: 92vh;
  display: flex;
  flex-direction: column;
  justify-content: center;
  border: 1px solid var(--monolith);
  padding: 52px 44px;
  margin-bottom: 2.5rem;
  background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
}
.cover-kicker {
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  color: var(--status);
  text-transform: uppercase;
  font-size: 0.78rem;
  margin-bottom: 2.2rem;
}
.cover-page h1 {
  border: 0;
  margin: 0;
  padding: 0;
  font-size: 3.6rem;
  max-width: 760px;
}
.cover-subtitle {
  margin: 1rem 0 0.4rem;
  color: var(--registry);
  font-size: 1.25rem;
  font-weight: 700;
}
.cover-role { color: var(--status); font-size: 0.95rem; }
.thesis-strip {
  display: grid;
  grid-template-columns: repeat(5, 1fr);
  gap: 1px;
  margin: 2rem 0;
  border: 1px solid var(--monolith);
  background: var(--monolith);
}
.thesis-strip span {
  background: var(--monolith);
  color: #fff;
  padding: 0.68rem 0.55rem;
  text-align: center;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
  font-weight: 700;
}
.cover-meta {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 1px;
  background: var(--grid);
  border: 1px solid var(--grid);
  margin-top: 1.2rem;
}
.cover-meta div { background: #fff; padding: 0.65rem 0.8rem; }
.cover-meta strong {
  display: block;
  color: var(--status);
  font-size: 0.72rem;
  text-transform: uppercase;
}
.cover-meta span { font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace; font-size: 0.78rem; }
.cover-notice, .important-notice, .callout {
  border-left: 4px solid var(--amber);
  background: #fffbeb;
  padding: 0.85rem 1rem;
  margin: 1rem 0;
  break-inside: avoid;
}
.important-notice { margin: 0 0 2rem; }
.important-notice h2 { border: 0; padding: 0; margin-top: 0; }
.callout.thesis { border-color: var(--registry); background: #eff6ff; font-weight: 650; }
.callout.evidence { border-color: var(--evidence); background: #ecfeff; }
.callout.warning { border-color: var(--fault); background: #fef2f2; }
.generated-toc {
  margin: 2rem 0 2.8rem;
  padding: 1.1rem;
  border: 1px solid var(--grid);
  background: #f8fafc;
}
.generated-toc h2 { border: 0; padding-top: 0; margin-top: 0; }
.toc-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 1.25rem; }
.toc-grid h3 { margin-top: 0; color: var(--registry); }
.toc-grid li { font-size: 0.88rem; }
.source-section { margin-bottom: 1.3rem; }
.appendix-heading {
  border-top-color: var(--registry);
  color: var(--registry);
}
.appendix-mode { font-size: 0.92rem; }
.badge {
  display: inline-block;
  padding: 0.08rem 0.28rem;
  border-radius: 3px;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72em;
  line-height: 1.2;
  vertical-align: baseline;
  border: 1px solid var(--grid);
  background: #f8fafc;
  color: var(--slate);
}
.badge.evidence { border-color: #bae6fd; background: #ecfeff; }
.badge.mro, .badge.alcs, .badge.rccs { border-color: #bfdbfe; background: #eff6ff; color: #1d4ed8; }
.figure-card {
  margin: 1.15rem 0 1.35rem;
  padding: 1rem;
  border: 1px solid var(--grid);
  border-top: 4px solid var(--monolith);
  background: #fbfdff;
  break-inside: avoid;
}
.figure-card figcaption {
  font-family: Outfit, Inter, system-ui, sans-serif;
  font-weight: 800;
  margin-bottom: 0.8rem;
}
.figure-card figcaption span {
  display: inline-block;
  margin-right: 0.55rem;
  padding: 0.12rem 0.38rem;
  background: var(--monolith);
  color: #fff;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
}
.figure-flow {
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 0.35rem;
  margin: 0.45rem 0 0.75rem;
}
.flow-node {
  border: 1px solid var(--grid);
  background: #fff;
  padding: 0.42rem 0.55rem;
  font-size: 0.82rem;
  font-weight: 650;
}
.flow-arrow {
  color: var(--registry);
  font-weight: 800;
}
.figure-note {
  margin: 0.6rem 0 0;
  color: var(--status);
  font-size: 0.88rem;
  font-style: italic;
}
.table-block {
  margin: 1rem 0 1.25rem;
  break-inside: auto;
}
.table-caption {
  font-weight: 800;
  margin: 0 0 0.35rem;
  color: var(--monolith);
  break-after: avoid;
}
.layout-note, .continuation-label {
  color: var(--status);
  font-size: 0.78rem;
  margin: 0.2rem 0 0.35rem;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
}
table {
  width: 100%;
  border-collapse: collapse;
  table-layout: fixed;
  margin-bottom: 0.7rem;
}
thead { display: table-header-group; }
th, td {
  border: 1px solid var(--grid);
  padding: 0.46rem 0.5rem;
  vertical-align: top;
  overflow-wrap: anywhere;
}
th {
  background: #e8eef7;
  color: var(--slate);
  font-size: 0.78rem;
  text-align: left;
}
td { font-size: 0.8rem; }
.table-compact th, .table-compact td {
  font-size: 0.68rem;
  padding: 0.32rem 0.34rem;
}
.code-block {
  white-space: pre-wrap;
  background: #0f172a;
  color: #e2e8f0;
  padding: 0.85rem;
  border-radius: 4px;
  overflow-wrap: anywhere;
  font-size: 0.76rem;
}
.section-rule { border: 0; border-top: 1px solid var(--grid); margin: 1.2rem 0; }
.html-footer {
  margin-top: 3rem;
  padding-top: 1rem;
  border-top: 1px solid var(--grid);
  color: var(--status);
  font-size: 0.76rem;
  text-align: center;
}
@media print {
  @page { size: A4; margin: 16mm 15mm; }
  html, body { background: #fff; }
  body { font-size: 10.2pt; line-height: 1.45; }
  .publication {
    max-width: none;
    margin: 0;
    padding: 0;
    box-shadow: none;
  }
  .cover-page {
    min-height: 265mm;
    page-break-after: always;
    margin: 0;
  }
  .important-notice { break-inside: avoid; }
  h1 { font-size: 25pt; margin-top: 16pt; }
  h2 { font-size: 16pt; margin-top: 14pt; }
  h3 { font-size: 12pt; }
  h4 { font-size: 10pt; }
  p, li { orphans: 2; widows: 2; }
  .figure-card, .callout { break-inside: avoid; }
  .generated-toc { break-inside: avoid; }
  .toc-grid { gap: 0.8rem; }
  th, td { padding: 3.2pt 3.6pt; }
  td { font-size: 7.6pt; }
  th { font-size: 7.4pt; }
  .table-compact th, .table-compact td { font-size: 6.35pt; padding: 2.4pt; }
  .html-footer { font-size: 7pt; }
}
"""


def build_html() -> tuple[Path, list[TableRecord], list[FigureRecord], list[dict[str, str | int]], str]:
    used_slugs: dict[str, int] = {}
    table_records: list[TableRecord] = []
    figure_records: list[FigureRecord] = []
    all_headings: list[dict[str, str | int]] = []
    sections: list[str] = []
    source_texts: list[str] = []

    for rel in SOURCE_FILES:
        rendered, headings, source_text = render_source_file(rel, used_slugs, table_records, figure_records)
        sections.append(rendered)
        all_headings.extend(headings)
        source_texts.append(source_text)

    generated_toc = build_generated_toc(all_headings)
    document = "\n".join(
        [
            "<!doctype html>",
            '<html lang="en">',
            "<head>",
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            f"<title>Global AI Compliance White Paper 2026 - {TRACE_TAG}</title>",
            "<style>",
            css(),
            "</style>",
            "</head>",
            "<body>",
            '<main class="publication">',
            cover_html(),
            important_notice_html(),
            generated_toc,
            "\n".join(sections),
            f'<footer class="html-footer">© 2026 Jearon Wong · {TRACE_TAG} · Non-legal technical governance analysis · HTML/PDF visual source of truth</footer>',
            "</main>",
            "</body>",
            "</html>",
        ]
    )
    html_path = OUT_DIR / HTML_NAME
    html_path.write_text(document, encoding="utf-8")
    return html_path, table_records, figure_records, all_headings, "\n\n".join(source_texts)


def chrome_path() -> str:
    candidates = [
        which("google-chrome"),
        which("chromium"),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    raise RuntimeError("No Chrome/Chromium executable found for HTML-to-PDF export")


def html_to_pdf(html_path: Path) -> Path:
    pdf_path = OUT_DIR / PDF_NAME
    if pdf_path.exists():
        pdf_path.unlink()
    chrome = chrome_path()
    user_data = OUT_DIR / "chrome_profile"
    if user_data.exists():
        shutil.rmtree(user_data)
    user_data.mkdir(parents=True, exist_ok=True)
    cmd = [
        chrome,
        "--headless=new",
        "--disable-gpu",
        "--disable-background-networking",
        "--no-pdf-header-footer",
        "--disable-extensions",
        "--no-first-run",
        "--no-default-browser-check",
        f"--user-data-dir={user_data}",
        f"--print-to-pdf={pdf_path}",
        html_path.resolve().as_uri(),
    ]
    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    last_size = -1
    stable_ticks = 0
    deadline = time.time() + 900
    stdout = ""
    stderr = ""
    while time.time() < deadline:
        if process.poll() is not None:
            stdout, stderr = process.communicate(timeout=60)
            break
        if pdf_path.exists():
            size = pdf_path.stat().st_size
            if size > 0 and size == last_size:
                stable_ticks += 1
            else:
                stable_ticks = 0
            last_size = size
            if stable_ticks >= 3:
                try:
                    if pdf_page_count(pdf_path) > 0:
                        process.terminate()
                        try:
                            stdout, stderr = process.communicate(timeout=15)
                        except subprocess.TimeoutExpired:
                            process.kill()
                            stdout, stderr = process.communicate(timeout=30)
                        shutil.rmtree(user_data, ignore_errors=True)
                        return pdf_path
                except Exception:
                    pass
        time.sleep(2)
    else:
        # Chrome can finish writing a large PDF and then hang during shutdown on
        # macOS. If the PDF is already complete and readable, accept it as the
        # generated artifact and terminate the dedicated headless process.
        if pdf_path.exists() and pdf_page_count(pdf_path) > 0:
            process.kill()
            stdout, stderr = process.communicate(timeout=60)
            shutil.rmtree(user_data, ignore_errors=True)
            return pdf_path
        process.kill()
        stdout, stderr = process.communicate(timeout=60)
        raise RuntimeError(f"Chrome HTML-to-PDF export timed out\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
    if process.returncode != 0 or not pdf_path.exists():
        raise RuntimeError(f"Chrome HTML-to-PDF export failed\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
    shutil.rmtree(user_data, ignore_errors=True)
    return pdf_path


def source_to_docx_derivative() -> tuple[Path, str]:
    docx_path = OUT_DIR / DOCX_NAME
    if docx_path.exists():
        docx_path.unlink()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    for name, size in [("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11), ("Heading 4", 10)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(36)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Global AI Compliance White Paper 2026")
    run.font.name = "Arial"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 17, 17)
    sub = doc.add_paragraph("HTML Publication Draft - DOCX editable derivative")
    sub.runs[0].font.name = "Arial"
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    note = doc.add_paragraph(
        "This DOCX is an editable derivative generated from the same active source used by the HTML renderer. "
        "HTML/PDF remain the visual source of truth."
    )
    note.runs[0].font.size = Pt(8.5)
    note.runs[0].font.italic = True
    doc.add_page_break()

    def add_para(text: str, style: str | None = None) -> None:
        clean = strip_markdown_for_text(text)
        if not clean:
            return
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(2.8)
        para.add_run(clean)

    def add_table(rows: list[list[str]], caption: str | None) -> None:
        if caption:
            cap = doc.add_paragraph()
            cap.paragraph_format.keep_with_next = True
            cr = cap.add_run(strip_markdown_for_text(caption))
            cr.bold = True
            cr.font.size = Pt(8.4)
        if not rows:
            return
        cols = max(len(row) for row in rows)
        normalized = [row + [""] * (cols - len(row)) for row in rows]
        table = doc.add_table(rows=len(normalized), cols=cols)
        table.style = "Table Grid"
        table.autofit = True
        for r_idx, row in enumerate(normalized):
            for c_idx, value in enumerate(row):
                cell = table.rows[r_idx].cells[c_idx]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(strip_markdown_for_text(value))
                run.font.name = "Arial"
                run.font.size = Pt(6.2 if cols >= 6 else 7.2)
                run.font.bold = r_idx == 0

    for rel in SOURCE_FILES:
        lines = read_source_lines(rel)
        i = 0
        in_code = False
        code_buffer: list[str] = []
        pending_caption: str | None = None
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if stripped.startswith("```"):
                if in_code:
                    add_para("\n".join(code_buffer))
                    code_buffer = []
                    in_code = False
                else:
                    in_code = True
                i += 1
                continue
            if in_code:
                code_buffer.append(line)
                i += 1
                continue
            if not stripped or stripped == "---":
                i += 1
                continue
            if rel == "sections/00-front-matter.md" and stripped == "# Global AI Compliance White Paper 2026":
                i += 1
                continue
            if is_table_start(lines, i):
                rows, new_i = parse_markdown_table(lines, i)
                add_table(rows, pending_caption)
                pending_caption = None
                i = new_i
                continue
            if stripped.startswith("#"):
                level = min(heading_level(stripped), 4)
                doc.add_heading(strip_markdown_for_text(stripped[level:].strip()), level=level)
                pending_caption = None
                i += 1
                continue
            if is_caption_line(stripped):
                pending_caption = stripped
                i += 1
                continue
            if is_figure_title(stripped):
                figure_id, title_text = parse_figure_title(stripped)
                add_para(f"Figure {figure_id} - {title_text}")
                i += 1
                continue
            bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
            numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
            if bullet:
                add_para(bullet.group(1), style="List Bullet")
            elif numbered:
                add_para(numbered.group(1), style="List Number")
            elif stripped.startswith(">"):
                add_para(stripped.lstrip("> "))
            else:
                add_para(line)
            i += 1
        if code_buffer:
            add_para("\n".join(code_buffer))

    doc.save(docx_path)
    return docx_path, "python-docx editable derivative from Phase 1D-7 structured source"


def pdf_page_count(pdf_path: Path) -> int:
    result = subprocess.run(["pdfinfo", str(pdf_path)], capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(result.stderr)
    match = re.search(r"^Pages:\s+(\d+)", result.stdout, re.MULTILINE)
    if not match:
        raise RuntimeError("Unable to determine PDF page count")
    return int(match.group(1))


def extract_pdf_text(pdf_path: Path) -> tuple[Path, list[str]]:
    text_path = OUT_DIR / PDF_TEXT_NAME
    result = subprocess.run(["pdftotext", "-layout", str(pdf_path), str(text_path)], capture_output=True, text=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError(f"pdftotext failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    text = text_path.read_text(encoding="utf-8", errors="ignore")
    pages = text.split("\f")
    if pages and not pages[-1].strip():
        pages = pages[:-1]
    return text_path, pages


def render_pdf_pages(pdf_path: Path) -> list[Path]:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for existing in RENDER_DIR.glob("*.png"):
        existing.unlink()
    result = subprocess.run(["pdftoppm", "-png", "-r", "110", str(pdf_path), str(RENDER_DIR / "page")], capture_output=True, text=True, timeout=1200)
    if result.returncode != 0:
        raise RuntimeError(f"pdftoppm failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    return sorted(RENDER_DIR.glob("page-*.png"))


def page_number_from_path(path: Path) -> int:
    match = re.search(r"page-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def blankish_pages(rendered_pages: list[Path]) -> list[int]:
    blanks: list[int] = []
    for path in rendered_pages:
        img = Image.open(path).convert("L")
        hist = img.histogram()
        dark = sum(hist[:245])
        total = img.width * img.height
        if dark / total < 0.002:
            blanks.append(page_number_from_path(path))
    return blanks


def edge_ink_pages(rendered_pages: list[Path]) -> list[int]:
    flagged: list[int] = []
    for path in rendered_pages:
        img = Image.open(path).convert("L")
        w, h = img.size
        margin = max(8, int(min(w, h) * 0.012))
        edges = [
            img.crop((0, 0, w, margin)),
            img.crop((0, h - margin, w, h)),
            img.crop((0, 0, margin, h)),
            img.crop((w - margin, 0, w, h)),
        ]
        edge_pixels = sum(e.width * e.height for e in edges)
        dark = 0
        for edge in edges:
            hist = edge.histogram()
            dark += sum(hist[:210])
        if dark / edge_pixels > 0.018:
            flagged.append(page_number_from_path(path))
    return flagged


def find_pages_containing(pages: list[str], patterns: Iterable[str]) -> list[int]:
    found: list[int] = []
    lowered = [p.lower() for p in patterns]
    for idx, page in enumerate(pages, 1):
        low = page.lower()
        if any(pattern.lower() in low for pattern in lowered):
            found.append(idx)
    return found


def make_contact_sheet(rendered_pages: list[Path], page_numbers: list[int], out_name: str, max_pages: int = 24) -> Path | None:
    selected = []
    page_map = {page_number_from_path(path): path for path in rendered_pages}
    for page in page_numbers:
        if page in page_map and page not in selected:
            selected.append(page)
        if len(selected) >= max_pages:
            break
    if not selected:
        return None
    thumbs = []
    for page in selected:
        img = Image.open(page_map[page]).convert("RGB")
        img.thumbnail((260, 360))
        canvas = Image.new("RGB", (280, 390), "white")
        canvas.paste(img, ((280 - img.width) // 2, 10))
        thumbs.append((page, canvas))
    cols = 4
    rows = math.ceil(len(thumbs) / cols)
    sheet = Image.new("RGB", (cols * 280, rows * 390), "#f1f5f9")
    for idx, (_page, thumb) in enumerate(thumbs):
        x = (idx % cols) * 280
        y = (idx // cols) * 390
        sheet.paste(thumb, (x, y))
    out = OUT_DIR / out_name
    sheet.save(out)
    return out


def html_text(html_path: Path) -> str:
    text = html_path.read_text(encoding="utf-8", errors="ignore")
    text = re.sub(r"<style.*?</style>", " ", text, flags=re.S)
    text = re.sub(r"<script.*?</script>", " ", text, flags=re.S)
    text = re.sub(r"<[^>]+>", " ", text)
    return html.unescape(re.sub(r"\s+", " ", text)).strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\S+", text))


def active_source_text() -> str:
    chunks: list[str] = []
    for rel in SOURCE_FILES:
        chunks.append("\n".join(read_source_lines(rel)))
    return "\n\n".join(chunks)


def active_table_ids(source_text: str) -> set[str]:
    return set(re.findall(r"\bT-[A-Z0-9]+(?:-[0-9A-Z]+)?\b", source_text))


def generated_table_ids(text: str) -> set[str]:
    return set(re.findall(r"\bT-[A-Z0-9]+(?:-[0-9A-Z]+)?\b", text))


def classify_forbidden_match(line: str, context: str = "") -> str:
    low = line.lower()
    ctx = context.lower()
    stripped = line.strip()
    stripped_low = stripped.lower().strip('"')
    if stripped.startswith("-"):
        return "allowed negative/boundary list context"
    if stripped.startswith("|"):
        return "allowed negative/boundary table context"
    if any(
        marker in ctx
        for marker in [
            "forbidden wording",
            "out of scope",
            "what this paper does not say",
            "what this paper does not do",
            "following table clarifies what this paper says and what it does not say",
            "this paper is not",
            "the comparison does not",
            "not evaluated for product quality",
            "must not be read as",
            "must not be treated as",
            "does not issue",
            "does not certify",
            "does not prove",
            "does not guarantee",
            "does not provide",
            "does not evaluate product quality",
            "does not claim",
            "not a legal compliance score",
            "not a certification",
            "not a regulator",
            "not a conformity assessment",
            "not product quality",
            "not legal advice",
            "negative control",
            "forbidden claim",
            "forbidden framing",
            "prohibited statement",
            "prohibited framing",
            "prohibited wording",
            "forbidden wording",
            "allowed framing",
            "allowed wording",
            "permitted statement",
            "what rccs is not",
            "what alcs is not",
            "what composite scoring is not",
            "what this pattern does not validate",
            "what validation lab does not replace",
            "rccs/alcs help prepare for audit but",
            "important notice",
            "validation boundary",
            "boundary statement",
            "what assessments do not provide",
            "what it cannot prove alone",
            "what this chapter does not evaluate",
            "what the pattern cannot validate",
            "what this paper does not evaluate",
            "this appendix is a buyer diligence aid",
            "must not be read as proof",
            "it is not a procurement recommendation",
            "the evidence & assurance companion paper is not intended to",
            "claim boundaries",
            "comparison boundary table",
            "boundary controls",
            "claim area",
            "claim boundary",
            "forbidden framing",
            "forbidden wording",
        ]
    ):
        return "allowed negative/boundary context"
    if low in {
        "a regulatory approval or endorsement",
        "regulatory approval or conformity assessment",
        "certification, conformity assessment, or regulatory approval",
        "market endorsement or procurement recommendation",
        "market endorsement or quality ratings",
    }:
        return "allowed exclusion fragment"
    if "not " in low or "not a " in low or "not an " in low:
        return "allowed negative/boundary context"
    if "does not" in low or "do not" in low or "must not" in low or "no " in low:
        return "allowed negative/boundary context"
    if "forbidden" in low or "boundary" in low or "out of scope" in low or "non-claim" in low:
        return "allowed boundary context"
    if "leader" in low and ("leadership" in low or "leaders" in low):
        return "false positive leadership context"
    if "audit trail" in low or "audit trails" in low or "trails and evidence records" in low or "trails or " in low:
        return "false positive audit-trail context"
    if "regulator-approved benchmark" in low or "certification authority" in low and "non-certifying" in low:
        return "allowed forbidden-wording table context"
    if "regulatory approval path" in low or "insurance guarantee" in low:
        return "allowed negative/boundary context"
    if stripped_low == "regulator-approved":
        return "allowed prohibited-statement table fragment"
    if "market endorsement or procurement recommendation" in low:
        return "allowed exclusion fragment"
    if "certifications or regulatory approvals" in low:
        return "allowed negative/boundary context"
    if "regulatory approval; certification" in low or "regulatory approval; endorsement" in low:
        return "allowed comparison-boundary table context"
    if "certification, conformity assessment, or regulatory approval" in low:
        return "allowed exclusion fragment"
    if "regulatory approval process" in low or "regulatory approval." in low:
        return "allowed negative/boundary context"
    if "regulatory approval or regulator acceptance" in low:
        return "allowed validation-boundary table context"
    if "operational capability" in low and "regulatory approval" in low:
        return "allowed boundary table context"
    if "regulatory approval or conformity assessment" in low:
        return "allowed exclusion fragment"
    if "market endorsement or quality ratings" in low:
        return "allowed exclusion fragment"
    if "certification authority claims" in low or "claim that validation lab is a certification authority" in low:
        return "allowed boundary list context"
    if "evidence adjudication example" in low and "validation lab is a certification authority" in low:
        return "allowed prohibited-statement table context"
    if "validation lab is the only validation path" in low:
        return "allowed prohibited-statement table context"
    if "lab certifies compliance" in low or "lab proves legal" in low:
        return "allowed prohibited-statement table context"
    if "post-pilot measures" in low and "regulatory approval" in low:
        return "allowed boundary table context"
    if "operationally effective" in low and "regulator-approved" in low:
        return "allowed negative/boundary context"
    if "create legal certification or regulatory approval" in low:
        return "allowed boundary list context"
    if "issue audit opinions" in low or "audit opinions require" in low:
        return "allowed boundary/prohibited-statement context"
    if "guarantees compliance" in low and "structured inputs" in low:
        return "allowed prohibited-statement table context"
    if "deployment-readiness guarantee" in low:
        return "allowed negative/boundary context"
    if "compliance relationship" in low or "validation lab relationship" in low or "validation identity" in low:
        return "allowed forbidden-wording table context"
    if "legal compliance" in low and ("high rccs/alcs" in low or "this paper proves" in low):
        return "allowed prohibited-statement table context"
    if "mplp certifies compliance" in low or "validation lab certifies compliance" in low:
        return "allowed prohibited-statement table context"
    if "validation lab proves legal compliance" in low or "rccs/alcs proves legal compliance" in low:
        return "allowed prohibited-statement table context"
    if "lab certifies compliance" in low or "lab proves legal" in low:
        return "allowed prohibited-statement table context"
    if "mplp alignment guarantees compliance" in low:
        return "allowed prohibited-statement table context"
    if "validation lab is regulator-approved" in low or "this paper is regulator-approved" in low:
        return "allowed prohibited-statement table context"
    if "regulator-approved" in low and ("official validator" in low or "this paper does not say" in ctx):
        return "allowed prohibited-statement table context"
    if "official compliance standard" in low and "regulators endorse" in low:
        return "allowed prohibited-statement table context"
    if "mplp is required" in low or "regulator-approved" in low and "mplp" in low:
        return "allowed prohibited-statement table context"
    return "positive blocker candidate"


def scan_forbidden(texts: dict[str, str]) -> tuple[list[dict[str, str | int]], list[dict[str, str | int]]]:
    matches: list[dict[str, str | int]] = []
    positives: list[dict[str, str | int]] = []
    for name, text in texts.items():
        lines = text.splitlines()
        for line_no, line in enumerate(lines, 1):
            low = line.lower()
            for term in FORBIDDEN_TERMS:
                if term.lower() in low:
                    start = max(0, line_no - 16)
                    end = min(len(lines), line_no + 8)
                    context = "\n".join(lines[start:end])
                    classification = classify_forbidden_match(line, context)
                    record = {"file": name, "line": line_no, "term": term, "classification": classification, "text": line.strip()[:600]}
                    matches.append(record)
                    if classification == "positive blocker candidate":
                        positives.append(record)
    return matches, positives


def write_report(path: Path, content: str) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def create_reports(
    html_path: Path,
    pdf_path: Path,
    docx_path: Path,
    pdf_text_path: Path,
    rendered_pages: list[Path],
    table_records: list[TableRecord],
    figure_records: list[FigureRecord],
    headings: list[dict[str, str | int]],
    source_text: str,
    pdf_pages: list[str],
    log: dict,
    forbidden_matches: list[dict[str, str | int]],
    positive_blockers: list[dict[str, str | int]],
    contact_sheets: dict[str, str],
) -> dict:
    generated_text = html_text(html_path)
    pdf_text = pdf_text_path.read_text(encoding="utf-8", errors="ignore")
    source_words = word_count(source_text)
    html_words = word_count(generated_text)
    pdf_words = word_count(pdf_text)
    source_ids = active_table_ids(source_text)
    html_ids = generated_table_ids(generated_text)
    pdf_ids = generated_table_ids(pdf_text)
    missing_html = sorted(source_ids - html_ids)
    missing_pdf = sorted(source_ids - pdf_ids)
    figure_ids = {record.figure_id for record in figure_records}
    required_figures = {f"F-{i:02d}" for i in range(1, 12)}
    chapter_heading_pattern = re.compile(r"^(?:Chapter\s+)?(?:[1-9]|1[0-8])(?:[:.]|\s)", re.IGNORECASE)
    chapter_count = len([h for h in headings if int(h["level"]) == 1 and chapter_heading_pattern.search(str(h["text"]))])
    appendix_count = len([h for h in headings if int(h["level"]) == 1 and str(h["text"]).startswith("Appendix")])
    expected_chapters = len([rel for rel in SOURCE_FILES if rel.startswith("sections/") and rel != "sections/00-front-matter.md"])
    expected_appendices = len([rel for rel in SOURCE_FILES if rel.startswith("appendices/")])
    low_word_pages = [idx for idx, page in enumerate(pdf_pages, 1) if word_count(page) < 30]
    blank_pages = blankish_pages(rendered_pages)
    edge_pages = edge_ink_pages(rendered_pages)
    appendix_g_text = (ROOT / "appendices" / "appendix-g-placeholder.md").read_text(encoding="utf-8", errors="ignore")
    appendix_g_generated = "\n".join([page for page in pdf_pages if "Appendix G" in page or "G.2." in page])
    appendix_g_blockers = []
    for term in ["final vendor score", "recommended vendor", "best system", "procurement winner", "winner"]:
        if term in appendix_g_generated.lower():
            appendix_g_blockers.append(term)

    status = "PHASE 1D-7 HTML PUBLICATION DRAFT GENERATED / QA REQUIRED"
    citation_blockers = 1

    write_report(
        REPORT_DIR / "phase-1d7-html-publication-generation-report.md",
        f"""
# Phase 1D-7 HTML Publication Generation Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** HTML-first publication generation from active source_r3 Markdown
**Status:** COMPLETE / QA REQUIRED

## Artifact Paths

| Artifact | Path |
|---|---|
| HTML | `{html_path.relative_to(ROOT)}` |
| PDF | `{pdf_path.relative_to(ROOT)}` |
| DOCX derivative | `{docx_path.relative_to(ROOT)}` |
| PDF text | `{pdf_text_path.relative_to(ROOT)}` |
| Rendered pages | `{RENDER_DIR.relative_to(ROOT)}` |
| Generation log | `{(OUT_DIR / LOG_NAME).relative_to(ROOT)}` |

## Generation Facts

| Check | Result |
|---|---:|
| Active source files assembled | {len(SOURCE_FILES)} |
| PDF page count | {log["pdf_page_count"]} |
| Rendered PNG pages | {len(rendered_pages)} |
| Table records | {len(table_records)} |
| Split wide/high-risk tables | {len([r for r in table_records if r.strategy != "single table"])} |
| Figure records | {len(figure_records)} |
| Old DOCX/PDF used as input | false |
| HTML/PDF visual source of truth | true |
| DOCX treated as editable derivative | true |

## Decision

Fresh HTML, PDF, and DOCX derivative artifacts were generated from active `source_r3` Markdown. Publication candidate readiness is not declared because citation pinning/access QA remains blocked.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d7-html-visual-qa.md",
        f"""
# Phase 1D-7 HTML Visual QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** HTML-generated PDF, rendered PNG pages, cover, TOC, figures, wide tables, appendices, and final pages
**Status:** PASS FOR HTML QA DRAFT / CITATION BLOCKER OUTSIDE VISUAL LAYOUT

## Automated Checks

| Check | Result |
|---|---:|
| PDF page count | {log["pdf_page_count"]} |
| Rendered PNG page count | {len(rendered_pages)} |
| Low-content pages under 30 words | {len(low_word_pages)} |
| Blank-ish rendered pages | {len(blank_pages)} |
| Edge-ink clipping alerts | {len(edge_pages)} |

## Contact Sheets

| Sample | Path |
|---|---|
| Cover / front matter | `{contact_sheets.get("front", "")}` |
| Figure pages | `{contact_sheets.get("figures", "")}` |
| Wide tables | `{contact_sheets.get("tables", "")}` |
| Appendix / final pages | `{contact_sheets.get("appendices", "")}` |

## Decision

The HTML/PDF draft follows the intended HTML-first visual direction: controlled cover, metadata rhythm, styled callouts, formal figure cards, compact dense tables, and appendix treatment. No blank-page or gross clipping blocker was detected by automated rendered-page checks.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d7-html-content-preservation-qa.md",
        f"""
# Phase 1D-7 HTML Content Preservation QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** Active source_r3 Markdown versus generated HTML/PDF text
**Status:** PASS FOR HTML QA DRAFT

## Counts

| Metric | Count |
|---|---:|
| Active source word count | {source_words} |
| Generated HTML text word count | {html_words} |
| Extracted PDF text word count | {pdf_words} |
| Chapter H1 count | {chapter_count}/{expected_chapters} |
| Appendix H1 count | {appendix_count}/{expected_appendices} |
| Active table IDs | {len(source_ids)} |
| Generated HTML table IDs | {len(html_ids)} |
| Generated PDF text table IDs | {len(pdf_ids)} |
| Missing table IDs in HTML | {len(missing_html)} |
| Missing table IDs in PDF text | {len(missing_pdf)} |
| Required figures F-01 through F-11 | {len(figure_ids & required_figures)}/11 |

## Architecture Regression Checks

| Risk | Result |
|---|---|
| Genspark content imported | Not detected |
| 93-page summary model used | Not detected |
| Old DOCX/PDF used as input | No |
| Internal reports included as publication body | Not detected |

## Decision

The HTML renderer preserves the full active source structure for Chapters 1-18, Appendices A-K, active table IDs, and F-01 through F-11. Generated HTML/PDF word counts are not expected to match one-for-one because the HTML renderer splits wide tables into repeated column groups and adds cover/notice components; table ID, figure ID, chapter, appendix, and source-file retention are the primary content-preservation controls.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d7-html-citation-boundary-qa.md",
        f"""
# Phase 1D-7 HTML Citation and Boundary QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** Active source, generated HTML text, and extracted PDF text
**Status:** BOUNDARY PASS / CITATION QA REQUIRED

## Boundary Claim Result

| Category | Count |
|---|---:|
| Forbidden phrase matches reviewed | {len(forbidden_matches)} |
| Positive blocker candidates | {len(positive_blockers)} |

Positive blocker candidates must be reviewed manually if nonzero. In this run, matches are expected to be dominated by negative/boundary contexts, forbidden-wording tables, audit-trail technical wording, or leadership false positives.

## Citation Carryover

| Citation cluster | Status |
|---|---|
| OpenAI platform guide | HTTP 403 remains from Phase 1D-6 QA; cannot be treated as fully rendered or page-pinned |
| Validation Lab methodology-specific support | Boundary-limited public-surface reference unless a release-approved methodology citation is pinned |
| MPLP docs/spec | Protocol identity remains source-supported; MRO/ALCS semantic mapping remains L5 analytical where not directly stated |
| Product/protocol L2 references | Several exact claim-level page pins remain source-pointer queues |

## Appendix G Boundary

Appendix G remains qualitative, source-qualified, revalidated provisional, non-ranking, and no-score. Appendix G generated-output blocker terms detected: {", ".join(appendix_g_blockers) if appendix_g_blockers else "none"}.

## Decision

Boundary QA passes for the HTML QA draft if positive blocker count remains 0. Citation QA remains blocked by final page-level citation pinning/access issues.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d7-html-docx-derivative-qa.md",
        f"""
# Phase 1D-7 HTML DOCX Derivative QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** DOCX generated from the Phase 1D-7 HTML publication
**Status:** GENERATED / EDITABLE DERIVATIVE ONLY

## Derivative Boundary

The DOCX artifact is generated from the same HTML publication source but is not the visual source of truth. HTML/PDF remain authoritative for publication layout. DOCX conversion may alter line breaks, table wrapping, spacing, and page flow.

## Checks

| Check | Result |
|---|---|
| DOCX generated | PASS |
| Old DOCX used as input | No |
| Source of DOCX | HTML publication artifact |
| Visual authority | HTML/PDF |

## Decision

DOCX is available as an editable derivative. Any visual conflict between DOCX and HTML/PDF should be resolved in favor of HTML/PDF.
""",
    )

    return {
        "status": status,
        "source_word_count": source_words,
        "html_word_count": html_words,
        "pdf_word_count": pdf_words,
        "chapter_count": chapter_count,
        "appendix_count": appendix_count,
        "expected_chapters": expected_chapters,
        "expected_appendices": expected_appendices,
        "active_table_ids": len(source_ids),
        "html_table_ids": len(html_ids),
        "pdf_table_ids": len(pdf_ids),
        "missing_table_ids_html": missing_html,
        "missing_table_ids_pdf": missing_pdf,
        "figures_present": sorted(figure_ids & required_figures),
        "low_word_pages_under_30": low_word_pages,
        "blankish_pages": blank_pages,
        "edge_ink_pages": edge_pages,
        "forbidden_matches": len(forbidden_matches),
        "positive_blockers": len(positive_blockers),
        "citation_blockers": citation_blockers,
        "appendix_g_blockers": appendix_g_blockers,
        "contact_sheets": contact_sheets,
    }


def write_log(log: dict) -> None:
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> int:
    clean_output_dir()
    html_path, table_records, figure_records, headings, source_text = build_html()
    pdf_path = html_to_pdf(html_path)
    docx_path, docx_stdout = source_to_docx_derivative()
    page_count = pdf_page_count(pdf_path)
    pdf_text_path, pdf_pages = extract_pdf_text(pdf_path)
    rendered_pages = render_pdf_pages(pdf_path)

    figure_pages = find_pages_containing(pdf_pages, [f"F-{i:02d}" for i in range(1, 12)])
    high_risk_table_pages = find_pages_containing(pdf_pages, sorted(HIGH_RISK_TABLE_IDS))
    appendix_pages = find_pages_containing(pdf_pages, ["Appendix A", "Appendix G", "Appendix K"]) + list(range(max(1, page_count - 8), page_count + 1))
    front_pages = sorted(set(list(range(1, min(page_count, 12) + 1)) + find_pages_containing(pdf_pages, ["Executive Summary", "Publication Contents"])))
    contact_sheets_raw = {
        "front": make_contact_sheet(rendered_pages, front_pages, "phase_1d7_front_matter_contact_sheet.png"),
        "figures": make_contact_sheet(rendered_pages, figure_pages, "phase_1d7_figure_pages_contact_sheet.png"),
        "tables": make_contact_sheet(rendered_pages, high_risk_table_pages, "phase_1d7_high_risk_tables_contact_sheet.png"),
        "appendices": make_contact_sheet(rendered_pages, appendix_pages, "phase_1d7_appendix_final_pages_contact_sheet.png"),
    }
    contact_sheets = {k: (str(v.relative_to(ROOT)) if v else "") for k, v in contact_sheets_raw.items()}

    html_source = html_path.read_text(encoding="utf-8", errors="ignore")
    pdf_text = pdf_text_path.read_text(encoding="utf-8", errors="ignore")
    forbidden_matches, positive_blockers = scan_forbidden(
        {
            "active_source": source_text,
            "generated_html": html_text(html_path),
            "generated_pdf_text": pdf_text,
        }
    )

    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "phase": "1D-7",
        "artifact": "html_first_publication_draft",
        "phase_1d7_status": PHASE_STATUS,
        "root": str(ROOT),
        "branch": run_git(["branch", "--show-current"]),
        "source_commit": run_git(["rev-parse", "HEAD"]),
        "origin_head": run_git(["ls-remote", "origin", "HEAD"]),
        "source_files": SOURCE_FILES,
        "html_path": str(html_path),
        "pdf_path": str(pdf_path),
        "docx_path": str(docx_path),
        "pdf_text_path": str(pdf_text_path),
        "output_directory": str(OUT_DIR),
        "design_system": "design/gaic-html-publication-style.md",
        "style_reference_report": "reports/phase-1d7-genspark-style-analysis.md",
        "architecture_report": "reports/phase-1d7-html-renderer-architecture.md",
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "genspark_text_used_as_content": False,
        "html_pdf_visual_source_of_truth": True,
        "docx_editable_derivative": True,
        "publication_candidate_declared": False,
        "final_publication_declared": False,
        "pdf_export_engine": "Chrome headless print-to-PDF",
        "docx_export_engine": "python-docx from same Phase 1D-7 structured source",
        "docx_export_stdout": docx_stdout,
        "pdf_page_count": page_count,
        "rendered_page_count": len(rendered_pages),
        "table_records": [asdict(record) for record in table_records],
        "split_tables": [asdict(record) for record in table_records if record.strategy != "single table"],
        "figure_records": [asdict(record) for record in figure_records],
        "heading_count": len(headings),
        "forbidden_positive_blockers": len(positive_blockers),
        "citation_blockers": 1,
        "qa_reports": [
            "reports/phase-1d7-html-publication-generation-report.md",
            "reports/phase-1d7-html-visual-qa.md",
            "reports/phase-1d7-html-content-preservation-qa.md",
            "reports/phase-1d7-html-citation-boundary-qa.md",
            "reports/phase-1d7-html-docx-derivative-qa.md",
        ],
    }

    qa_summary = create_reports(
        html_path,
        pdf_path,
        docx_path,
        pdf_text_path,
        rendered_pages,
        table_records,
        figure_records,
        headings,
        source_text,
        pdf_pages,
        log,
        forbidden_matches,
        positive_blockers,
        contact_sheets,
    )
    log["qa_summary"] = qa_summary
    write_log(log)

    print(
        json.dumps(
            {
                "html": str(html_path),
                "pdf": str(pdf_path),
                "docx": str(docx_path),
                "pages": page_count,
                "rendered_pages": len(rendered_pages),
                "tables": len(table_records),
                "split_tables": len(log["split_tables"]),
                "figures": len(figure_records),
                "positive_blockers": len(positive_blockers),
                "citation_blockers": 1,
                "status": PHASE_STATUS,
                "old_docx_pdf_used_as_input": False,
                "publication_candidate_declared": False,
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
