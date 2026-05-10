#!/usr/bin/env python3
"""
Phase 1D-3 technical evidence pack generator for GAIC-2026 v0.3.2 FRC-R3.

This script creates the dense companion artifact for the public white paper.
It uses current source_r3 Markdown only and does not read old DOCX/PDF files as
source inputs.
"""

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import generate_r3_phase1d_docx_pdf as base
import generate_r3_phase1d2_publication_draft as pub


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "phase_1d3" / "evidence_pack"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
DOCX_NAME = "GAIC-2026-v0.3.2-FRC-R3-Technical-Evidence-Pack.docx"
PDF_NAME = "GAIC-2026-v0.3.2-FRC-R3-Technical-Evidence-Pack.pdf"
MD_NAME = "GAIC-2026-v0.3.2-FRC-R3-Technical-Evidence-Pack.md"
LOG_NAME = "technical_evidence_pack_generation_log.json"


EVIDENCE_PACK_FILES = [
    "GAIC-2026-v0.3.2-FRC-R3-SOURCE-MASTER.md",
    *base.SOURCE_FILES,
    "inventories/table-inventory.md",
    "inventories/figure-inventory.md",
    "inventories/citation-inventory.md",
    "inventories/source-coverage-matrix.md",
    "inventories/claim-evidence-register.md",
    "inventories/page-level-citation-map.md",
    "inventories/citation-rendering-qa-checklist.md",
    "inventories/forbidden-claim-context-whitelist.md",
    "reports/appendix-g-no-score-proofing-report.md",
    "reports/phase-1c-cleanup-report.md",
    "reports/phase-1c-claim-level-revalidation-report.md",
    "reports/phase-1c-final-citation-pinning-report.md",
    "reports/phase-1d-generation-and-qa-report.md",
    "reports/phase-1d-citation-rendering-qa.md",
    "reports/phase-1d-table-layout-qa.md",
    "reports/phase-1d-figure-qa.md",
    "reports/phase-1d-appendix-g-proofing-report.md",
    "reports/phase-1d-forbidden-claim-sweep.md",
    "reports/phase-1d2-publication-design-audit.md",
    "reports/phase-1d2-figure-production-plan.md",
    "reports/phase-1d2-table-reflow-plan.md",
    "reports/phase-1d2-generation-and-qa-report.md",
    "reports/phase-1d3-publication-architecture-decision.md",
    "reports/phase-1d3-source-split-plan.md",
    "reports/phase-1d3-narrative-reflow-report.md",
    "reports/phase-1d3-table-compression-report.md",
    "reports/phase-1d3-public-whitepaper-qa.md",
    "reports/phase-1d3-evidence-pack-qa.md",
    "reports/phase-1d3-final-status-report.md",
]


def configure_base() -> None:
    base.OUT_DIR = OUT_DIR
    base.RENDER_DIR = RENDER_DIR
    base.DOCX_NAME = DOCX_NAME
    base.PDF_NAME = PDF_NAME
    base.ASSEMBLED_MD_NAME = MD_NAME
    base.LOG_NAME = LOG_NAME


def write_evidence_markdown() -> Path:
    chunks = [
        "# GAIC-2026 v0.3.2 FRC-R3 Technical Evidence Pack",
        "",
        f"**Generated:** {datetime.now().strftime('%B %d, %Y')}",
        "**Phase:** 1D-3 Publication Compression and Narrative Reflow",
        "**Artifact Role:** Dense technical evidence companion to the public white paper",
        "",
        "This evidence pack preserves the technical substrate intentionally removed from the public white paper: full body source, Appendices A-K, full rubrics, detailed mappings, source registers, citation ledgers, claim evidence register, inventories, and QA reports.",
        "",
        "It is not a legal compliance determination, certification, regulatory approval, procurement recommendation, vendor ranking, or final vendor assessment.",
        "",
        "## Evidence Pack Source Index",
        "",
    ]
    for rel in EVIDENCE_PACK_FILES:
        chunks.append(f"- `{rel}`")
    chunks.append("")

    for rel in EVIDENCE_PACK_FILES:
        path = ROOT / rel
        if not path.exists():
            chunks.append(f"\n\n<!-- MISSING SOURCE: {rel} -->\n")
            continue
        chunks.append(f"\n\n<!-- SOURCE: {rel} -->\n\n")
        chunks.append(path.read_text(encoding="utf-8"))

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = OUT_DIR / MD_NAME
    md_path.write_text("\n".join(chunks).strip() + "\n", encoding="utf-8")
    return md_path


def set_evidence_defaults(doc: Document) -> None:
    pub.set_publication_defaults(doc)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("GAIC-2026 v0.3.2-FRC-R3 | Technical Evidence Pack | Dense Review Artifact")
        run.font.name = "Arial"
        run.font.size = Pt(7.0)
        run.font.color.rgb = pub.ANODIZED_SLATE


def build_docx(md_path: Path) -> tuple[Path, list[base.TableRecord], str]:
    doc = Document()
    set_evidence_defaults(doc)
    records: list[base.TableRecord] = []
    text = pub.add_file_to_publication_doc(doc, md_path, records)
    docx_path = OUT_DIR / DOCX_NAME
    doc.save(docx_path)
    return docx_path, records, text


def write_log(docx_path: Path, pdf_path: Path, rendered_pages: list[Path], records: list[base.TableRecord], source_text: str) -> dict:
    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "phase": "1D-3",
        "artifact": "technical_evidence_pack",
        "root": str(ROOT),
        "source_files": EVIDENCE_PACK_FILES,
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "assembled_markdown_path": str(OUT_DIR / MD_NAME),
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "full_appendices_a_k_included": True,
        "public_whitepaper_artifact": False,
        "pdf_page_count": base.pdf_page_count(pdf_path),
        "rendered_page_count": len(rendered_pages),
        "docx_stats": base.docx_stats(docx_path),
        "table_count": len(records),
        "table_records": [base.asdict(record) for record in records],
        "high_risk_table_count": len([record for record in records if record.high_risk or (record.table_id and record.table_id in base.HIGH_RISK_TABLE_IDS)]),
        "source_text_character_count": len(source_text),
    }
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    return log


def main() -> int:
    configure_base()
    md_path = write_evidence_markdown()
    docx_path, records, source_text = build_docx(md_path)
    pdf_path = base.convert_docx_to_pdf(docx_path)
    rendered_pages = base.render_pdf_pages(pdf_path)
    log = write_log(docx_path, pdf_path, rendered_pages, records, source_text)
    print(
        json.dumps(
            {
                "docx": str(docx_path),
                "pdf": str(pdf_path),
                "pages": log["pdf_page_count"],
                "tables": len(records),
                "rendered_pages": len(rendered_pages),
                "log": str(OUT_DIR / LOG_NAME),
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
