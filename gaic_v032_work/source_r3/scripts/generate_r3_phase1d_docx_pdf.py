#!/usr/bin/env python3
"""
Fresh Phase 1D DOCX/PDF generator for GAIC-2026 v0.3.2 FRC-R3.

This script assembles the active source_r3 Markdown files only. It does not
read prior DOCX/PDF artifacts as input.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "phase_1d"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
DOCX_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D-Draft.docx"
PDF_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D-Draft.pdf"
ASSEMBLED_MD_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D-Draft.md"
LOG_NAME = "phase_1d_generation_log.json"

SOURCE_FILES = [
    "sections/00-front-matter.md",
    "sections/01-scope-methodology.md",
    "sections/02-missing-layer.md",
    "sections/03-agentic-ai-breaks-model-compliance.md",
    "sections/04-regulatory-standards-baseline.md",
    "sections/05-regulatory-language-to-engineering-objects.md",
    "sections/06-missing-regulatory-objects.md",
    "sections/07-rccs-regulatory-compliance-coverage.md",
    "sections/08-alcs-agentic-lifecycle-conformance.md",
    "sections/09-composite-scoring-method.md",
    "sections/10-enterprise-control-crosswalk.md",
    "sections/11-comparative-field.md",
    "sections/12-detailed-system-mappings.md",
    "sections/13-mplp-deep-mapping.md",
    "sections/14-evidence-validation-placeholder.md",
    "sections/15-failure-scenarios-placeholder.md",
    "sections/16-companion-paper-boundary.md",
    "sections/17-adoption-roadmap-placeholder.md",
    "sections/18-conclusion-placeholder.md",
    "appendices/appendix-a-placeholder.md",
    "appendices/appendix-b-placeholder.md",
    "appendices/appendix-c-placeholder.md",
    "appendices/appendix-d-placeholder.md",
    "appendices/appendix-e-placeholder.md",
    "appendices/appendix-f-placeholder.md",
    "appendices/appendix-g-placeholder.md",
    "appendices/appendix-h-placeholder.md",
    "appendices/appendix-i-placeholder.md",
    "appendices/appendix-j-placeholder.md",
    "appendices/appendix-k-placeholder.md",
]

PUBLICATION_INVENTORY_INPUTS = [
    "GAIC-2026-v0.3.2-FRC-R3-SOURCE-MASTER.md",
    "inventories/table-inventory.md",
    "inventories/figure-inventory.md",
    "inventories/citation-inventory.md",
    "inventories/claim-evidence-register.md",
    "inventories/page-level-citation-map.md",
    "inventories/citation-rendering-qa-checklist.md",
    "inventories/forbidden-claim-context-whitelist.md",
]

FORBIDDEN_TERMS = [
    "certifies compliance",
    "certification authority",
    "regulator-approved",
    "regulatory approval",
    "proves legal compliance",
    "guarantees compliance",
    "only validation path",
    "official compliance standard",
    "conformity assessment body",
    "market endorsement",
    "recommended vendor",
    "best system",
    "procurement winner",
    "winner",
    "leader",
    "trails",
    "final vendor score",
    "MPLP is required",
    "MPLP certifies",
    "Validation Lab certifies compliance",
]

APPENDIX_G_STRICT_TERMS = [
    "Composite Score",
    "Ranking",
    "score",
    "ranked",
    "leads",
    "trails",
    "recommended vendor",
    "best system",
    "winner",
    "leader",
    "trailer",
    "final vendor score",
]

HIGH_RISK_TABLE_IDS = [
    "T-D-01",
    "T-13-02",
    "T-13-04",
    "T-14-01",
    "T-14-04",
    "T-15-01",
    "T-15-02",
    "T-15-03",
    "T-15-04",
    "T-17-01",
    "T-17-02",
    "T-17-03",
    "T-17-04",
    "T-17-05",
    "T-18-01",
    "T-18-02",
    "T-G-01",
    "T-G-02",
    "T-G-03",
    "T-G-04",
    "T-G-05",
    "T-G-06",
    "T-G-07",
    "T-G-08",
    "T-G-09",
    "T-H-01",
    "T-H-02",
    "T-H-03",
    "T-H-04",
    "T-I-01",
    "T-I-02",
    "T-I-03",
    "T-I-04",
    "T-J-01",
    "T-J-02",
    "T-J-03",
    "T-K-01",
    "T-K-02",
    "T-K-03",
]


@dataclass
class TableRecord:
    table_id: str | None
    caption: str | None
    source_file: str
    rows: int
    cols: int
    orientation: str
    font_size: float
    high_risk: bool


def which(name: str) -> str | None:
    return shutil.which(name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_text_direction(cell, direction: str = "lrTb") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    text_dir = OxmlElement("w:textDirection")
    text_dir.set(qn("w:val"), direction)
    tc_pr.append(text_dir)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
        ("Heading 4", 10.5, RGBColor(64, 64, 64)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(5)
    styles["Heading 1"].paragraph_format.page_break_before = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("GAIC-2026 v0.3.2-FRC-R3 Phase 1D Draft")


def clean_inline_markdown(text: str) -> str:
    text = text.replace("<br>", " / ").replace("<br/>", " / ").replace("<br />", " / ")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.strip()


def add_markdown_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    text = clean_inline_markdown(text)
    if not text:
        doc.add_paragraph()
        return
    para = doc.add_paragraph(style=style)
    para.add_run(text)


def add_code_block(doc: Document, code: list[str]) -> None:
    if not code:
        return
    para = doc.add_paragraph()
    run = para.add_run("\n".join(code))
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(5)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for raw in table_lines:
        parts = [p.strip() for p in raw.strip("|").split("|")]
        if parts and all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            continue
        rows.append([clean_inline_markdown(p) for p in parts])
    return rows, i


def is_table_start(lines: list[str], i: int) -> bool:
    if i + 1 >= len(lines):
        return False
    a = lines[i].strip()
    b = lines[i + 1].strip()
    return a.startswith("|") and a.endswith("|") and b.startswith("|") and b.endswith("|") and "---" in b


def detect_table_id(caption: str | None, rows: list[list[str]]) -> str | None:
    candidates = []
    if caption:
        candidates.append(caption)
    for row in rows[:3]:
        candidates.extend(row)
    joined = " ".join(candidates)
    match = re.search(r"\bT-[A-Z0-9]+(?:-[0-9A-Z]+)?\b", joined)
    return match.group(0) if match else None


def choose_table_layout(table_id: str | None, cols: int, caption: str | None) -> tuple[str, float]:
    high_risk = bool(table_id and table_id in HIGH_RISK_TABLE_IDS)
    text = f"{caption or ''} {table_id or ''}".lower()
    landscape = cols >= 5 or high_risk or "landscape" in text
    if cols >= 6 or high_risk:
        return ("landscape" if landscape else "portrait", 5.8)
    if cols == 5:
        return ("landscape", 6.5)
    if cols == 4:
        return ("portrait", 7.2)
    return ("portrait", 8.0)


def add_section_for_orientation(doc: Document, orientation: str) -> None:
    section = doc.add_section()
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT


def add_table_to_doc(
    doc: Document,
    rows: list[list[str]],
    source_file: str,
    caption: str | None,
    records: list[TableRecord],
) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table_id = detect_table_id(caption, rows)
    orientation, font_size = choose_table_layout(table_id, cols, caption)
    high_risk = bool(table_id and table_id in HIGH_RISK_TABLE_IDS)

    add_section_for_orientation(doc, orientation)

    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    page_width = 10.0 if orientation == "landscape" else 7.1
    col_width = max(650, int((page_width / cols) * 1440))

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(tr)
        set_table_cant_split(tr)
        for c_idx, value in enumerate(row):
            cell = tr.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, col_width)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.bold = r_idx == 0
            if cols >= 6 and len(value) <= 20 and r_idx == 0:
                set_text_direction(cell)

    records.append(
        TableRecord(
            table_id=table_id,
            caption=caption,
            source_file=source_file,
            rows=len(rows),
            cols=cols,
            orientation=orientation,
            font_size=font_size,
            high_risk=high_risk,
        )
    )
    add_section_for_orientation(doc, "portrait")


def heading_level(line: str) -> int:
    return len(line) - len(line.lstrip("#"))


def is_caption_line(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^\*\*Table\s+T-[^*]+\*\*$", stripped, re.IGNORECASE))


def add_figure_placeholder(doc: Document, line: str) -> None:
    text = clean_inline_markdown(line.strip().strip("[]"))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.size = Pt(9)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)


def add_file_to_doc(doc: Document, path: Path, records: list[TableRecord]) -> str:
    relative = path.relative_to(ROOT).as_posix()
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if is_table_start(lines, i):
            rows, new_i = parse_markdown_table(lines, i)
            add_table_to_doc(doc, rows, relative, pending_caption, records)
            pending_caption = None
            i = new_i
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = min(heading_level(stripped), 4)
            text = clean_inline_markdown(stripped[level:].strip())
            doc.add_heading(text, level=level)
            pending_caption = None
            i += 1
            continue

        if is_caption_line(stripped):
            pending_caption = clean_inline_markdown(stripped)
            para = doc.add_paragraph()
            para.paragraph_format.keep_with_next = True
            run = para.add_run(pending_caption)
            run.bold = True
            run.font.size = Pt(8.5)
            i += 1
            continue

        if "FIGURE" in stripped.upper() or re.match(r"^\[FIGURE", stripped, re.IGNORECASE):
            add_figure_placeholder(doc, stripped)
            pending_caption = None
            i += 1
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.22)
            run = para.add_run(clean_inline_markdown(stripped.lstrip("> ")))
            run.italic = True
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bullet_match:
            add_markdown_paragraph(doc, bullet_match.group(2), style="List Bullet")
            i += 1
            continue
        if number_match:
            # Preserve source numbering literally. Word's built-in List Number
            # style can continue numbering across separated Markdown lists after
            # LibreOffice conversion, which is worse than a plain numbered line.
            add_markdown_paragraph(doc, f"{number_match.group(2)}. {number_match.group(3)}")
            i += 1
            continue

        add_markdown_paragraph(doc, line)
        i += 1

    if code_buffer:
        add_code_block(doc, code_buffer)
    return text


def assemble_markdown() -> str:
    chunks = []
    for rel in SOURCE_FILES:
        path = ROOT / rel
        chunks.append(f"\n\n<!-- SOURCE: {rel} -->\n\n")
        chunks.append(path.read_text(encoding="utf-8"))
    assembled = "\n".join(chunks)
    (OUT_DIR / ASSEMBLED_MD_NAME).write_text(assembled, encoding="utf-8")
    return assembled


def historical_artifacts() -> list[str]:
    artifacts = []
    base = ROOT.parent
    for path in base.rglob("*"):
        if path.is_file() and path.suffix.lower() in {".docx", ".pdf"}:
            try:
                artifacts.append(path.relative_to(ROOT).as_posix())
            except ValueError:
                artifacts.append(str(path))
    return sorted(artifacts)


def build_docx() -> tuple[Path, list[TableRecord], str]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_defaults(doc)
    records: list[TableRecord] = []

    source_texts = []
    for rel in SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            raise FileNotFoundError(f"Required source file missing: {rel}")
        source_texts.append(add_file_to_doc(doc, path, records))

    docx_path = OUT_DIR / DOCX_NAME
    doc.save(docx_path)
    return docx_path, records, "\n".join(source_texts)


def convert_docx_to_pdf(docx_path: Path) -> Path:
    soffice = which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if not Path(soffice).exists() and not which("soffice"):
        raise RuntimeError("LibreOffice soffice not found")
    with tempfile.TemporaryDirectory(prefix="gaic_lo_profile_") as profile:
        cmd = [
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT_DIR),
            str(docx_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    generated = OUT_DIR / (docx_path.stem + ".pdf")
    final_pdf = OUT_DIR / PDF_NAME
    if generated != final_pdf:
        if final_pdf.exists():
            final_pdf.unlink()
        generated.rename(final_pdf)
    return final_pdf


def render_pdf_pages(pdf_path: Path) -> list[Path]:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for existing in RENDER_DIR.glob("*.png"):
        existing.unlink()
    pdftoppm = which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm not found")
    prefix = RENDER_DIR / "page"
    cmd = [pdftoppm, "-png", "-r", "120", str(pdf_path), str(prefix)]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError(f"pdftoppm render failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    return sorted(RENDER_DIR.glob("page-*.png"))


def pdf_page_count(pdf_path: Path) -> int | None:
    pdfinfo = which("pdfinfo")
    if not pdfinfo:
        return None
    result = subprocess.run([pdfinfo, str(pdf_path)], capture_output=True, text=True, timeout=60)
    if result.returncode != 0:
        return None
    match = re.search(r"^Pages:\s+(\d+)", result.stdout, re.MULTILINE)
    return int(match.group(1)) if match else None


def docx_stats(docx_path: Path) -> dict:
    doc = Document(docx_path)
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            headings.append({"style": para.style.name, "text": para.text[:160]})
    return {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "heading_count": len(headings),
        "heading_sample": headings[:30],
    }


def scan_terms(paths: Iterable[Path], terms: list[str]) -> dict[str, list[dict[str, str | int]]]:
    results: dict[str, list[dict[str, str | int]]] = {term: [] for term in terms}
    for path in paths:
        text = path.read_text(encoding="utf-8", errors="ignore")
        for line_no, line in enumerate(text.splitlines(), 1):
            low = line.lower()
            for term in terms:
                if term.lower() in low:
                    results[term].append(
                        {
                            "file": path.relative_to(ROOT).as_posix(),
                            "line": line_no,
                            "text": line.strip()[:260],
                        }
                    )
    return results


def classify_forbidden_match(line: str) -> str:
    low = line.lower()
    boundary_markers = [
        "not ",
        "does not",
        "do not",
        "no ",
        "forbidden",
        "boundary",
        "non-claim",
        "what assessments do not provide",
        "must not",
        "negative",
        "disclaimer",
    ]
    report_markers = ["grep", "sweep", "whitelist", "historical", "completion report", "proofing report"]
    if any(marker in low for marker in report_markers):
        return "historical/report or grep-verification context"
    if any(marker in low for marker in boundary_markers):
        return "negative boundary / non-claim context"
    return "requires review"


def appendix_g_term_scan() -> dict[str, list[dict[str, str | int]]]:
    return scan_terms([ROOT / "appendices/appendix-g-placeholder.md"], APPENDIX_G_STRICT_TERMS)


def active_source_paths() -> list[Path]:
    return [ROOT / rel for rel in SOURCE_FILES] + [ROOT / rel for rel in PUBLICATION_INVENTORY_INPUTS]


def write_log(
    docx_path: Path,
    pdf_path: Path,
    table_records: list[TableRecord],
    rendered_pages: list[Path],
    source_text: str,
) -> dict:
    forbidden = scan_terms(active_source_paths(), FORBIDDEN_TERMS)
    forbidden_review = []
    for term, matches in forbidden.items():
        for match in matches:
            forbidden_review.append({**match, "term": term, "classification": classify_forbidden_match(str(match["text"]))})
    positive_like = [m for m in forbidden_review if m["classification"] == "requires review"]

    app_g = appendix_g_term_scan()
    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "root": str(ROOT),
        "source_files": SOURCE_FILES,
        "publication_inventory_inputs_for_qa": PUBLICATION_INVENTORY_INPUTS,
        "old_artifacts_not_used_as_input": historical_artifacts(),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "assembled_markdown_path": str(OUT_DIR / ASSEMBLED_MD_NAME),
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "docx_stats": docx_stats(docx_path),
        "pdf_page_count": pdf_page_count(pdf_path),
        "rendered_page_count": len(rendered_pages),
        "rendered_pages_sample": [str(p) for p in rendered_pages[:10]],
        "table_records": [asdict(r) for r in table_records],
        "high_risk_tables_rendered": [asdict(r) for r in table_records if r.high_risk or (r.table_id and r.table_id in HIGH_RISK_TABLE_IDS)],
        "forbidden_scan": forbidden,
        "forbidden_requires_review": positive_like,
        "appendix_g_strict_scan": app_g,
        "source_text_character_count": len(source_text),
    }
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    return log


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    assemble_markdown()
    docx_path, table_records, source_text = build_docx()
    pdf_path = convert_docx_to_pdf(docx_path)
    rendered_pages = render_pdf_pages(pdf_path)
    log = write_log(docx_path, pdf_path, table_records, rendered_pages, source_text)
    print(json.dumps({
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "pages": log["pdf_page_count"],
        "tables": len(table_records),
        "rendered_pages": len(rendered_pages),
        "log": str(OUT_DIR / LOG_NAME),
    }, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
