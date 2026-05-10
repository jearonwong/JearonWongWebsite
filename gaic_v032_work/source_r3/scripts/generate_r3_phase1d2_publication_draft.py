#!/usr/bin/env python3
"""
Phase 1D-2 publication-design draft generator for GAIC-2026 v0.3.2 FRC-R3.

This script uses the current source_r3 Markdown assets only. It writes fresh
Phase 1D-2 artifacts into out/phase_1d2 and does not read previous DOCX/PDF
artifacts as source content.
"""

from __future__ import annotations

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import generate_r3_phase1d_docx_pdf as base


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "phase_1d2"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
DOCX_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D2-Publication-Draft.docx"
PDF_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D2-Publication-Draft.pdf"
ASSEMBLED_MD_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Phase1D2-Publication-Draft.md"
LOG_NAME = "phase_1d2_generation_log.json"

MONOLITH_CHARCOAL = RGBColor(17, 19, 24)
ANODIZED_SLATE = RGBColor(60, 72, 88)
PROTOCOL_BLUE = RGBColor(31, 78, 121)
EVIDENCE_GREEN = RGBColor(53, 107, 86)
BOUNDARY_AMBER = RGBColor(138, 106, 32)
GRID_GRAY = "D9DEE5"
HEADER_FILL = "EEF2F6"
CHARCOAL_FILL = "111318"
PANEL_FILL = "F7F7F4"


FIGURE_PANELS = {
    "F-01": {
        "title": "AI Agent Lifecycle Governance Stack",
        "lanes": ["Intent", "Context", "Plan", "Authority", "Agent Work", "Tool Actions", "Evidence", "Review", "Accepted Outcome", "Dispute / Remediate"],
        "note": "Lifecycle responsibility is proven across the work unit, not only at model output.",
    },
    "F-02": {
        "title": "Governance Stack / Missing Layer",
        "lanes": ["Enterprise Trust", "AI Agent Lifecycle Governance", "Agent Orchestration", "Model Governance", "Infrastructure"],
        "note": "The missing layer sits between execution mechanics and enterprise accountability.",
    },
    "F-03": {
        "title": "MRO Topology",
        "lanes": ["Responsibility", "Authority", "Evidence", "Privacy", "Substitution", "Closure"],
        "note": "Sixteen MROs are grouped by the lifecycle control surface they make testable.",
    },
    "F-04": {
        "title": "RCCS/ALCS Dual Scoring Model",
        "lanes": ["RCCS", "Evidence Multiplier", "ALCS", "Boundary Review", "Analytical Composite"],
        "note": "Scoring is analytical and does not prove legal compliance or certification.",
    },
    "F-05": {
        "title": "Lifecycle Conformance Mapping Strength Heatmap",
        "lanes": ["Visible Primitive", "Partial Object", "Implementer Integration", "L5 Inference", "Phase 1C Source Binding"],
        "note": "The heatmap is qualitative and non-ranking.",
    },
    "F-06": {
        "title": "Enterprise Control Overlay",
        "lanes": ["Audit", "Privacy", "Security", "Procurement", "Board Reporting", "Incident Response"],
        "note": "Lifecycle objects extend existing controls; they do not replace them.",
    },
    "F-07": {
        "title": "Model Risk in Agentic Lifecycle",
        "lanes": ["Model Change", "Authority Boundary", "Tool Action", "Evidence Chain", "Accepted Outcome"],
        "note": "Model risk becomes enterprise risk when embedded in delegated workflows.",
    },
    "F-08": {
        "title": "Evidence-Based Validation Pattern Flow",
        "lanes": ["Evidence Pack", "Hash Manifest", "Ruleset", "Evidence Pointers", "Verdict Hash", "Replay / Recheck"],
        "note": "Validation Lab remains a non-certifying evidence adjudication example.",
    },
    "F-09": {
        "title": "Enterprise Failure Scenario Chain",
        "lanes": ["Trigger", "Missing Object", "Evidence Gap", "Enterprise Consequence", "Remediation Closure"],
        "note": "The chain is an analytical pattern, not an incident probability model.",
    },
    "F-10": {
        "title": "Companion Paper Boundary Map",
        "lanes": ["Main Framework", "Evidence & Assurance", "Insurability", "Reserved Topics"],
        "note": "The main white paper defines the framework; companion papers carry deeper assurance and insurance work.",
    },
    "F-11": {
        "title": "Stage 0-to-7 Adoption Roadmap",
        "lanes": ["0 Inventory", "1 Authority", "2 Responsibility", "3 Evidence", "4 Privacy", "5 Validation", "6 Remediation", "7 Scale"],
        "note": "The roadmap is an adoption path, not a certification checklist.",
    },
}


def patch_base_paths() -> None:
    base.OUT_DIR = OUT_DIR
    base.RENDER_DIR = RENDER_DIR
    base.DOCX_NAME = DOCX_NAME
    base.PDF_NAME = PDF_NAME
    base.ASSEMBLED_MD_NAME = ASSEMBLED_MD_NAME
    base.LOG_NAME = LOG_NAME


def set_cell_border(cell, color: str = "D9DEE5", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_publication_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = MONOLITH_CHARCOAL
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before in [
        ("Title", 28, MONOLITH_CHARCOAL, 4),
        ("Heading 1", 17, MONOLITH_CHARCOAL, 16),
        ("Heading 2", 13.5, PROTOCOL_BLUE, 11),
        ("Heading 3", 11.5, ANODIZED_SLATE, 8),
        ("Heading 4", 10.5, ANODIZED_SLATE, 6),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(5)
    styles["Heading 1"].paragraph_format.page_break_before = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("GAIC-2026 v0.3.2-FRC-R3 | Phase 1D-2 Publication Draft | Not Final Publication Candidate")
    run.font.name = "Arial"
    run.font.size = Pt(7.2)
    run.font.color.rgb = ANODIZED_SLATE


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(92)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Global AI Compliance White Paper 2026")
    run.font.name = "Arial"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = MONOLITH_CHARCOAL

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    srun = subtitle.add_run("From Model Governance to Agentic Lifecycle Conformance")
    srun.font.name = "Arial"
    srun.font.size = Pt(14)
    srun.font.color.rgb = ANODIZED_SLATE

    strip = doc.add_table(rows=2, cols=5)
    strip.alignment = WD_TABLE_ALIGNMENT.LEFT
    strip.autofit = True
    labels = ["Evidence", "Authority", "Agent Work", "Accepted Outcome", "Remediation"]
    for idx, label in enumerate(labels):
        cell = strip.rows[0].cells[idx]
        base.set_cell_shading(cell, CHARCOAL_FILL)
        set_cell_border(cell, "111318", "8")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(247, 247, 244)
    for idx, label in enumerate(["source-bound", "delegated", "tool-linked", "reviewed", "closed"]):
        cell = strip.rows[1].cells[idx]
        base.set_cell_shading(cell, PANEL_FILL)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(7.5)
        r.font.color.rgb = ANODIZED_SLATE

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(22)
    meta.paragraph_format.space_after = Pt(4)
    mrun = meta.add_run("GACWP-2026-v0.3.2-FRC-R3 | Phase 1D-2 Publication Design Draft | May 2026")
    mrun.font.name = "Arial"
    mrun.font.size = Pt(8.8)
    mrun.font.color.rgb = ANODIZED_SLATE

    boundary = doc.add_paragraph()
    brun = boundary.add_run("Non-legal technical governance analysis. Not legal advice, certification, regulatory approval, procurement recommendation, or final vendor assessment.")
    brun.font.name = "Arial"
    brun.font.size = Pt(8.4)
    brun.font.color.rgb = BOUNDARY_AMBER
    doc.add_page_break()


def figure_id_from_line(line: str) -> str | None:
    match = re.search(r"\bF-\d{2}\b", line)
    if match:
        return match.group(0)
    if "Figure 2:" in line or "GOVERNANCE STACK" in line.upper():
        return "F-02"
    return None


def add_publication_figure_panel(doc: Document, line: str) -> None:
    figure_id = figure_id_from_line(line) or "F-01"
    panel = FIGURE_PANELS.get(figure_id, FIGURE_PANELS["F-01"])

    doc.add_paragraph()
    caption = doc.add_paragraph()
    caption.paragraph_format.keep_with_next = True
    crun = caption.add_run(f"Figure {figure_id}: {panel['title']}")
    crun.font.name = "Arial"
    crun.font.size = Pt(9)
    crun.font.bold = True
    crun.font.color.rgb = MONOLITH_CHARCOAL

    table = doc.add_table(rows=2, cols=len(panel["lanes"]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for idx, label in enumerate(panel["lanes"]):
        top = table.rows[0].cells[idx]
        base.set_cell_shading(top, CHARCOAL_FILL)
        set_cell_border(top, "111318", "6")
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.name = "Arial"
        run.font.size = Pt(7.4 if len(panel["lanes"]) > 6 else 8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(247, 247, 244)

        bottom = table.rows[1].cells[idx]
        base.set_cell_shading(bottom, PANEL_FILL)
        set_cell_border(bottom)
        q = bottom.paragraphs[0]
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker = "object" if idx % 2 == 0 else "control"
        br = q.add_run(marker)
        br.font.name = "Arial"
        br.font.size = Pt(6.8)
        br.font.color.rgb = ANODIZED_SLATE

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nrun = note.add_run(panel["note"])
    nrun.font.name = "Arial"
    nrun.font.size = Pt(7.8)
    nrun.italic = True
    nrun.font.color.rgb = BOUNDARY_AMBER if "not " in panel["note"].lower() else ANODIZED_SLATE


def is_publication_figure_marker(stripped: str) -> bool:
    if stripped.startswith("[FIGURE"):
        return True
    if "FIGURE PLACEHOLDER" in stripped.upper():
        return True
    if re.match(r"^\*\*Figure F-\d{2}:", stripped):
        return True
    if re.match(r"^Figure F-\d{2}:", stripped):
        return True
    if re.match(r"^\[FIGURE PLACEHOLDER F-\d{2}:", stripped, re.IGNORECASE):
        return True
    return False


def add_publication_table_to_doc(doc: Document, rows, source_file: str, caption: str | None, records) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table_id = base.detect_table_id(caption, rows)
    orientation, font_size = base.choose_table_layout(table_id, cols, caption)
    high_risk = bool(table_id and table_id in base.HIGH_RISK_TABLE_IDS)
    if high_risk and font_size < 6.2:
        font_size = 6.2

    base.add_section_for_orientation(doc, orientation)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    page_width = 10.0 if orientation == "landscape" else 7.1
    col_width = max(650, int((page_width / cols) * 1440))

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        if r_idx == 0:
            base.set_repeat_table_header(tr)
        base.set_table_cant_split(tr)
        for c_idx, value in enumerate(row):
            cell = tr.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            base.set_cell_width(cell, col_width)
            set_cell_border(cell)
            if r_idx == 0:
                base.set_cell_shading(cell, HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.bold = r_idx == 0
            run.font.color.rgb = MONOLITH_CHARCOAL if r_idx else ANODIZED_SLATE

    records.append(
        base.TableRecord(
            table_id=table_id,
            caption=caption,
            source_file=source_file,
            rows=len(rows),
            cols=cols,
            orientation=orientation,
            font_size=font_size,
            high_risk=high_risk,
        )
    )
    base.add_section_for_orientation(doc, "portrait")


def add_file_to_publication_doc(doc: Document, path: Path, records) -> str:
    relative = path.relative_to(ROOT).as_posix()
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                base.add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if base.is_table_start(lines, i):
            rows, new_i = base.parse_markdown_table(lines, i)
            add_publication_table_to_doc(doc, rows, relative, pending_caption, records)
            pending_caption = None
            i = new_i
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = min(base.heading_level(stripped), 4)
            heading_text = base.clean_inline_markdown(stripped[level:].strip())
            if relative == "sections/00-front-matter.md" and heading_text == "Global AI Compliance White Paper 2026":
                i += 1
                continue
            doc.add_heading(heading_text, level=level)
            pending_caption = None
            i += 1
            continue

        if base.is_caption_line(stripped):
            pending_caption = base.clean_inline_markdown(stripped)
            para = doc.add_paragraph()
            para.paragraph_format.keep_with_next = True
            run = para.add_run(pending_caption)
            run.bold = True
            run.font.size = Pt(8.8)
            run.font.color.rgb = MONOLITH_CHARCOAL
            i += 1
            continue

        if is_publication_figure_marker(stripped):
            add_publication_figure_panel(doc, stripped)
            pending_caption = None
            i += 1
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.22)
            run = para.add_run(base.clean_inline_markdown(stripped.lstrip("> ")))
            run.italic = True
            run.font.color.rgb = ANODIZED_SLATE
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bullet_match:
            base.add_markdown_paragraph(doc, bullet_match.group(2), style="List Bullet")
            i += 1
            continue
        if number_match:
            base.add_markdown_paragraph(doc, f"{number_match.group(2)}. {number_match.group(3)}")
            i += 1
            continue

        base.add_markdown_paragraph(doc, line)
        i += 1

    if code_buffer:
        base.add_code_block(doc, code_buffer)
    return text


def build_publication_docx() -> tuple[Path, list[base.TableRecord], str]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_publication_defaults(doc)
    add_cover(doc)
    records: list[base.TableRecord] = []
    source_texts = []

    for rel in base.SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            raise FileNotFoundError(f"Required source file missing: {rel}")
        source_texts.append(add_file_to_publication_doc(doc, path, records))

    docx_path = OUT_DIR / DOCX_NAME
    doc.save(docx_path)
    return docx_path, records, "\n".join(source_texts)


def augment_log(log: dict) -> dict:
    log.update(
        {
            "phase": "1D-2",
            "publication_design_system": "design/gaic-whitepaper-vi-system.md",
            "design_position": "JearonWong publication design draft; architectural, evidence-first, restrained protocol visual system",
            "output_directory": str(OUT_DIR),
            "phase_1d_artifacts_mutated": False,
            "publication_candidate_declared": False,
        }
    )
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    return log


def main() -> int:
    patch_base_paths()
    base.set_document_defaults = set_publication_defaults
    base.add_figure_placeholder = add_publication_figure_panel
    base.add_table_to_doc = add_publication_table_to_doc
    base.add_file_to_doc = add_file_to_publication_doc

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    base.assemble_markdown()
    docx_path, table_records, source_text = build_publication_docx()
    pdf_path = base.convert_docx_to_pdf(docx_path)
    rendered_pages = base.render_pdf_pages(pdf_path)
    log = base.write_log(docx_path, pdf_path, table_records, rendered_pages, source_text)
    log = augment_log(log)
    print(
        json.dumps(
            {
                "docx": str(docx_path),
                "pdf": str(pdf_path),
                "pages": log["pdf_page_count"],
                "tables": len(table_records),
                "rendered_pages": len(rendered_pages),
                "log": str(OUT_DIR / LOG_NAME),
                "publication_candidate_declared": False,
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
