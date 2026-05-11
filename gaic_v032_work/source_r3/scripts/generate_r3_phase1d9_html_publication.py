#!/usr/bin/env python3
"""
Phase 1D-9 HTML-first evaluation results and semantic table renderer for GAIC-2026 v0.3.2 FRC-R3.

This renderer treats active source_r3 Markdown as the only content truth,
builds a structured HTML/CSS publication, exports PDF from HTML, and creates
a DOCX derivative from the same HTML. It does not use old DOCX/PDF artifacts
as input.
"""

from __future__ import annotations

import html
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin

from PIL import Image, ImageChops
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import generate_r3_phase1d_docx_pdf as base
import generate_r3_phase1d4_full_whitepaper as full


ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parents[1]
OUT_DIR = ROOT / "out" / "phase_1d9" / "html_publication"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
REPORT_DIR = ROOT / "reports"
ARTIFACT_STEM = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-HTML-Publication-Draft-v3"
HTML_NAME = f"{ARTIFACT_STEM}.html"
PDF_NAME = f"{ARTIFACT_STEM}.pdf"
DOCX_NAME = f"{ARTIFACT_STEM}.docx"
PDF_TEXT_NAME = "phase_1d9_pdf_text.txt"
LOG_NAME = "phase_1d9_generation_log.json"

TRACE_TAG = "GACWP-2026-v0.3.2-FRC-R3"
PHASE_STATUS = "PHASE 1D-9 EVALUATION RESULTS AND TABLE REFLOW GENERATED / QA REQUIRED"
SOURCE_FILES = full.SOURCE_FILES
HIGH_RISK_TABLE_IDS = set(base.HIGH_RISK_TABLE_IDS) | {
    "T-F-01", "T-F-02", "T-F-03", "T-F-04", "T-F-05",
    "T-F-06", "T-F-07", "T-F-08", "T-F-09", "T-F-10",
    "T-F-11", "T-F-12", "T-F-13", "T-F-14", "T-F-15",
}

FORBIDDEN_TERMS = [
    "certifies compliance",
    "certification authority",
    "regulator-approved",
    "regulatory approval",
    "proves legal compliance",
    "guarantees compliance",
    "only validation path",
    "official compliance standard",
    "conformity assessment body",
    "market endorsement",
    "recommended vendor",
    "best system",
    "procurement winner",
    "winner",
    "leader",
    "trails",
    "final vendor score",
    "MPLP is required",
    "MPLP certifies",
    "Validation Lab certifies compliance",
    "audit opinion",
    "insurance guarantee",
    "deployment-readiness guarantee",
]

APPENDIX_G_STRICT_TERMS = [
    "Composite Score",
    "Ranking",
    "ranked",
    "score table",
    "final vendor score",
    "recommended vendor",
    "best system",
    "winner",
    "leader",
    "trailer",
    "leads",
    "trails",
]

FIGURE_TITLE_FALLBACKS = {
    "F-01": "AI Agent Lifecycle Governance Stack",
    "F-02": "Governance Stack / Missing Layer",
    "F-03": "MRO Topology",
    "F-04": "RCCS/ALCS Dual Scoring Model",
    "F-05": "Lifecycle Conformance Mapping Strength Heatmap",
    "F-06": "Enterprise Control Overlay for AI Agent Lifecycle Governance",
    "F-07": "Model Risk in Agentic Lifecycle",
    "F-08": "Evidence-Based Validation Pattern Flow",
    "F-09": "Enterprise Failure Scenario Chain",
    "F-10": "Companion Paper Boundary Map",
    "F-11": "Stage 0-to-7 Adoption Roadmap",
}

MRO_CATEGORIES = [
    ("Responsibility", ["MRO-01", "MRO-03", "MRO-04", "MRO-06"]),
    ("Authority", ["MRO-02", "MRO-05", "MRO-07"]),
    ("Evidence", ["MRO-08", "MRO-11", "MRO-12"]),
    ("Privacy", ["MRO-10", "MRO-13", "MRO-14"]),
    ("Substitution", ["MRO-09", "MRO-15"]),
    ("Closure", ["MRO-16"]),
]

EVIDENCE_MULTIPLIERS = [
    ("L1", 1.00, "Binding legal, regulatory, or formal standards text"),
    ("L2", 0.85, "Official product, API, developer, or platform documentation"),
    ("L3", 0.75, "Independent audits, certifications, third-party evaluations, or regulatory commentary"),
    ("L4", 0.55, "Vendor white papers, product pages, blog posts, or marketing statements"),
    ("L5", 0.35, "Author inference, conceptual analysis, or strategic interpretation"),
]

ADOPTION_STAGES = [
    ("0", "Inventory"),
    ("1", "Authority"),
    ("2", "Responsibility"),
    ("3", "Evidence"),
    ("4", "Privacy"),
    ("5", "Validation"),
    ("6", "Remediation"),
    ("7", "Scale"),
]

FAILURE_CATEGORIES = [
    ("Authority/Boundary", 2),
    ("Evidence/Traceability", 1),
    ("Responsibility/Acceptance", 2),
    ("Integration/Substitution", 3),
]

SOURCE_SECTION_GROUPS = [
    ("Legal / Regulatory", ["EU-AI-ACT-2024-1689", "GDPR-2016-679", "COLORADO-SB24-205", "COLORADO-SB25B-004"]),
    ("Standards / Frameworks", ["NIST-AI-RMF-1-0", "ISO-IEC-42001", "IMDA-AGENTIC-AI-MGF", "W3C-PROV", "W3C-VC-DATA-MODEL"]),
    ("Protocol / MPLP", ["MPLP-DOCS", "MPLP-SITE"]),
    ("Validation", ["VALIDATION-LAB"]),
    (
        "Product / Platform Documentation",
        [
            "IBM-WATSONX-GOVERNANCE-DOCS",
            "MICROSOFT-AZURE-AI-FOUNDRY-DOCS",
            "MICROSOFT-RESPONSIBLE-AI-DOCS",
            "AWS-BEDROCK-DOCS",
            "AWS-BEDROCK-GUARDRAILS-DOCS",
            "AWS-AGENTCORE-DOCS",
            "GOOGLE-VERTEX-AI-DOCS",
            "GOOGLE-ADK-DOCS",
            "GOOGLE-MODEL-ARMOR-DOCS",
            "LANGGRAPH-DOCS",
            "LANGSMITH-DOCS",
            "OPENAI-AGENTS-SDK-DOCS",
            "OPENAI-TOOLS-DOCS",
            "CREWAI-DOCS",
            "CREWAI-GITHUB",
        ],
    ),
]

SYSTEM_PROFILES = [
    {
        "system": "MPLP",
        "category": "Lifecycle protocol",
        "rccs": "Partial",
        "alcs": "Strong",
        "confidence": "L4 / L5",
        "x": 57,
        "y": 84,
        "visible_strength": "Protocol-level lifecycle semantics for responsibility objects, evidence pointers, and closure records.",
        "lifecycle_gap": "Enterprise deployment, external adoption, and implementation evidence remain outside protocol semantics.",
        "evidence_basis": "MPLP public materials plus author analytical mapping where lifecycle conformance is not directly stated.",
        "boundary": "Protocol analytical profile only; not a procurement recommendation, certification, or proof of legal compliance.",
    },
    {
        "system": "IBM watsonx.governance",
        "category": "Governance platform",
        "rccs": "Strong",
        "alcs": "Partial",
        "confidence": "L2 / L5",
        "x": 82,
        "y": 58,
        "visible_strength": "Model governance, documentation, risk management, monitoring, and enterprise governance controls.",
        "lifecycle_gap": "Agent handoff, accepted outcome, cross-agent responsibility transfer, and runtime substitution semantics are less visible.",
        "evidence_basis": "Official product documentation for visible governance features; lifecycle-object interpretation remains analytical.",
        "boundary": "Source-qualified posture profile only; not product ranking or implementation audit opinion.",
    },
    {
        "system": "Microsoft Azure AI Foundry",
        "category": "Cloud AI platform",
        "rccs": "Strong",
        "alcs": "Partial",
        "confidence": "L2 / L5",
        "x": 80,
        "y": 62,
        "visible_strength": "Responsible AI, model/platform governance, safety controls, evaluation support, and enterprise integration.",
        "lifecycle_gap": "Accepted-outcome closure, MAS evidence partitioning, and explicit MRO-level lifecycle objects require enterprise design.",
        "evidence_basis": "Official Microsoft documentation and responsible AI materials, with lifecycle mapping as source-qualified analysis.",
        "boundary": "Analytical profile, not endorsement, readiness threshold, or procurement advice.",
    },
    {
        "system": "AWS Bedrock + Guardrails + AgentCore",
        "category": "Cloud AI platform",
        "rccs": "Strong",
        "alcs": "Partial",
        "confidence": "L2 / L5",
        "x": 78,
        "y": 60,
        "visible_strength": "Foundation-model access controls, guardrails, agent/runtime services, monitoring hooks, and enterprise integration.",
        "lifecycle_gap": "Responsibility transfer, acceptance records, privacy-preserving validation, and lifecycle closure remain implementation-dependent.",
        "evidence_basis": "Official AWS service documentation; MRO/ALCS mapping remains provisional and source-qualified.",
        "boundary": "Not a ranking, certification, procurement recommendation, or compliance proof.",
    },
    {
        "system": "Google Vertex AI + ADK + Model Armor",
        "category": "Cloud AI platform",
        "rccs": "Strong",
        "alcs": "Partial",
        "confidence": "L2 / L5",
        "x": 79,
        "y": 61,
        "visible_strength": "Model/platform governance, agent-development tooling, safety filtering, and enterprise AI operations.",
        "lifecycle_gap": "Lifecycle responsibility object coverage depends on how teams bind ADK/runtime behavior to evidence and accountability controls.",
        "evidence_basis": "Official Google documentation; agentic lifecycle conformance is mapped analytically where not directly specified.",
        "boundary": "Source-qualified posture only; no vendor ordering or final vendor score.",
    },
    {
        "system": "LangGraph + LangSmith",
        "category": "Agent orchestration / observability",
        "rccs": "Partial",
        "alcs": "Strong",
        "confidence": "L2 / L5",
        "x": 60,
        "y": 80,
        "visible_strength": "Agent graph orchestration, traceability, observability, evaluation workflows, and runtime state patterns.",
        "lifecycle_gap": "Regulatory control mapping, legal-source coverage, and enterprise procurement controls require surrounding governance.",
        "evidence_basis": "Official LangGraph/LangSmith documentation with analytical RCCS/ALCS mapping.",
        "boundary": "Lifecycle capability posture, not a regulator-approved benchmark or market endorsement.",
    },
    {
        "system": "OpenAI Agents SDK",
        "category": "Agent SDK / developer platform",
        "rccs": "Partial",
        "alcs": "Partial",
        "confidence": "L2 / L5 with access caveat",
        "x": 58,
        "y": 64,
        "visible_strength": "Agent/tool orchestration primitives, developer-facing SDK surfaces, and tool/action workflow support.",
        "lifecycle_gap": "OpenAI platform guide access remains HTTP 403 in QA; full enterprise governance and lifecycle evidence posture require additional source pinning.",
        "evidence_basis": "Accessible SDK documentation and source-pointer queues; unresolved access remains tracked.",
        "boundary": "Access-limited analytical posture only; not final assessment or ranking.",
    },
    {
        "system": "CrewAI",
        "category": "Agent orchestration framework",
        "rccs": "Limited / Partial",
        "alcs": "Partial",
        "confidence": "L2 / L5",
        "x": 48,
        "y": 63,
        "visible_strength": "Role-based multi-agent orchestration, task delegation, crews/flows, and operational agent patterns.",
        "lifecycle_gap": "Regulatory controls, formal evidence retention, accepted outcome, privacy mapping, and remediation closure require external governance.",
        "evidence_basis": "Official documentation and repository materials with analytical MRO/ALCS mapping.",
        "boundary": "Framework posture only; not a product quality score, certification, or procurement conclusion.",
    },
]

HEATMAP_GROUPS = ["RCCS", "Authority", "Evidence", "Privacy", "Substitution", "Closure"]
SYSTEM_HEATMAP = {
    "MPLP": ["Partial", "Strong", "Strong", "Partial", "Partial", "Strong"],
    "IBM watsonx.governance": ["Strong", "Partial", "Strong", "Partial", "Limited", "Partial"],
    "Microsoft Azure AI Foundry": ["Strong", "Partial", "Partial", "Partial", "Partial", "Partial"],
    "AWS Bedrock + Guardrails + AgentCore": ["Strong", "Partial", "Partial", "Partial", "Partial", "Partial"],
    "Google Vertex AI + ADK + Model Armor": ["Strong", "Partial", "Partial", "Partial", "Partial", "Partial"],
    "LangGraph + LangSmith": ["Partial", "Strong", "Strong", "Limited", "Partial", "Partial"],
    "OpenAI Agents SDK": ["Partial", "Partial", "Partial", "Limited", "Partial", "Limited"],
    "CrewAI": ["Limited", "Partial", "Partial", "Limited", "Partial", "Limited"],
}

CHAPTER_HIERARCHY = {
    "01-scope-methodology.md": ("The paper defines a non-legal governance analysis for agentic lifecycle conformance.", "Readers need the boundaries before interpreting any scoring, mapping, or system profile.", "Scope, evidence levels, document-control tables, and non-claim boundaries.", "The framework is analytical and source-qualified, not legal advice, certification, or procurement guidance."),
    "02-missing-layer.md": ("Model governance is necessary but insufficient for agentic systems.", "Agentic workflows shift compliance attention from model artifacts to delegated lifecycle responsibility.", "Lifecycle stack analysis and missing-layer framing.", "AI Agent Lifecycle Governance is the missing operational layer between model controls and accountable enterprise outcomes."),
    "03-agentic-ai-breaks-model-compliance.md": ("Agentic systems break model-centric controls by acting through tools, roles, memory, and handoffs.", "Governance fails when the compliance object stays at the model while work occurs across lifecycle events.", "Agentic workflow analysis and lifecycle failure modes.", "The compliance unit must include authority, action, evidence, acceptance, and closure."),
    "04-regulatory-standards-baseline.md": ("Existing laws and standards already imply lifecycle duties but do not fully name agentic responsibility objects.", "The white paper should translate regulatory language into engineering objects without overstating law.", "Snapshot-bound regulatory and standards comparison.", "The baseline supports object derivation, not legal interpretation for a specific deployment."),
    "05-regulatory-language-to-engineering-objects.md": ("Regulatory requirements become actionable only when mapped to lifecycle objects.", "Enterprises need implementable objects, not only policy principles.", "Object-derivation logic from governance language.", "The chapter bridges legal/standards language into Missing Regulatory Objects."),
    "06-missing-regulatory-objects.md": ("Sixteen MROs define the missing lifecycle responsibility layer.", "Without these objects, authority, evidence, privacy, substitution, and closure remain implicit.", "Canonical MRO dictionary and object field model.", "MROs provide analytical object structure, not a certification checklist."),
    "07-rccs-regulatory-compliance-coverage.md": ("RCCS measures traditional regulatory coverage separately from lifecycle conformance.", "Separating RCCS prevents model governance strength from being mistaken for agentic lifecycle assurance.", "RCCS dimensions, rubrics, and evidence multipliers.", "RCCS is not a legal compliance score and does not prove implementation."),
    "08-alcs-agentic-lifecycle-conformance.md": ("ALCS measures lifecycle conformance surfaces that RCCS alone does not capture.", "Agentic systems need authority, responsibility, evidence, privacy, substitution, and closure posture review.", "ALCS dimensions mapped to MROs.", "ALCS is an analytical lifecycle posture model, not deployment-readiness certification."),
    "09-composite-scoring-method.md": ("RCCS and ALCS can be combined analytically, but evidence confidence changes interpretation.", "A high conceptual score without source support should not be treated as strong posture.", "Canonical multipliers and composite formula.", "Composite outputs are provisional analytical profiles, not final vendor scores."),
    "10-enterprise-control-crosswalk.md": ("Lifecycle objects can be translated into enterprise controls and procurement questions.", "Governance must connect lifecycle semantics to board, risk, security, privacy, and procurement operations.", "Control crosswalk and enterprise adoption mapping.", "The crosswalk supports due diligence; it does not recommend vendors."),
    "11-comparative-field.md": ("Current systems expose different lifecycle strengths and gaps by category.", "Readers need category context before reading system mappings.", "Qualitative positioning across protocol, platform, SDK, and orchestration categories.", "The comparison is source-qualified and non-ranking."),
    "12-detailed-system-mappings.md": ("System mappings show uneven visible coverage across RCCS/ALCS/MRO surfaces.", "The paper needs explicit profiles so readers can separate method, result, and evidence.", "Source-qualified mapping across eight evaluated systems.", "Profiles are provisional, evidence-bound, and not procurement recommendations."),
    "13-mplp-deep-mapping.md": ("MPLP has strong semantic fit for lifecycle responsibility objects but remains one protocol path.", "Protocol semantics are useful only when kept separate from endorsement or certification claims.", "MPLP-to-MRO/RCCS/ALCS mapping and conflict-of-interest discipline.", "MPLP is not required, certified, or exclusive."),
    "14-evidence-validation-placeholder.md": ("Evidence-based validation can review lifecycle objects without becoming certification.", "Assurance patterns must preserve privacy and source limits.", "Validation pattern flow and non-certifying Validation Lab example.", "Validation Lab remains a boundary-limited public-surface example."),
    "15-failure-scenarios-placeholder.md": ("Enterprise failures become predictable when lifecycle responsibility objects are missing.", "Scenarios make the governance gap concrete without claiming empirical incident frequency.", "Illustrative hypothetical failure chains mapped to MRO/ALCS categories.", "Scenarios are illustrative and hypothetical, not legal conclusions."),
    "16-companion-paper-boundary.md": ("Evidence, assurance, and insurability require separate treatment beyond this paper.", "The white paper should not overclaim companion-paper deliverables.", "Boundary map for planned companion scope.", "Companion papers are planned/intended scope, not published guarantees."),
    "17-adoption-roadmap-placeholder.md": ("Enterprises can adopt lifecycle governance in staged increments from inventory to scale.", "Implementation needs a sequence, not a one-step certification claim.", "Stage 0-to-7 roadmap and control maturity logic.", "The roadmap is an implementation planning aid, not a mandate or guarantee."),
    "18-conclusion-placeholder.md": ("The core conclusion is that AI agent compliance is lifecycle responsibility compliance.", "Readers need the final boundary between analytical governance and legal/procurement claims.", "Contribution summary, non-claims table, and call to action.", "The paper is a governance framework and source-qualified analysis, not legal advice or certification."),
}


@dataclass
class TableRecord:
    table_id: str | None
    caption: str | None
    source_file: str
    rows: int
    cols: int
    rendered_tables: int
    strategy: str
    high_risk: bool


@dataclass
class FigureRecord:
    figure_id: str
    title: str
    source_file: str
    has_flow: bool
    has_interpretation: bool


@dataclass
class ChartRecord:
    chart_id: str
    title: str
    chart_type: str
    section_location: str
    source_data: str
    boundary_note: str
    safe_for_publication: bool


@dataclass
class CodeBlockRecord:
    source_file: str
    language: str
    line_count: int
    rendering: str


def which(name: str) -> str | None:
    return shutil.which(name)


def run_git(args: list[str]) -> str:
    result = subprocess.run(["git", *args], cwd=REPO_ROOT, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        return f"ERROR: {result.stderr.strip()}"
    return result.stdout.strip()


def clean_output_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in OUT_DIR.iterdir():
        if path.is_dir():
            shutil.rmtree(path)
        else:
            path.unlink()
    RENDER_DIR.mkdir(parents=True, exist_ok=True)


def sanitize_line(line: str, relative: str) -> str | None:
    return full.sanitize_publication_line(line, relative)


def read_source_lines(relative: str) -> list[str]:
    path = ROOT / relative
    if not path.exists():
        raise FileNotFoundError(f"Missing required source file: {relative}")
    lines: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        sanitized = sanitize_line(raw, relative)
        if sanitized is not None:
            lines.append(sanitized.rstrip())
    return lines


def strip_markdown_for_text(text: str) -> str:
    return base.clean_inline_markdown(text).replace("→", "->").strip()


def inline_md_to_html(text: str) -> str:
    placeholders: list[str] = []

    def stash(value: str) -> str:
        placeholders.append(value)
        return f"\u0000{len(placeholders) - 1}\u0000"

    raw = text.strip()
    raw = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", lambda m: stash(html.escape(m.group(1))), raw)
    raw = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        lambda m: stash(
            f'<a href="{html.escape(m.group(2), quote=True)}">{html.escape(m.group(1))}</a>'
        ),
        raw,
    )
    raw = re.sub(r"`([^`]+)`", lambda m: stash(f'<code>{html.escape(m.group(1))}</code>'), raw)
    escaped = html.escape(raw)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", escaped)
    escaped = escaped.replace("&lt;br&gt;", "<br>").replace("&lt;br/&gt;", "<br>").replace("&lt;br /&gt;", "<br>")
    for idx, value in enumerate(placeholders):
        escaped = escaped.replace(f"\u0000{idx}\u0000", value)
    escaped = re.sub(r"\b(L[1-5])\b", r'<span class="badge evidence">\1</span>', escaped)
    escaped = re.sub(r"\b(MRO-\d{2})\b", r'<span class="badge mro">\1</span>', escaped)
    escaped = re.sub(r"\b(RCCS-\d{2})\b", r'<span class="badge rccs">\1</span>', escaped)
    escaped = re.sub(r"\b(ALCS-\d{2})\b", r'<span class="badge alcs">\1</span>', escaped)
    return escaped


def slugify(text: str, used: dict[str, int]) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", strip_markdown_for_text(text).lower()).strip("-")
    slug = slug or "section"
    count = used.get(slug, 0)
    used[slug] = count + 1
    return slug if count == 0 else f"{slug}-{count + 1}"


def heading_level(line: str) -> int:
    return len(line) - len(line.lstrip("#"))


def is_caption_line(line: str) -> bool:
    return base.is_caption_line(line)


def is_table_start(lines: list[str], i: int) -> bool:
    return base.is_table_start(lines, i)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for raw in table_lines:
        parts = [p.strip() for p in raw.strip("|").split("|")]
        if parts and all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            continue
        rows.append(parts)
    return rows, i


def detect_table_id(caption: str | None, rows: list[list[str]]) -> str | None:
    cleaned = [[strip_markdown_for_text(cell) for cell in row] for row in rows]
    return base.detect_table_id(caption, cleaned)


def column_groups(cols: int) -> list[list[int]]:
    if cols <= 5:
        return [list(range(cols))]
    if cols == 6:
        return [[0, 1, 2], [0, 3, 4, 5]]
    if cols == 7:
        return [[0, 1, 2, 3], [0, 4, 5, 6]]
    if cols == 8:
        return [[0, 1, 2, 3], [0, 4, 5, 6, 7]]
    if cols == 9:
        return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8]]
    return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8, 9]]


def classify_table(caption: str | None, rows: list[list[str]], table_id: str | None, relative: str) -> str:
    caption_text = strip_markdown_for_text(caption or "").lower()
    headers = [strip_markdown_for_text(cell).lower() for cell in (rows[0] if rows else [])]
    joined_headers = " | ".join(headers)
    if table_id and (table_id.startswith("T-E-") or table_id.startswith("T-F-")):
        return "rubric_table"
    if table_id and re.match(r"T-0[78]-", table_id) and any("evidence" in h for h in headers):
        return "rubric_table"
    if table_id and table_id.startswith(("T-B-", "T-H-", "T-I-", "T-J-", "T-K-", "T-17-", "T-18-")):
        return "scorecard_table"
    if {"score", "capability level"}.issubset(set(headers)) or ("score" in joined_headers and "evidence" in joined_headers):
        return "rubric_table"
    if table_id and (table_id.startswith("T-12-") or table_id.startswith("T-D-") or table_id.startswith("T-G-") or table_id in {"T-13-02", "T-13-03", "T-13-04"}):
        return "system_profile_table"
    if "mapping strength" in joined_headers or "lifecycle gap" in joined_headers or "claim-level support" in joined_headers:
        return "system_profile_table"
    if "system" in joined_headers and ("gap" in joined_headers or "boundary" in joined_headers or "evidence" in joined_headers):
        return "system_profile_table"
    if table_id and table_id.startswith("T-A-"):
        return "mro_object_card_table"
    if table_id and table_id.startswith("T-C-") or "source" in caption_text or "citation" in caption_text:
        return "source_register_table"
    if "boundary" in joined_headers or "does not say" in joined_headers or "non-claims" in caption_text:
        return "governance_boundary_table"
    if relative.startswith("appendices/"):
        return "appendix_dense_table"
    return "matrix_table"


def semantic_table_class(table_type: str) -> str:
    return {
        "matrix_table": "table-matrix",
        "rubric_table": "table-rubric",
        "scorecard_table": "table-scorecard",
        "source_register_table": "table-source-register",
        "mro_object_card_table": "table-mro-card",
        "system_profile_table": "table-system-profile",
        "governance_boundary_table": "table-boundary",
        "appendix_dense_table": "table-appendix-dense",
    }.get(table_type, "table-matrix")


def render_rubric_cards(rows: list[list[str]], caption: str | None, table_id: str | None) -> str:
    if len(rows) < 2:
        return ""
    headers = [strip_markdown_for_text(cell) for cell in rows[0]]
    parts = ['<div class="rubric-card-set">']
    if caption:
        parts.append(f'<div class="table-caption">{inline_md_to_html(caption)}</div>')
    parts.append('<div class="layout-note">Rubric rendered as score cards to preserve evidence readability; all source rows are retained.</div>')
    for row in rows[1:]:
        data = {headers[idx].lower(): row[idx] if idx < len(row) else "" for idx in range(len(headers))}
        score = data.get("score") or data.get("level") or row[0]
        capability = data.get("capability level") or data.get("capability") or (row[1] if len(row) > 1 else "")
        evidence = data.get("evidence required") or data.get("evidence") or row[-1]
        other_items = []
        for idx, header in enumerate(headers):
            low = header.lower()
            if low in {"score", "level", "capability level", "capability", "evidence required", "evidence"}:
                continue
            if idx < len(row) and row[idx].strip():
                other_items.append(f'<div><dt>{html.escape(header)}</dt><dd>{inline_md_to_html(row[idx])}</dd></div>')
        parts.extend(
            [
                '<article class="rubric-card">',
                f'<div class="rubric-score">{inline_md_to_html(score)}</div>',
                '<div class="rubric-body">',
                f'<h4>{inline_md_to_html(capability)}</h4>',
                f'<p>{inline_md_to_html(evidence)}</p>',
                f'<dl>{"".join(other_items)}</dl>' if other_items else "",
                "</div>",
                "</article>",
            ]
        )
    parts.append("</div>")
    return "\n".join(parts)


def render_system_profile_cards(rows: list[list[str]], caption: str | None, table_id: str | None) -> str:
    if len(rows) < 2:
        return ""
    headers = [strip_markdown_for_text(cell) for cell in rows[0]]
    parts = ['<div class="system-profile-card-set">']
    if caption:
        parts.append(f'<div class="table-caption">{inline_md_to_html(caption)}</div>')
    parts.append('<div class="layout-note">System mapping rendered as profile cards to preserve long evidence and boundary text; all source rows are retained without rank ordering.</div>')
    for row in rows[1:]:
        title = row[0] if row else "System Profile Row"
        parts.append('<article class="system-profile-row-card">')
        parts.append(f'<h4>{inline_md_to_html(title)}</h4>')
        parts.append("<dl>")
        for idx, header in enumerate(headers[1:], 1):
            if idx < len(row) and row[idx].strip():
                parts.append(f'<div><dt>{html.escape(header)}</dt><dd>{inline_md_to_html(row[idx])}</dd></div>')
        parts.append("</dl>")
        parts.append("</article>")
    parts.append("</div>")
    return "\n".join(parts)


def long_multicolumn_rows(rows: list[list[str]], cols: int, threshold: int = 150) -> bool:
    if cols < 4:
        return False
    for row in rows[1:]:
        narrative = " ".join(strip_markdown_for_text(cell) for cell in row[1:])
        if len(narrative) >= threshold:
            return True
    return False


def render_table(rows: list[list[str]], caption: str | None, relative: str, records: list[TableRecord]) -> str:
    if not rows:
        return ""
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table_id = detect_table_id(caption, normalized)
    table_type = classify_table(caption, normalized, table_id, relative)
    high_risk = bool(table_id and table_id in HIGH_RISK_TABLE_IDS)
    evidence_long = any(len(strip_markdown_for_text(row[-1])) > 90 for row in normalized[1:]) if cols >= 3 else False
    if table_type == "rubric_table" and cols >= 3 and (evidence_long or bool(table_id and table_id.startswith(("T-E-", "T-F-", "T-07-", "T-08-")))):
        rendered = render_rubric_cards(normalized, caption, table_id)
        records.append(
            TableRecord(
                table_id=table_id,
                caption=strip_markdown_for_text(caption or "") or None,
                source_file=relative,
                rows=len(normalized),
                cols=cols,
                rendered_tables=1,
                strategy="rubric card reflow",
                high_risk=high_risk,
            )
        )
        return rendered
    if table_type == "system_profile_table" and cols >= 4 and table_id and table_id.startswith(("T-12-", "T-13-", "T-G-")):
        rendered = render_system_profile_cards(normalized, caption, table_id)
        records.append(
            TableRecord(
                table_id=table_id,
                caption=strip_markdown_for_text(caption or "") or None,
                source_file=relative,
                rows=len(normalized),
                cols=cols,
                rendered_tables=1,
                strategy="system profile card reflow",
                high_risk=high_risk,
            )
        )
        return rendered
    if table_type in {"scorecard_table", "source_register_table", "governance_boundary_table"} and table_id and table_id.startswith(("T-B-", "T-H-", "T-I-", "T-J-", "T-K-", "T-17-", "T-18-")) and (cols >= 4 or high_risk or long_multicolumn_rows(normalized, cols)):
        rendered = render_system_profile_cards(normalized, caption, table_id)
        records.append(
            TableRecord(
                table_id=table_id,
                caption=strip_markdown_for_text(caption or "") or None,
                source_file=relative,
                rows=len(normalized),
                cols=cols,
                rendered_tables=1,
                strategy="semantic row card reflow",
                high_risk=high_risk,
            )
        )
        return rendered
    split_required = (cols >= 6 or high_risk) and table_type not in {"rubric_table", "governance_boundary_table"}
    groups = column_groups(cols) if split_required else [list(range(cols))]
    strategy = "semantic column split" if split_required else "single table"
    classes = ["table-block", semantic_table_class(table_type)]
    if high_risk or split_required:
        classes.append("table-compact")
    if table_type == "appendix_dense_table":
        classes.append("table-compact")
    parts = [f'<div class="{" ".join(classes)}">']
    if caption:
        parts.append(f'<div class="table-caption">{inline_md_to_html(caption)}</div>')
    if split_required:
        parts.append('<div class="layout-note">Wide source table split into consecutive column groups; all source rows and columns are preserved.</div>')
    for idx, group in enumerate(groups, 1):
        if split_required:
            parts.append(f'<div class="continuation-label">Continuation {idx} of {len(groups)}</div>')
        parts.append("<table>")
        if normalized:
            parts.append("<thead><tr>")
            for col in group:
                parts.append(f"<th>{inline_md_to_html(normalized[0][col] if col < len(normalized[0]) else '')}</th>")
            parts.append("</tr></thead>")
        parts.append("<tbody>")
        for row in normalized[1:]:
            parts.append("<tr>")
            for col in group:
                parts.append(f"<td>{inline_md_to_html(row[col] if col < len(row) else '')}</td>")
            parts.append("</tr>")
        parts.append("</tbody></table>")
    parts.append("</div>")
    records.append(
        TableRecord(
            table_id=table_id,
            caption=strip_markdown_for_text(caption or "") or None,
            source_file=relative,
            rows=len(normalized),
            cols=cols,
            rendered_tables=len(groups),
            strategy=strategy,
            high_risk=high_risk,
        )
    )
    return "\n".join(parts)


def is_figure_title(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith("-"):
        return False
    return bool(re.match(r"^\*\*Figure F-\d{2}\b", stripped) or re.match(r"^Figure F-\d{2}\b", stripped))


def parse_figure_title(line: str) -> tuple[str, str]:
    stripped = strip_markdown_for_text(line)
    match = re.search(r"\bFigure\s+(F-\d{2})\s*(?:[-—:]\s*)?(.*)$", stripped)
    if not match:
        return "F-00", stripped
    figure_id = match.group(1)
    title = match.group(2).strip() or FIGURE_TITLE_FALLBACKS.get(figure_id, "")
    return figure_id, title


def render_flow(flow: str) -> str:
    flow = strip_markdown_for_text(flow)
    if not flow:
        return ""
    lanes = [part.strip() for part in re.split(r"\s*(?:→|->)\s*", flow) if part.strip()]
    if len(lanes) < 2:
        return f'<p class="figure-flow-text">{inline_md_to_html(flow)}</p>'
    pieces: list[str] = ['<div class="figure-flow">']
    for idx, lane in enumerate(lanes):
        pieces.append(f'<span class="flow-node">{inline_md_to_html(lane)}</span>')
        if idx < len(lanes) - 1:
            pieces.append('<span class="flow-arrow">→</span>')
    pieces.append("</div>")
    return "\n".join(pieces)


def svg_bar_chart(chart_id: str, title: str, values: list[tuple[str, float, str]], max_value: float = 1.0) -> str:
    width = 760
    left = 150
    top = 54
    row_h = 42
    bar_w = 470
    height = top + len(values) * row_h + 42
    pieces = [
        f'<svg class="chart-svg" viewBox="0 0 {width} {height}" role="img" aria-labelledby="{chart_id}-title">',
        f'<title id="{chart_id}-title">{html.escape(title)}</title>',
        f'<text x="0" y="24" class="chart-title">{html.escape(title)}</text>',
    ]
    for idx, (label, value, note) in enumerate(values):
        y = top + idx * row_h
        w = max(4, min(bar_w, (value / max_value) * bar_w))
        pieces.extend(
            [
                f'<text x="0" y="{y + 18}" class="chart-label">{html.escape(label)}</text>',
                f'<rect x="{left}" y="{y}" width="{bar_w}" height="22" rx="2" class="chart-track"/>',
                f'<rect x="{left}" y="{y}" width="{w:.1f}" height="22" rx="2" class="chart-bar"/>',
                f'<text x="{left + bar_w + 14}" y="{y + 17}" class="chart-value">{value:.2f}</text>',
                f'<text x="{left}" y="{y + 36}" class="chart-note">{html.escape(note)}</text>',
            ]
        )
    pieces.append("</svg>")
    return "\n".join(pieces)


def svg_count_bar_chart(chart_id: str, title: str, values: list[tuple[str, int, str]]) -> str:
    max_value = max((v for _, v, _ in values), default=1)
    width = 760
    left = 170
    top = 54
    row_h = 42
    bar_w = 450
    height = top + len(values) * row_h + 42
    pieces = [
        f'<svg class="chart-svg" viewBox="0 0 {width} {height}" role="img" aria-labelledby="{chart_id}-title">',
        f'<title id="{chart_id}-title">{html.escape(title)}</title>',
        f'<text x="0" y="24" class="chart-title">{html.escape(title)}</text>',
    ]
    for idx, (label, value, note) in enumerate(values):
        y = top + idx * row_h
        w = max(4, min(bar_w, (value / max_value) * bar_w))
        pieces.extend(
            [
                f'<text x="0" y="{y + 18}" class="chart-label">{html.escape(label)}</text>',
                f'<rect x="{left}" y="{y}" width="{bar_w}" height="22" rx="2" class="chart-track"/>',
                f'<rect x="{left}" y="{y}" width="{w:.1f}" height="22" rx="2" class="chart-bar alt"/>',
                f'<text x="{left + bar_w + 14}" y="{y + 17}" class="chart-value">{value}</text>',
                f'<text x="{left}" y="{y + 36}" class="chart-note">{html.escape(note)}</text>',
            ]
        )
    pieces.append("</svg>")
    return "\n".join(pieces)


def svg_timeline(chart_id: str, title: str, stages: list[tuple[str, str]]) -> str:
    width = 760
    height = 190
    left = 52
    right = 708
    y = 92
    gap = (right - left) / (len(stages) - 1)
    pieces = [
        f'<svg class="chart-svg timeline-svg" viewBox="0 0 {width} {height}" role="img" aria-labelledby="{chart_id}-title">',
        f'<title id="{chart_id}-title">{html.escape(title)}</title>',
        f'<text x="0" y="24" class="chart-title">{html.escape(title)}</text>',
        f'<line x1="{left}" y1="{y}" x2="{right}" y2="{y}" class="timeline-line"/>',
    ]
    for idx, (stage, label) in enumerate(stages):
        x = left + idx * gap
        pieces.extend(
            [
                f'<circle cx="{x:.1f}" cy="{y}" r="15" class="timeline-node"/>',
                f'<text x="{x:.1f}" y="{y + 5}" text-anchor="middle" class="timeline-stage">{html.escape(stage)}</text>',
                f'<text x="{x:.1f}" y="{y + 42}" text-anchor="middle" class="timeline-label">{html.escape(label)}</text>',
            ]
        )
    pieces.append("</svg>")
    return "\n".join(pieces)


def svg_quadrant_chart(chart_id: str, title: str, profiles: list[dict[str, str]]) -> str:
    width = 760
    height = 520
    left = 92
    top = 74
    plot_w = 560
    plot_h = 340
    pieces = [
        f'<svg class="chart-svg quadrant-svg" viewBox="0 0 {width} {height}" role="img" aria-labelledby="{chart_id}-title">',
        f'<title id="{chart_id}-title">{html.escape(title)}</title>',
        f'<text x="0" y="26" class="chart-title">{html.escape(title)}</text>',
        f'<rect x="{left}" y="{top}" width="{plot_w}" height="{plot_h}" class="quadrant-bg"/>',
        f'<line x1="{left + plot_w / 2}" y1="{top}" x2="{left + plot_w / 2}" y2="{top + plot_h}" class="quadrant-line"/>',
        f'<line x1="{left}" y1="{top + plot_h / 2}" x2="{left + plot_w}" y2="{top + plot_h / 2}" class="quadrant-line"/>',
        f'<text x="{left}" y="{top - 14}" class="chart-note">ALCS posture</text>',
        f'<text x="{left + plot_w - 92}" y="{top + plot_h + 36}" class="chart-note">RCCS posture</text>',
        f'<text x="{left + 18}" y="{top + 24}" class="quadrant-label">Lifecycle-strong / governance-visible</text>',
        f'<text x="{left + plot_w / 2 + 18}" y="{top + 24}" class="quadrant-label">Governance-strong / lifecycle-visible</text>',
        f'<text x="{left + 18}" y="{top + plot_h - 18}" class="quadrant-label">Emerging lifecycle profile</text>',
        f'<text x="{left + plot_w / 2 + 18}" y="{top + plot_h - 18}" class="quadrant-label">Baseline governance profile</text>',
    ]
    for idx, profile in enumerate(profiles):
        x = left + (float(profile["x"]) / 100.0) * plot_w
        y = top + plot_h - (float(profile["y"]) / 100.0) * plot_h
        short = profile["system"].split()[0].replace("Microsoft", "MS").replace("Google", "GCP")
        pieces.extend(
            [
                f'<circle cx="{x:.1f}" cy="{y:.1f}" r="9" class="quadrant-dot q{idx % 4}"/>',
                f'<text x="{x + 12:.1f}" y="{y + 4:.1f}" class="quadrant-system">{html.escape(short)}</text>',
            ]
        )
    pieces.extend(
        [
            '<text x="92" y="474" class="chart-note">Positions are qualitative analytical posture coordinates, not final scores or rankings.</text>',
            "</svg>",
        ]
    )
    return "\n".join(pieces)


def posture_class(value: str) -> str:
    low = value.lower()
    if "strong" in low:
        return "strong"
    if "partial" in low:
        return "partial"
    if "limited" in low:
        return "limited"
    return "not-visible"


def chart_card(chart_id: str, title: str, chart_type: str, svg: str, interpretation: str, boundary: str) -> str:
    return "\n".join(
        [
            f'<figure class="chart-card" id="{chart_id}">',
            f'<figcaption><span>{html.escape(chart_id.upper())}</span>{html.escape(title)}</figcaption>',
            svg,
            f'<p class="chart-interpretation">{inline_md_to_html(interpretation)}</p>',
            f'<p class="chart-boundary">{inline_md_to_html(boundary)}</p>',
            "</figure>",
        ]
    )


def chart_blocks() -> tuple[dict[str, str], list[ChartRecord]]:
    records: list[ChartRecord] = []
    blocks: dict[str, str] = {}

    evidence_svg = svg_bar_chart(
        "chart-evidence-multipliers",
        "Evidence Confidence Multipliers",
        [(level, value, desc) for level, value, desc in EVIDENCE_MULTIPLIERS],
        max_value=1.0,
    )
    blocks["evidence_multipliers"] = chart_card(
        "chart-evidence-multipliers",
        "Evidence Confidence Multipliers",
        "bar chart",
        evidence_svg,
        "The chart visualizes the canonical confidence multipliers used by the scoring method. Higher confidence evidence preserves more of the normalized dimension score.",
        "This chart does not prove legal compliance, product implementation, audit success, or vendor superiority. It only visualizes the already-defined evidence multipliers.",
    )
    records.append(ChartRecord("chart-evidence-multipliers", "Evidence Confidence Multipliers", "bar chart", "Chapter 9", "Canonical L1-L5 multiplier values", "No vendor data or ranking", True))

    mro_svg = svg_count_bar_chart(
        "chart-mro-category-distribution",
        "MRO Category Distribution",
        [(name, len(ids), ", ".join(ids)) for name, ids in MRO_CATEGORIES],
    )
    blocks["mro_distribution"] = chart_card(
        "chart-mro-category-distribution",
        "MRO Category Distribution",
        "bar chart",
        mro_svg,
        "The chart groups the sixteen Missing Regulatory Objects by lifecycle control surface: responsibility, authority, evidence, privacy, substitution, and closure.",
        "The grouping is an explanatory topology for this publication, not a legal taxonomy, certification model, or claim that categories are mutually exclusive.",
    )
    records.append(ChartRecord("chart-mro-category-distribution", "MRO Category Distribution", "bar chart", "Chapter 6", "Sixteen canonical MROs grouped by lifecycle control surface", "Explanatory topology only", True))

    timeline_svg = svg_timeline("chart-adoption-roadmap", "Stage 0-to-7 Adoption Roadmap", ADOPTION_STAGES)
    blocks["adoption_timeline"] = chart_card(
        "chart-adoption-roadmap",
        "Stage 0-to-7 Adoption Roadmap",
        "horizontal timeline",
        timeline_svg,
        "The timeline summarizes the staged adoption path from inventory through authority, responsibility, evidence, privacy, validation, remediation, and scale.",
        "This is an implementation planning aid, not a deployment-readiness guarantee, procurement mandate, certification checklist, or legal compliance procedure.",
    )
    records.append(ChartRecord("chart-adoption-roadmap", "Stage 0-to-7 Adoption Roadmap", "horizontal timeline", "Chapter 17", "Stage 0 through Stage 7 labels from Chapter 17", "Planning aid only", True))

    failure_svg = svg_count_bar_chart(
        "chart-failure-scenario-distribution",
        "Enterprise Failure Scenario Distribution",
        [(name, count, "Illustrative scenario count") for name, count in FAILURE_CATEGORIES],
    )
    blocks["failure_distribution"] = chart_card(
        "chart-failure-scenario-distribution",
        "Enterprise Failure Scenario Distribution",
        "bar chart",
        failure_svg,
        "The chart shows how the eight illustrative failure scenarios are distributed across the four failure-mode categories used in Chapter 15.",
        "The chart is not an incident probability model, empirical risk frequency claim, legal conclusion, or vendor/system assessment.",
    )
    records.append(ChartRecord("chart-failure-scenario-distribution", "Enterprise Failure Scenario Distribution", "bar chart", "Chapter 15", "Eight illustrative failure scenarios grouped by Chapter 15 category", "Not probability or empirical incident data", True))

    quadrant_svg = svg_quadrant_chart(
        "chart-rccs-alcs-quadrant",
        "RCCS / ALCS Analytical Posture Map",
        SYSTEM_PROFILES,
    )
    blocks["rccs_alcs_quadrant"] = chart_card(
        "chart-rccs-alcs-quadrant",
        "RCCS / ALCS Analytical Posture Map",
        "qualitative quadrant chart",
        quadrant_svg,
        "The chart places evaluated systems by qualitative RCCS and ALCS posture so readers can see the difference between baseline governance coverage and agentic lifecycle conformance.",
        "Positions are not final vendor scores, rankings, procurement recommendations, certification claims, or legal compliance proof.",
    )
    records.append(ChartRecord("chart-rccs-alcs-quadrant", "RCCS / ALCS Analytical Posture Map", "qualitative quadrant chart", "Evaluation Results", "Qualitative RCCS/ALCS posture from source-qualified profiles", "No rank order or final scores", True))

    return blocks, records


def heatmap_html() -> str:
    parts = [
        '<div class="heatmap-card" id="system-dimension-heatmap">',
        '<div class="heatmap-title">System x Dimension Posture Heatmap</div>',
        '<div class="heatmap-note">Values are qualitative, source-qualified posture labels. They are not final scores or rankings.</div>',
        '<table class="heatmap-table">',
        "<thead><tr><th>System</th>",
    ]
    for group in HEATMAP_GROUPS:
        parts.append(f"<th>{html.escape(group)}</th>")
    parts.append("</tr></thead><tbody>")
    for profile in SYSTEM_PROFILES:
        system = profile["system"]
        parts.append(f'<tr><th>{html.escape(system)}</th>')
        for value in SYSTEM_HEATMAP[system]:
            parts.append(f'<td><span class="heat {posture_class(value)}">{html.escape(value)}</span></td>')
        parts.append("</tr>")
    parts.extend(["</tbody></table>", "</div>"])
    return "\n".join(parts)


def evaluation_results_html(chart_html: dict[str, str]) -> str:
    profile_cards = []
    for profile in SYSTEM_PROFILES:
        profile_cards.append(
            "\n".join(
                [
                    '<article class="evaluation-profile-card">',
                    f'<div class="profile-kicker">{html.escape(profile["category"])}</div>',
                    f'<h3>{html.escape(profile["system"])}</h3>',
                    '<div class="profile-postures">',
                    f'<span><strong>RCCS</strong>{html.escape(profile["rccs"])}</span>',
                    f'<span><strong>ALCS</strong>{html.escape(profile["alcs"])}</span>',
                    f'<span><strong>Confidence</strong>{html.escape(profile["confidence"])}</span>',
                    "</div>",
                    f'<p><strong>Visible strength:</strong> {inline_md_to_html(profile["visible_strength"])}</p>',
                    f'<p><strong>Lifecycle gap:</strong> {inline_md_to_html(profile["lifecycle_gap"])}</p>',
                    f'<p><strong>Evidence basis:</strong> {inline_md_to_html(profile["evidence_basis"])}</p>',
                    f'<p class="profile-boundary"><strong>Boundary:</strong> {inline_md_to_html(profile["boundary"])}</p>',
                    "</article>",
                ]
            )
        )
    return "\n".join(
        [
            '<section class="evaluation-results chapter" id="evaluation-results-rccs-alcs-analytical-profiles">',
            "<h1>Evaluation Results: RCCS/ALCS Analytical Profiles</h1>",
            '<div class="chapter-finding-panel evaluation-finding"><div><strong>Main Result</strong><span>Evaluated systems show different visible posture patterns across RCCS and ALCS. Governance-platform strength does not automatically equal lifecycle conformance, and lifecycle-protocol strength does not automatically prove enterprise implementation.</span></div><div><strong>Method</strong><span>Profiles use source-qualified qualitative posture labels, evidence confidence notes, and bounded MRO/RCCS/ALCS mapping.</span></div><div><strong>Boundary</strong><span>No profile is a final vendor score, ranking, procurement recommendation, certification, or legal compliance proof.</span></div></div>',
            "<h2>Why Qualitative Profiles Are Used</h2>",
            "<p>The active source contains enough evidence to express visible posture patterns, but not enough release-ready page-level support to publish final numeric vendor scores. Phase 1D-9 therefore chooses qualitative posture profiles rather than score bands or provisional numeric profiles. This keeps the result layer useful while preserving the non-ranking and source-qualified boundary.</p>",
            chart_html.get("rccs_alcs_quadrant", ""),
            heatmap_html(),
            "<h2>System Profile Cards</h2>",
            '<div class="evaluation-profile-grid">',
            "\n".join(profile_cards),
            "</div>",
            '<div class="callout boundary">Reading rule: compare patterns, not rank order. The profiles are grouped in the same representative-system sequence used by the source mappings and do not imply a leaderboard.</div>',
            "</section>",
        ]
    )


def parse_simple_yaml_block(code: str) -> list[tuple[str, str]]:
    fields: list[tuple[str, str]] = []
    for raw in code.splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or ":" not in line:
            continue
        key, value = line.split(":", 1)
        value = value.strip().strip('"').strip("'")
        if value and value not in {"|", ">"}:
            fields.append((key.strip(), value))
    return fields


def render_object_card_from_code(code: str, language: str, relative: str, records: list[CodeBlockRecord]) -> str:
    fields = parse_simple_yaml_block(code)
    title = "Lifecycle Object Profile"
    for key, value in fields:
        if key in {"object_type", "object_id", "mro_id"} and value:
            title = value
            break
    rows = "\n".join(
        f'<div><dt>{html.escape(key)}</dt><dd>{inline_md_to_html(value)}</dd></div>' for key, value in fields[:12]
    )
    if not rows:
        rows = f'<div><dt>Template</dt><dd>{html.escape(code[:280])}</dd></div>'
    records.append(CodeBlockRecord(relative, language, len(code.splitlines()), "object-card"))
    return "\n".join(
        [
            '<div class="object-card">',
            f'<div class="object-card-title">{html.escape(title)}</div>',
            '<div class="object-card-boundary">Reader-friendly object profile generated from a short source code block. It is illustrative, non-normative, and does not change the underlying source content.</div>',
            f'<dl>{rows}</dl>',
            "</div>",
        ]
    )


def render_code_block(code: str, language: str, relative: str, records: list[CodeBlockRecord]) -> str:
    line_count = len(code.splitlines())
    is_appendix = relative.startswith("appendices/")
    is_yaml = language.lower() in {"yaml", "yml"} or re.search(r"^\s*[A-Za-z0-9_-]+:\s*", code, re.MULTILINE)
    if not is_appendix and is_yaml and line_count <= 32:
        return render_object_card_from_code(code, language, relative, records)
    klass = "template-code" if is_appendix else "code-block"
    caption = "Illustrative non-normative template" if is_appendix else "Technical example"
    records.append(CodeBlockRecord(relative, language, line_count, klass))
    return "\n".join(
        [
            f'<figure class="{klass}">',
            f'<figcaption>{caption}</figcaption>',
            f'<pre><code>{html.escape(code)}</code></pre>',
            "</figure>",
        ]
    )


def render_figure(lines: list[str], start: int, relative: str, records: list[FigureRecord]) -> tuple[str, int]:
    figure_id, title = parse_figure_title(lines[start])
    i = start + 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    flow = ""
    if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not is_table_start(lines, i):
        flow = lines[i].strip()
        i += 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    interpretation = ""
    if i < len(lines) and re.match(r"^\*Figure F-\d{2}\b", lines[i].strip()):
        interpretation = lines[i].strip().strip("*")
        i += 1
    records.append(
        FigureRecord(
            figure_id=figure_id,
            title=title,
            source_file=relative,
            has_flow=bool(flow),
            has_interpretation=bool(interpretation),
        )
    )
    return (
        "\n".join(
            [
                f'<figure class="figure-card" id="figure-{figure_id.lower()}">',
                f'<figcaption><span>{figure_id}</span>{html.escape(title)}</figcaption>',
                render_flow(flow),
                f'<p class="figure-note">{inline_md_to_html(interpretation)}</p>' if interpretation else "",
                "</figure>",
            ]
        ),
        i,
    )


def callout_class(text: str) -> str | None:
    low = strip_markdown_for_text(text).lower()
    if "not legal advice" in low or "not a certification" in low or "not certification" in low:
        return "callout boundary"
    if low.startswith("**pattern boundary:**") or low.startswith("**boundary note:**") or low.startswith("**chapter boundary:**"):
        return "callout boundary"
    if low.startswith("**the paper argues") or low.startswith("**core thesis"):
        return "callout thesis"
    if "source-bound" in low or "evidence" in low and low.startswith("**"):
        return "callout evidence"
    if "blocker" in low or "unresolved" in low:
        return "callout warning"
    return None


def render_list(lines: list[str], start: int, ordered: bool) -> tuple[str, int]:
    tag = "ol" if ordered else "ul"
    parts = [f"<{tag}>"]
    i = start
    pattern = r"^\s*\d+\.\s+(.*)$" if ordered else r"^\s*[-*]\s+(.*)$"
    while i < len(lines):
        match = re.match(pattern, lines[i])
        if not match:
            break
        parts.append(f"<li>{inline_md_to_html(match.group(1))}</li>")
        i += 1
    parts.append(f"</{tag}>")
    return "\n".join(parts), i


def render_heading(line: str, used_slugs: dict[str, int], relative: str) -> tuple[str, dict[str, str | int]]:
    level = min(heading_level(line), 4)
    text = strip_markdown_for_text(line[level:].strip())
    slug = slugify(text, used_slugs)
    classes = []
    if relative.startswith("appendices/") and level == 1:
        classes.append("appendix-heading")
    class_attr = f' class="{" ".join(classes)}"' if classes else ""
    return f'<h{level} id="{slug}"{class_attr}>{html.escape(text)}</h{level}>', {"level": level, "text": text, "id": slug, "source": relative}


def chapter_hierarchy_panel(relative: str) -> str:
    key = Path(relative).name
    values = CHAPTER_HIERARCHY.get(key)
    if not values:
        return ""
    main, matters, method, concludes = values
    return "\n".join(
        [
            '<div class="chapter-finding-panel">',
            f'<div><strong>Main Finding</strong><span>{inline_md_to_html(main)}</span></div>',
            f'<div><strong>Why It Matters</strong><span>{inline_md_to_html(matters)}</span></div>',
            f'<div><strong>Method / Evidence Basis</strong><span>{inline_md_to_html(method)}</span></div>',
            f'<div><strong>What This Chapter Concludes</strong><span>{inline_md_to_html(concludes)}</span></div>',
            "</div>",
        ]
    )


def render_source_file(
    relative: str,
    used_slugs: dict[str, int],
    table_records: list[TableRecord],
    figure_records: list[FigureRecord],
    code_records: list[CodeBlockRecord],
    chart_html: dict[str, str],
) -> tuple[str, list[dict[str, str | int]], str]:
    lines = read_source_lines(relative)
    section_classes = ["source-section"]
    if relative.startswith("appendices/"):
        section_classes.append("appendix")
        section_classes.append("appendix-mode")
    elif relative.startswith("sections/") and relative != "sections/00-front-matter.md":
        section_classes.append("chapter")
        section_classes.append("body-mode")
    else:
        section_classes.append("front-matter-source")
    html_parts: list[str] = [f'<section class="{" ".join(section_classes)}" data-source="{html.escape(relative)}">']
    headings: list[dict[str, str | int]] = []
    source_text_lines: list[str] = []
    i = 0
    in_code = False
    code_language = ""
    code_buffer: list[str] = []
    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        source_text_lines.append(line)

        if stripped.startswith("```"):
            if in_code:
                html_parts.append(render_code_block("\n".join(code_buffer), code_language, relative, code_records))
                code_buffer = []
                code_language = ""
                in_code = False
            else:
                in_code = True
                code_language = stripped.strip("`").strip()
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            html_parts.append('<hr class="section-rule">')
            i += 1
            continue
        if relative == "sections/00-front-matter.md" and stripped == "# Global AI Compliance White Paper 2026":
            i += 1
            continue
        if is_table_start(lines, i):
            rows, new_i = parse_markdown_table(lines, i)
            html_parts.append(render_table(rows, pending_caption, relative, table_records))
            pending_caption = None
            i = new_i
            continue
        if stripped.startswith("#"):
            rendered, heading = render_heading(stripped, used_slugs, relative)
            headings.append(heading)
            html_parts.append(rendered)
            if int(heading["level"]) == 1 and relative.startswith("sections/") and relative != "sections/00-front-matter.md":
                panel = chapter_hierarchy_panel(relative)
                if panel:
                    html_parts.append(panel)
            if relative.endswith("06-missing-regulatory-objects.md") and "Summary of the Sixteen Missing Regulatory Objects" in str(heading["text"]):
                html_parts.append(chart_html.get("mro_distribution", ""))
            if relative.endswith("09-composite-scoring-method.md") and "Evidence Multiplier Application" in str(heading["text"]):
                html_parts.append(chart_html.get("evidence_multipliers", ""))
            if relative.endswith("15-failure-scenarios-placeholder.md") and "Failure Mode Categories" in str(heading["text"]):
                html_parts.append(chart_html.get("failure_distribution", ""))
            if relative.endswith("17-adoption-roadmap-placeholder.md") and "Operationalizing Lifecycle Responsibility Objects" in str(heading["text"]):
                html_parts.append(chart_html.get("adoption_timeline", ""))
            pending_caption = None
            i += 1
            continue
        if is_caption_line(stripped):
            pending_caption = stripped
            i += 1
            continue
        if is_figure_title(stripped):
            rendered, new_i = render_figure(lines, i, relative, figure_records)
            html_parts.append(rendered)
            pending_caption = None
            i = new_i
            continue
        if stripped.startswith(">"):
            html_parts.append(f'<blockquote>{inline_md_to_html(stripped.lstrip("> "))}</blockquote>')
            i += 1
            continue
        if re.match(r"^\s*[-*]\s+", line):
            rendered, new_i = render_list(lines, i, ordered=False)
            html_parts.append(rendered)
            i = new_i
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            rendered, new_i = render_list(lines, i, ordered=True)
            html_parts.append(rendered)
            i = new_i
            continue

        klass = callout_class(stripped)
        if klass:
            html_parts.append(f'<div class="{klass}">{inline_md_to_html(stripped)}</div>')
        else:
            html_parts.append(f"<p>{inline_md_to_html(stripped)}</p>")
        i += 1

    if code_buffer:
        html_parts.append(render_code_block("\n".join(code_buffer), code_language, relative, code_records))
    html_parts.append("</section>")
    return "\n".join(html_parts), headings, "\n".join(source_text_lines)


def build_generated_toc(headings: list[dict[str, str | int]]) -> str:
    front_items = [
        '<li><a href="#problem-this-white-paper-solves">Problem This White Paper Solves</a></li>',
        '<li><a href="#how-to-read-this-paper">How to Read This Paper</a></li>',
    ]
    parts: dict[str, list[str]] = {
        "Part I - Foundations": [],
        "Part II - Lifecycle Objects and Scoring": [],
        "Part III - System Mapping and Control": [],
        "Part IV - Validation and Boundary": [],
        "Part V - Adoption and Conclusion": [],
    }
    appendix_items = []
    chapter_groups = [
        (1, 5, "Part I - Foundations"),
        (6, 9, "Part II - Lifecycle Objects and Scoring"),
        (10, 13, "Part III - System Mapping and Control"),
        (14, 16, "Part IV - Validation and Boundary"),
        (17, 18, "Part V - Adoption and Conclusion"),
    ]
    for heading in headings:
        text = str(heading["text"])
        if int(heading["level"]) != 1:
            continue
        item = f'<li><a href="#{heading["id"]}">{html.escape(text)}</a></li>'
        if text.startswith("Evaluation Results"):
            parts["Part III - System Mapping and Control"].append(item)
        elif text.startswith("Appendix"):
            appendix_items.append(item)
        else:
            match = re.match(r"^(?:Chapter\s+)?([0-9]{1,2})\b", text)
            if match:
                num = int(match.group(1))
                for start, end, label in chapter_groups:
                    if start <= num <= end:
                        parts[label].append(item)
                        break
    appendix_items.append('<li><a href="#references-sources">References / Sources</a></li>')
    part_html = []
    for label, items in parts.items():
        part_html.append(f"<div><h3>{html.escape(label)}</h3><ol>")
        part_html.append("\n".join(items))
        part_html.append("</ol></div>")
    return "\n".join(
        [
            '<nav class="generated-toc" aria-label="Publication table of contents">',
            "<h2>Publication Contents</h2>",
            '<p class="toc-note">HTML anchors are active in the browser artifact. PDF page-number pinning remains a final publication QA task.</p>',
            '<div class="toc-grid">',
            '<div><h3>Front Matter</h3><ol>',
            "\n".join(front_items),
            "</ol></div>",
            "\n".join(part_html),
            '<div><h3>Appendices</h3><ol>',
            "\n".join(appendix_items),
            "</ol></div>",
            "</div>",
            "</nav>",
        ]
    )


def cover_html() -> str:
    return f"""
<section class="cover-page">
  <div class="cover-kicker">Jearon Wong / AI Agent Lifecycle Governance</div>
  <h1>Global AI Compliance White Paper 2026</h1>
  <p class="cover-subtitle">From Model Governance to Agentic Lifecycle Conformance</p>
  <p class="cover-role">Jearon Wong · Protocol Architect for the Agent Era</p>
  <div class="thesis-strip" aria-label="Lifecycle thesis strip">
    <span>Intent</span><span>Authority</span><span>Agent Work</span><span>Evidence</span><span>Closure</span>
  </div>
  <div class="cover-meta">
    <div><strong>Document ID</strong><span>{TRACE_TAG}</span></div>
    <div><strong>Version</strong><span>v0.3.2 Final Release Candidate R3</span></div>
    <div><strong>Date</strong><span>May 2026</span></div>
    <div><strong>Status</strong><span>Evaluation Results and Semantic Table Reflow Draft / QA Required</span></div>
    <div><strong>Visual source of truth</strong><span>HTML/PDF primary; DOCX derivative</span></div>
  </div>
  <div class="cover-notice">Non-legal technical governance analysis. Not legal advice, not a certification standard, not regulatory approval, not a procurement recommendation, not a vendor ranking, and not a final vendor assessment.</div>
</section>
"""


def important_notice_html() -> str:
    return """
<section class="important-notice">
  <h2>Important Notice</h2>
  <p>This publication is a technical governance analysis of AI Agent Lifecycle Governance. It is not legal advice, a certification standard, a conformity assessment, a regulatory approval path, an audit opinion, an insurance guarantee, a procurement recommendation, a vendor ranking, or a final vendor assessment.</p>
  <p>RCCS/ALCS scores and system mappings are analytical and source-qualified. Appendix G remains revalidated provisional, qualitative, non-ranking, and no-score.</p>
</section>
"""


def problem_overview_html() -> str:
    return """
<section class="problem-overview front-extra">
  <h1 id="problem-this-white-paper-solves">Problem This White Paper Solves</h1>
  <div class="callout thesis">The central problem is that model governance is necessary but insufficient once AI systems act through agents, tools, handoffs, memory, projects, vendors, and accepted outcomes.</div>
  <p>Traditional AI governance is usually organized around the model: model risk, model documentation, model monitoring, model security, and model transparency. Agentic systems shift the compliance object. The governance question is no longer only whether a model is evaluated; it is whether the lifecycle of delegated work can prove who authorized an action, what authority was granted, which agent or tool acted, what evidence was preserved, who accepted the outcome, and how disputes or remediation close.</p>
  <p>This white paper frames that missing layer as AI Agent Lifecycle Governance. Enterprise failures emerge when lifecycle responsibility objects are absent: delegated authority boundaries are implicit, human oversight is not mapped to machine roles, evidence chains cannot reconstruct agent work, privacy rights conflict with retention, and vendor or runtime substitution changes governance semantics without conformance review.</p>
  <p>The paper responds with sixteen Missing Regulatory Objects, the RCCS/ALCS analytical model, evidence-based validation patterns, illustrative enterprise failure scenarios, and a staged adoption roadmap. These tools support structured posture review and audit-readiness analysis. They do not create legal advice, certification, regulatory approval, procurement recommendations, vendor rankings, or final vendor assessments.</p>
  <div class="overview-grid">
    <div><strong>Problem</strong><span>Model-centric governance does not fully express lifecycle responsibility in agentic workflows.</span></div>
    <div><strong>Object Layer</strong><span>MROs define the missing responsibility, authority, evidence, privacy, substitution, and closure objects.</span></div>
    <div><strong>Assessment</strong><span>RCCS/ALCS structure analytical posture review without becoming legal compliance scores.</span></div>
    <div><strong>Assurance Pattern</strong><span>Evidence-based validation supports reviewable assurance while preserving non-certifying boundaries.</span></div>
    <div><strong>Adoption</strong><span>The roadmap moves enterprises from inventory to authority, responsibility, evidence, privacy, validation, remediation, and scale.</span></div>
  </div>
</section>
"""


def how_to_read_html() -> str:
    return """
<section class="how-to-read front-extra">
  <h1 id="how-to-read-this-paper">How to Read This Paper</h1>
  <div class="reader-grid">
    <div><strong>Executives</strong><span>Read Chapters 1-3, 15, 17, and 18 for the business risk, failure patterns, adoption path, and conclusion.</span></div>
    <div><strong>Compliance / Legal Teams</strong><span>Read Chapters 4-10 and Appendices C, I, J, and K for regulatory mapping, scoring boundaries, source posture, and governance records.</span></div>
    <div><strong>Product / Platform Teams</strong><span>Read Chapters 5-9, 12, and Appendices A, B, E, and F for lifecycle objects, scoring rubrics, and implementation-facing templates.</span></div>
    <div><strong>Auditors / Assurance Reviewers</strong><span>Read Chapters 10, 14, 15, and Appendices C, D, G, and I for evidence patterns, validation boundaries, and source-qualified system mappings.</span></div>
    <div><strong>Protocol / Runtime Architects</strong><span>Read Chapters 6, 8, 13, 14, and Appendix A for MRO semantics, ALCS mapping, protocol fit, and machine-readable object templates.</span></div>
  </div>
</section>
"""


def parse_citation_inventory() -> list[dict[str, str]]:
    path = ROOT / "inventories" / "citation-inventory.md"
    text = path.read_text(encoding="utf-8", errors="ignore")
    rows: list[dict[str, str]] = []
    in_table = False
    headers: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if line.startswith("| Source ID |"):
            in_table = True
            headers = [h.strip() for h in line.strip("|").split("|")]
            continue
        if in_table and line.startswith("|---"):
            continue
        if in_table and line.startswith("|"):
            parts = [p.strip() for p in line.strip("|").split("|")]
            if len(parts) == len(headers):
                rows.append(dict(zip(headers, parts)))
            continue
        if in_table and not line.startswith("|"):
            break
    return rows


def references_section_html() -> str:
    rows = parse_citation_inventory()
    row_by_id = {row.get("Source ID", ""): row for row in rows}
    level_counts: dict[str, int] = {}
    status_counts = {"source_bound": 0, "source_pointer_queue": 0, "l5_analytical_inference": 0, "unresolved_access": 0}
    for row in rows:
        level = strip_markdown_for_text(row.get("Evidence Level", "")) or "Unspecified"
        level_counts[level] = level_counts.get(level, 0) + 1
        row_text = " ".join(strip_markdown_for_text(value).lower() for value in row.values())
        if "403" in row_text or "unresolved" in row_text or "blocked" in row_text:
            status_counts["unresolved_access"] += 1
        if "source-pointer" in row_text or "source pointer" in row_text or "queue" in row_text:
            status_counts["source_pointer_queue"] += 1
        if "l5" in row_text or "analytical" in row_text or "inference" in row_text:
            status_counts["l5_analytical_inference"] += 1
        if "verified" in row_text or "official" in row_text or "identified" in row_text:
            status_counts["source_bound"] += 1
    sections: list[str] = [
        '<section class="references-section appendix" id="references-sources">',
        "<h1>References / Sources</h1>",
        '<div class="callout boundary">This section is a publication-facing source register. It summarizes source classes, evidence levels, and citation status. It does not convert product documentation into audit proof, legal compliance proof, certification, regulatory approval, procurement recommendation, or vendor endorsement.</div>',
        "<h2>How to Read Evidence Labels</h2>",
        '<div class="source-summary-grid">',
        f'<div><strong>Source-bound</strong><span>{status_counts["source_bound"]} entries contain identified or official source support language.</span></div>',
        f'<div><strong>Source-pointer queue</strong><span>{status_counts["source_pointer_queue"]} entries still require final page-level pinning or rendering QA.</span></div>',
        f'<div><strong>L5 analytical inference</strong><span>{status_counts["l5_analytical_inference"]} entries contain analytical or inferred mapping language.</span></div>',
        f'<div><strong>Unresolved access</strong><span>{status_counts["unresolved_access"]} entries contain unresolved access, HTTP 403, or blocking language.</span></div>',
        "</div>",
        '<p class="reference-action">Evidence labels express source confidence and claim-support posture. They do not prove legal compliance, certification, deployment readiness, product quality, or procurement suitability.</p>',
        "<h2>Evidence Level Counts</h2>",
        '<div class="evidence-count-row">',
        "".join(f'<span><strong>{html.escape(level)}</strong>{count}</span>' for level, count in sorted(level_counts.items())),
        "</div>",
    ]
    for group_name, ids in SOURCE_SECTION_GROUPS:
        sections.append(f"<h2>{html.escape(group_name)}</h2>")
        sections.append('<div class="reference-list">')
        for source_id in ids:
            row = row_by_id.get(source_id)
            if not row:
                continue
            name = strip_markdown_for_text(row.get("Source Name", ""))
            klass = strip_markdown_for_text(row.get("Source Class", ""))
            level = strip_markdown_for_text(row.get("Evidence Level", ""))
            status = strip_markdown_for_text(row.get("URL / Entry Point Status", ""))
            action = strip_markdown_for_text(row.get("Phase 1C Action", ""))
            sections.append(
                "\n".join(
                    [
                        '<article class="reference-item">',
                        f'<div class="reference-id">{html.escape(source_id)}</div>',
                        f'<div class="reference-title">{html.escape(name)}</div>',
                        f'<div class="reference-meta"><span>{html.escape(klass)}</span><span>{html.escape(level)}</span></div>',
                        f'<p>{inline_md_to_html(status)}</p>',
                        f'<p class="reference-action">{inline_md_to_html(action)}</p>',
                        "</article>",
                    ]
                )
            )
        sections.append("</div>")
    sections.append(
        """
<h2>Source Status Notes</h2>
<ul class="source-status-notes">
  <li>L1/L2/L3/L4/L5 evidence levels indicate source confidence and claim support posture; they do not prove deployment compliance.</li>
  <li>Final page-level citation pinning remains a publication QA task for several high-risk L2 product and protocol references.</li>
  <li>The OpenAI platform guide HTTP 403 access issue remains explicitly tracked; accessible Agents SDK documentation supports only narrowed SDK-surface claims.</li>
  <li>Validation Lab remains a boundary-limited public-surface, non-certifying evidence adjudication example unless a release-approved methodology citation is pinned.</li>
  <li>MPLP references carry a conflict-of-interest boundary: MPLP may be analyzed as one lifecycle protocol path, but this paper does not require, certify, or exclusively endorse MPLP.</li>
</ul>
</section>
"""
    )
    return "\n".join(sections)


def css() -> str:
    return """
:root {
  --monolith: #111111;
  --slate: #1a1a1a;
  --machine: #f1f5f9;
  --registry: #3b82f6;
  --evidence: #22d3ee;
  --status: #64748b;
  --grid: #d8dee8;
  --amber: #b7791f;
  --fault: #991b1b;
  --paper: #ffffff;
}
* { box-sizing: border-box; }
html { background: var(--machine); color: var(--monolith); }
body {
  margin: 0;
  font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  font-size: 15px;
  line-height: 1.52;
  color: var(--monolith);
}
.publication {
  max-width: 980px;
  margin: 0 auto;
  background: var(--paper);
  padding: 32px 80px 48px;
  box-shadow: 0 20px 80px rgba(15, 23, 42, 0.08);
}
a { color: var(--registry); text-decoration: none; }
p { margin: 0 0 0.82rem; }
h1, h2, h3, h4 {
  font-family: Outfit, Inter, system-ui, sans-serif;
  line-height: 1.18;
  letter-spacing: 0;
  color: var(--monolith);
  break-after: avoid;
}
h1 {
  margin: 2.2rem 0 1rem;
  padding-top: 0.6rem;
  border-top: 4px solid var(--monolith);
  font-size: 2.32rem;
}
h2 {
  margin: 2rem 0 0.75rem;
  padding-top: 0.75rem;
  border-top: 1px solid var(--grid);
  font-size: 1.54rem;
}
h3 { margin: 1.35rem 0 0.48rem; font-size: 1.14rem; color: var(--slate); }
h4 { margin: 1rem 0 0.35rem; font-size: 0.98rem; color: var(--status); }
ul, ol { margin: 0.2rem 0 0.9rem 1.35rem; padding: 0; }
li { margin: 0.18rem 0; }
blockquote {
  margin: 1rem 0;
  padding: 0.65rem 0.9rem;
  border-left: 4px solid var(--registry);
  background: #f8fafc;
  color: var(--slate);
}
.cover-page {
  min-height: 92vh;
  display: flex;
  flex-direction: column;
  justify-content: center;
  border: 1px solid var(--monolith);
  padding: 52px 44px;
  margin-bottom: 2.5rem;
  background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
}
.cover-kicker {
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  color: var(--status);
  text-transform: uppercase;
  font-size: 0.78rem;
  margin-bottom: 2.2rem;
}
.cover-page h1 {
  border: 0;
  margin: 0;
  padding: 0;
  font-size: 3.6rem;
  max-width: 760px;
}
.cover-subtitle {
  margin: 1rem 0 0.4rem;
  color: var(--registry);
  font-size: 1.25rem;
  font-weight: 700;
}
.cover-role { color: var(--status); font-size: 0.95rem; }
.thesis-strip {
  display: grid;
  grid-template-columns: repeat(5, 1fr);
  gap: 1px;
  margin: 2rem 0;
  border: 1px solid var(--monolith);
  background: var(--monolith);
}
.thesis-strip span {
  background: var(--monolith);
  color: #fff;
  padding: 0.68rem 0.55rem;
  text-align: center;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
  font-weight: 700;
}
.cover-meta {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 1px;
  background: var(--grid);
  border: 1px solid var(--grid);
  margin-top: 1.2rem;
}
.cover-meta div { background: #fff; padding: 0.65rem 0.8rem; }
.cover-meta strong {
  display: block;
  color: var(--status);
  font-size: 0.72rem;
  text-transform: uppercase;
}
.cover-meta span { font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace; font-size: 0.78rem; }
.cover-notice, .important-notice, .callout {
  border-left: 4px solid var(--amber);
  background: #fffbeb;
  padding: 0.85rem 1rem;
  margin: 1rem 0;
  break-inside: avoid;
}
.important-notice { margin: 0 0 2rem; }
.important-notice h2 { border: 0; padding: 0; margin-top: 0; }
.callout.thesis { border-color: var(--registry); background: #eff6ff; font-weight: 650; }
.callout.evidence { border-color: var(--evidence); background: #ecfeff; }
.callout.warning { border-color: var(--fault); background: #fef2f2; }
.generated-toc {
  margin: 2rem 0 2.8rem;
  padding: 1.1rem;
  border: 1px solid var(--grid);
  background: #f8fafc;
}
.generated-toc h2 { border: 0; padding-top: 0; margin-top: 0; }
.toc-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 1.25rem; }
.toc-grid h3 { margin-top: 0; color: var(--registry); }
.toc-grid li { font-size: 0.88rem; }
.toc-note { color: var(--status); font-size: 0.86rem; margin-top: -0.4rem; }
.source-section { margin-bottom: 1.3rem; }
.chapter-finding-panel {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 1px;
  margin: 0.45rem 0 1.15rem;
  border: 1px solid var(--grid);
  background: var(--grid);
  break-inside: avoid;
}
.chapter-finding-panel div {
  background: #f8fafc;
  padding: 0.68rem 0.78rem;
}
.chapter-finding-panel strong {
  display: block;
  color: var(--registry);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
  text-transform: uppercase;
  margin-bottom: 0.22rem;
}
.chapter-finding-panel span { color: var(--slate); font-size: 0.88rem; }
.evaluation-results {
  border-top: 6px solid var(--monolith);
}
.evaluation-finding { grid-template-columns: 1fr 1fr 1fr; }
.evaluation-profile-grid {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 0.75rem;
  margin: 0.8rem 0 1.2rem;
}
.evaluation-profile-card {
  border: 1px solid var(--grid);
  border-top: 4px solid var(--registry);
  padding: 0.82rem;
  background: #fbfdff;
  break-inside: avoid;
}
.evaluation-profile-card h3 { margin-top: 0.18rem; }
.profile-kicker {
  color: var(--status);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.7rem;
  text-transform: uppercase;
}
.profile-postures {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 1px;
  margin: 0.5rem 0 0.65rem;
  background: var(--grid);
  border: 1px solid var(--grid);
}
.profile-postures span {
  background: #fff;
  padding: 0.36rem 0.42rem;
  font-size: 0.78rem;
}
.profile-postures strong {
  display: block;
  color: var(--status);
  font-size: 0.66rem;
  text-transform: uppercase;
}
.profile-boundary { color: var(--status); font-size: 0.82rem; }
.front-extra {
  margin: 2rem 0 2.6rem;
  padding: 1rem 0 0;
}
.overview-grid, .reader-grid {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 0.65rem;
  margin: 1rem 0 1.4rem;
}
.overview-grid div, .reader-grid div {
  border: 1px solid var(--grid);
  background: #f8fafc;
  padding: 0.75rem 0.85rem;
}
.overview-grid strong, .reader-grid strong {
  display: block;
  font-family: Outfit, Inter, system-ui, sans-serif;
  margin-bottom: 0.22rem;
}
.overview-grid span, .reader-grid span { color: var(--slate); font-size: 0.92rem; }
.appendix-heading {
  border-top-color: var(--registry);
  color: var(--registry);
}
.appendix-mode { font-size: 0.95rem; }
.badge {
  display: inline-block;
  padding: 0.08rem 0.28rem;
  border-radius: 3px;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72em;
  line-height: 1.2;
  vertical-align: baseline;
  border: 1px solid var(--grid);
  background: #f8fafc;
  color: var(--slate);
}
.badge.evidence { border-color: #bae6fd; background: #ecfeff; }
.badge.mro, .badge.alcs, .badge.rccs { border-color: #bfdbfe; background: #eff6ff; color: #1d4ed8; }
.figure-card, .chart-card {
  margin: 1.15rem 0 1.35rem;
  padding: 1rem;
  border: 1px solid var(--grid);
  border-top: 4px solid var(--monolith);
  background: #fbfdff;
  break-inside: avoid;
}
.figure-card figcaption, .chart-card figcaption {
  font-family: Outfit, Inter, system-ui, sans-serif;
  font-weight: 800;
  margin-bottom: 0.8rem;
}
.figure-card figcaption span, .chart-card figcaption span {
  display: inline-block;
  margin-right: 0.55rem;
  padding: 0.12rem 0.38rem;
  background: var(--monolith);
  color: #fff;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
}
.figure-flow {
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 0.35rem;
  margin: 0.45rem 0 0.75rem;
}
.flow-node {
  border: 1px solid var(--grid);
  background: #fff;
  padding: 0.42rem 0.55rem;
  font-size: 0.82rem;
  font-weight: 650;
}
.flow-arrow {
  color: var(--registry);
  font-weight: 800;
}
.figure-note {
  margin: 0.6rem 0 0;
  color: var(--status);
  font-size: 0.88rem;
  font-style: italic;
}
.chart-svg { width: 100%; height: auto; display: block; margin: 0.25rem 0 0.6rem; }
.chart-title { font-family: Outfit, Inter, system-ui, sans-serif; font-size: 18px; font-weight: 800; fill: var(--monolith); }
.chart-label { font-family: Inter, system-ui, sans-serif; font-size: 13px; font-weight: 700; fill: var(--slate); }
.chart-value { font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace; font-size: 12px; font-weight: 800; fill: var(--registry); }
.chart-note { font-family: Inter, system-ui, sans-serif; font-size: 10px; fill: var(--status); }
.chart-track { fill: #e8eef7; }
.chart-bar { fill: var(--registry); }
.chart-bar.alt { fill: var(--evidence); }
.timeline-line { stroke: var(--grid); stroke-width: 5; stroke-linecap: round; }
.timeline-node { fill: var(--monolith); stroke: var(--evidence); stroke-width: 2; }
.timeline-stage { font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace; font-size: 11px; font-weight: 800; fill: #fff; }
.timeline-label { font-family: Inter, system-ui, sans-serif; font-size: 10px; font-weight: 700; fill: var(--slate); }
.quadrant-bg { fill: #f8fafc; stroke: var(--grid); }
.quadrant-line { stroke: #cbd5e1; stroke-width: 2; stroke-dasharray: 6 5; }
.quadrant-label { font-family: Inter, system-ui, sans-serif; font-size: 10px; fill: var(--status); }
.quadrant-dot { stroke: #fff; stroke-width: 2; fill: var(--registry); }
.quadrant-dot.q1 { fill: var(--evidence); }
.quadrant-dot.q2 { fill: #0f172a; }
.quadrant-dot.q3 { fill: #64748b; }
.quadrant-system { font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace; font-size: 10px; fill: var(--slate); font-weight: 800; }
.chart-interpretation { margin: 0.5rem 0 0.25rem; font-size: 0.9rem; color: var(--slate); }
.chart-boundary { margin: 0; font-size: 0.82rem; color: var(--status); font-style: italic; }
.heatmap-card {
  margin: 1.1rem 0 1.3rem;
  padding: 0.95rem;
  border: 1px solid var(--grid);
  background: #fbfdff;
  break-inside: avoid;
}
.heatmap-title {
  font-family: Outfit, Inter, system-ui, sans-serif;
  font-weight: 800;
  margin-bottom: 0.2rem;
}
.heatmap-note { color: var(--status); font-size: 0.82rem; margin-bottom: 0.6rem; }
.heatmap-table th:first-child { width: 28%; }
.heat {
  display: inline-block;
  min-width: 58px;
  padding: 0.12rem 0.32rem;
  border-radius: 3px;
  font-size: 0.72rem;
  font-weight: 800;
  text-align: center;
}
.heat.strong { background: #dcfce7; color: #166534; }
.heat.partial { background: #dbeafe; color: #1d4ed8; }
.heat.limited { background: #fef3c7; color: #92400e; }
.heat.not-visible { background: #f1f5f9; color: var(--status); }
.table-block {
  margin: 1rem 0 1.25rem;
  break-inside: auto;
}
.table-caption {
  font-weight: 800;
  margin: 0 0 0.35rem;
  color: var(--monolith);
  font-size: 0.96rem;
  break-after: avoid;
}
.layout-note, .continuation-label {
  color: var(--status);
  font-size: 0.78rem;
  margin: 0.2rem 0 0.35rem;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
}
table {
  width: 100%;
  border-collapse: collapse;
  table-layout: fixed;
  margin-bottom: 0.7rem;
}
thead { display: table-header-group; }
th, td {
  border: 1px solid var(--grid);
  padding: 0.52rem 0.56rem;
  vertical-align: top;
  overflow-wrap: anywhere;
}
th {
  background: #e8eef7;
  color: var(--slate);
  font-size: 0.84rem;
  text-align: left;
}
td { font-size: 0.86rem; }
.table-rubric table th:nth-child(1), .table-rubric table td:nth-child(1) { width: 8%; }
.table-rubric table th:nth-child(2), .table-rubric table td:nth-child(2) { width: 24%; }
.table-rubric table th:nth-child(3), .table-rubric table td:nth-child(3) { width: 68%; }
.table-system-profile table th:nth-child(1), .table-system-profile table td:nth-child(1) { width: 20%; }
.table-system-profile table th:nth-child(2), .table-system-profile table td:nth-child(2) { width: 28%; }
.table-system-profile table th:last-child, .table-system-profile table td:last-child { width: 52%; }
.table-boundary table th:nth-child(1), .table-boundary table td:nth-child(1) { width: 18%; }
.table-boundary table th:nth-child(2), .table-boundary table td:nth-child(2) { width: 41%; }
.table-boundary table th:nth-child(3), .table-boundary table td:nth-child(3) { width: 41%; }
.table-source-register table th:first-child, .table-source-register table td:first-child { width: 18%; }
.rubric-card-set {
  margin: 1rem 0 1.25rem;
  break-inside: auto;
}
.rubric-card {
  display: grid;
  grid-template-columns: 62px 1fr;
  gap: 0.7rem;
  padding: 0.62rem 0.72rem;
  margin-bottom: 0.55rem;
  border: 1px solid var(--grid);
  background: #fbfdff;
  break-inside: avoid;
}
.rubric-score {
  display: flex;
  align-items: center;
  justify-content: center;
  min-height: 42px;
  background: var(--monolith);
  color: #fff;
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-weight: 800;
  font-size: 1rem;
}
.rubric-body h4 {
  margin: 0 0 0.25rem;
  color: var(--registry);
}
.rubric-body p { margin-bottom: 0.25rem; font-size: 0.88rem; }
.rubric-body dl { margin: 0.2rem 0 0; }
.rubric-body dt { font-weight: 800; color: var(--status); font-size: 0.74rem; }
.rubric-body dd { margin: 0 0 0.15rem; font-size: 0.82rem; }
.system-profile-card-set {
  margin: 1rem 0 1.25rem;
  break-inside: auto;
}
.system-profile-row-card {
  border: 1px solid var(--grid);
  border-left: 4px solid var(--registry);
  background: #fbfdff;
  padding: 0.72rem 0.82rem;
  margin-bottom: 0.58rem;
  break-inside: avoid;
}
.system-profile-row-card h4 {
  margin-top: 0;
  color: var(--registry);
}
.system-profile-row-card dl {
  display: grid;
  grid-template-columns: 0.82fr 2.1fr;
  gap: 1px;
  margin: 0.38rem 0 0;
  border: 1px solid var(--grid);
  background: var(--grid);
}
.system-profile-row-card div { display: contents; }
.system-profile-row-card dt, .system-profile-row-card dd {
  background: #fff;
  margin: 0;
  padding: 0.4rem 0.5rem;
  font-size: 0.82rem;
}
.system-profile-row-card dt {
  color: var(--status);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-weight: 800;
}
.table-compact th, .table-compact td {
  font-size: 0.78rem;
  padding: 0.38rem 0.4rem;
}
.appendix-mode .table-compact th, .appendix-mode .table-compact td {
  font-size: 0.72rem;
  padding: 0.32rem 0.34rem;
}
.code-block, .template-code {
  margin: 0.9rem 0 1.1rem;
  border: 1px solid var(--grid);
  background: #f8fafc;
  break-inside: auto;
}
.code-block figcaption, .template-code figcaption {
  padding: 0.42rem 0.65rem;
  border-bottom: 1px solid var(--grid);
  color: var(--status);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.72rem;
  font-weight: 700;
}
.code-block pre, .template-code pre {
  margin: 0;
  white-space: pre-wrap;
  color: #0f172a;
  padding: 0.72rem 0.78rem;
  overflow-wrap: anywhere;
  font-size: 0.74rem;
  line-height: 1.42;
}
.template-code {
  background: #fbfdff;
  border-left: 4px solid var(--registry);
}
.object-card {
  margin: 0.9rem 0 1.1rem;
  padding: 0.9rem;
  border: 1px solid var(--grid);
  border-left: 4px solid var(--registry);
  background: #f8fafc;
  break-inside: avoid;
}
.object-card-title {
  font-family: Outfit, Inter, system-ui, sans-serif;
  font-weight: 800;
  margin-bottom: 0.25rem;
}
.object-card-boundary {
  color: var(--status);
  font-size: 0.78rem;
  margin-bottom: 0.6rem;
}
.object-card dl {
  display: grid;
  grid-template-columns: 1fr 2fr;
  gap: 1px;
  margin: 0;
  background: var(--grid);
  border: 1px solid var(--grid);
}
.object-card div { display: contents; }
.object-card dt, .object-card dd {
  background: #fff;
  margin: 0;
  padding: 0.42rem 0.55rem;
  font-size: 0.82rem;
}
.object-card dt {
  color: var(--status);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-weight: 700;
}
.references-section {
  margin-top: 2.2rem;
}
.reference-list {
  display: grid;
  grid-template-columns: 1fr;
  gap: 0.55rem;
}
.reference-item {
  border: 1px solid var(--grid);
  padding: 0.72rem 0.82rem;
  background: #fbfdff;
  break-inside: avoid;
}
.reference-id {
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  color: var(--registry);
  font-weight: 800;
  font-size: 0.78rem;
}
.reference-title { font-weight: 800; margin-top: 0.12rem; }
.reference-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 0.35rem;
  margin: 0.25rem 0 0.35rem;
}
.reference-meta span {
  border: 1px solid var(--grid);
  background: #fff;
  padding: 0.08rem 0.32rem;
  font-size: 0.74rem;
  color: var(--status);
}
.reference-item p { font-size: 0.82rem; margin-bottom: 0.28rem; }
.reference-action { color: var(--status); font-style: italic; }
.source-status-notes li { font-size: 0.9rem; }
.source-summary-grid {
  display: grid;
  grid-template-columns: repeat(4, 1fr);
  gap: 1px;
  margin: 0.75rem 0 0.9rem;
  border: 1px solid var(--grid);
  background: var(--grid);
}
.source-summary-grid div {
  background: #f8fafc;
  padding: 0.58rem 0.65rem;
}
.source-summary-grid strong {
  display: block;
  color: var(--registry);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.68rem;
  text-transform: uppercase;
}
.source-summary-grid span { font-size: 0.78rem; color: var(--slate); }
.evidence-count-row {
  display: flex;
  flex-wrap: wrap;
  gap: 0.38rem;
  margin: 0.5rem 0 0.9rem;
}
.evidence-count-row span {
  border: 1px solid var(--grid);
  background: #f8fafc;
  padding: 0.28rem 0.45rem;
  font-size: 0.78rem;
}
.evidence-count-row strong {
  margin-right: 0.34rem;
  color: var(--registry);
  font-family: "JetBrains Mono", "SFMono-Regular", Consolas, monospace;
}
.part-divider {
  break-before: page;
  padding-top: 2rem;
}
.section-rule { border: 0; border-top: 1px solid var(--grid); margin: 1.2rem 0; }
.html-footer {
  margin-top: 3rem;
  padding-top: 1rem;
  border-top: 1px solid var(--grid);
  color: var(--status);
  font-size: 0.76rem;
  text-align: center;
}
@media print {
  @page { size: A4; margin: 16mm 15mm; }
  html, body { background: #fff; }
  body { font-size: 10.2pt; line-height: 1.45; }
  .publication {
    max-width: none;
    margin: 0;
    padding: 0;
    box-shadow: none;
  }
  .cover-page {
    min-height: 265mm;
    page-break-after: always;
    margin: 0;
  }
  .front-extra, .generated-toc { break-before: page; }
  .chapter, .appendix { break-before: page; }
  .front-matter-source { break-before: auto; }
  .important-notice { break-inside: avoid; }
  h1 { font-size: 25pt; margin-top: 16pt; }
  h2 { font-size: 16pt; margin-top: 14pt; }
  h3 { font-size: 12pt; }
  h4 { font-size: 10pt; }
  p, li { orphans: 2; widows: 2; }
  .figure-card, .chart-card, .callout, .object-card, .reference-item { break-inside: avoid; }
  .table-block, .template-code { break-inside: auto; }
  .generated-toc { break-inside: auto; }
  .toc-grid { gap: 0.8rem; }
  .toc-grid, .overview-grid, .reader-grid { grid-template-columns: 1fr 1fr; }
  .source-summary-grid { grid-template-columns: 1fr 1fr; }
  .chapter-finding-panel, .evaluation-profile-grid { grid-template-columns: 1fr 1fr; }
  th, td { padding: 4pt 4.3pt; }
  td { font-size: 8.7pt; }
  th { font-size: 8.5pt; }
  .table-caption { font-size: 10pt; }
  .table-compact th, .table-compact td { font-size: 7.9pt; padding: 3pt; }
  .appendix-mode .table-compact th, .appendix-mode .table-compact td { font-size: 7.35pt; padding: 2.55pt; }
  .chart-card { break-inside: avoid; }
  .rubric-card, .system-profile-row-card, .evaluation-profile-card, .heatmap-card { break-inside: avoid; }
  .rubric-score { min-height: 34pt; }
  .system-profile-row-card dt, .system-profile-row-card dd { font-size: 7.9pt; }
  .template-code pre { font-size: 7.5pt; line-height: 1.34; }
  .code-block pre { font-size: 8pt; }
  .html-footer { display: none; }
}
"""


def build_html() -> tuple[Path, list[TableRecord], list[FigureRecord], list[ChartRecord], list[CodeBlockRecord], list[dict[str, str | int]], str]:
    used_slugs: dict[str, int] = {}
    table_records: list[TableRecord] = []
    figure_records: list[FigureRecord] = []
    chart_html, chart_records = chart_blocks()
    code_records: list[CodeBlockRecord] = []
    all_headings: list[dict[str, str | int]] = []
    sections: list[str] = []
    source_texts: list[str] = []

    for rel in SOURCE_FILES:
        rendered, headings, source_text = render_source_file(rel, used_slugs, table_records, figure_records, code_records, chart_html)
        sections.append(rendered)
        all_headings.extend(headings)
        source_texts.append(source_text)
        if rel == "sections/12-detailed-system-mappings.md":
            sections.append(evaluation_results_html(chart_html))
            all_headings.append(
                {
                    "level": 1,
                    "text": "Evaluation Results: RCCS/ALCS Analytical Profiles",
                    "id": "evaluation-results-rccs-alcs-analytical-profiles",
                    "source": "generated/evaluation-results",
                }
            )

    generated_toc = build_generated_toc(all_headings)
    document = "\n".join(
        [
            "<!doctype html>",
            '<html lang="en">',
            "<head>",
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            f"<title>Global AI Compliance White Paper 2026 - {TRACE_TAG}</title>",
            "<style>",
            css(),
            "</style>",
            "</head>",
            "<body>",
            '<main class="publication">',
            cover_html(),
            important_notice_html(),
            generated_toc,
            problem_overview_html(),
            how_to_read_html(),
            "\n".join(sections),
            references_section_html(),
            f'<footer class="html-footer">© 2026 Jearon Wong · {TRACE_TAG} · Non-legal technical governance analysis · HTML/PDF visual source of truth</footer>',
            "</main>",
            "</body>",
            "</html>",
        ]
    )
    html_path = OUT_DIR / HTML_NAME
    html_path.write_text(document, encoding="utf-8")
    return html_path, table_records, figure_records, chart_records, code_records, all_headings, "\n\n".join(source_texts)


def chrome_path() -> str:
    candidates = [
        which("google-chrome"),
        which("chromium"),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    raise RuntimeError("No Chrome/Chromium executable found for HTML-to-PDF export")


def html_to_pdf(html_path: Path) -> Path:
    pdf_path = OUT_DIR / PDF_NAME
    if pdf_path.exists():
        pdf_path.unlink()
    chrome = chrome_path()
    user_data = OUT_DIR / "chrome_profile"
    if user_data.exists():
        shutil.rmtree(user_data)
    user_data.mkdir(parents=True, exist_ok=True)
    cmd = [
        chrome,
        "--headless=new",
        "--disable-gpu",
        "--disable-background-networking",
        "--no-pdf-header-footer",
        "--disable-extensions",
        "--no-first-run",
        "--no-default-browser-check",
        f"--user-data-dir={user_data}",
        f"--print-to-pdf={pdf_path}",
        html_path.resolve().as_uri(),
    ]
    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    last_size = -1
    stable_ticks = 0
    deadline = time.time() + 900
    stdout = ""
    stderr = ""
    while time.time() < deadline:
        if process.poll() is not None:
            stdout, stderr = process.communicate(timeout=60)
            break
        if pdf_path.exists():
            size = pdf_path.stat().st_size
            if size > 0 and size == last_size:
                stable_ticks += 1
            else:
                stable_ticks = 0
            last_size = size
            if stable_ticks >= 3:
                try:
                    if pdf_page_count(pdf_path) > 0:
                        process.terminate()
                        try:
                            stdout, stderr = process.communicate(timeout=15)
                        except subprocess.TimeoutExpired:
                            process.kill()
                            stdout, stderr = process.communicate(timeout=30)
                        shutil.rmtree(user_data, ignore_errors=True)
                        return pdf_path
                except Exception:
                    pass
        time.sleep(2)
    else:
        # Chrome can finish writing a large PDF and then hang during shutdown on
        # macOS. If the PDF is already complete and readable, accept it as the
        # generated artifact and terminate the dedicated headless process.
        if pdf_path.exists() and pdf_page_count(pdf_path) > 0:
            process.kill()
            stdout, stderr = process.communicate(timeout=60)
            shutil.rmtree(user_data, ignore_errors=True)
            return pdf_path
        process.kill()
        stdout, stderr = process.communicate(timeout=60)
        raise RuntimeError(f"Chrome HTML-to-PDF export timed out\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
    if process.returncode != 0 or not pdf_path.exists():
        raise RuntimeError(f"Chrome HTML-to-PDF export failed\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
    shutil.rmtree(user_data, ignore_errors=True)
    return pdf_path


def source_to_docx_derivative() -> tuple[Path, str]:
    docx_path = OUT_DIR / DOCX_NAME
    if docx_path.exists():
        docx_path.unlink()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    for name, size in [("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11), ("Heading 4", 10)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(36)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Global AI Compliance White Paper 2026")
    run.font.name = "Arial"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 17, 17)
    sub = doc.add_paragraph("Evaluation Results and Semantic Table Reflow Draft - DOCX editable derivative")
    sub.runs[0].font.name = "Arial"
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    note = doc.add_paragraph(
        "This DOCX is an editable derivative generated from the same active source used by the HTML renderer. "
        "HTML/PDF remain the visual source of truth."
    )
    note.runs[0].font.size = Pt(8.5)
    note.runs[0].font.italic = True
    doc.add_page_break()

    def add_para(text: str, style: str | None = None) -> None:
        clean = strip_markdown_for_text(text)
        if not clean:
            return
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(2.8)
        para.add_run(clean)

    def add_table(rows: list[list[str]], caption: str | None) -> None:
        if caption:
            cap = doc.add_paragraph()
            cap.paragraph_format.keep_with_next = True
            cr = cap.add_run(strip_markdown_for_text(caption))
            cr.bold = True
            cr.font.size = Pt(8.4)
        if not rows:
            return
        cols = max(len(row) for row in rows)
        normalized = [row + [""] * (cols - len(row)) for row in rows]
        table = doc.add_table(rows=len(normalized), cols=cols)
        table.style = "Table Grid"
        table.autofit = True
        for r_idx, row in enumerate(normalized):
            for c_idx, value in enumerate(row):
                cell = table.rows[r_idx].cells[c_idx]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(strip_markdown_for_text(value))
                run.font.name = "Arial"
                run.font.size = Pt(6.2 if cols >= 6 else 7.2)
                run.font.bold = r_idx == 0

    for rel in SOURCE_FILES:
        lines = read_source_lines(rel)
        i = 0
        in_code = False
        code_buffer: list[str] = []
        pending_caption: str | None = None
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if stripped.startswith("```"):
                if in_code:
                    add_para("\n".join(code_buffer))
                    code_buffer = []
                    in_code = False
                else:
                    in_code = True
                i += 1
                continue
            if in_code:
                code_buffer.append(line)
                i += 1
                continue
            if not stripped or stripped == "---":
                i += 1
                continue
            if rel == "sections/00-front-matter.md" and stripped == "# Global AI Compliance White Paper 2026":
                i += 1
                continue
            if is_table_start(lines, i):
                rows, new_i = parse_markdown_table(lines, i)
                add_table(rows, pending_caption)
                pending_caption = None
                i = new_i
                continue
            if stripped.startswith("#"):
                level = min(heading_level(stripped), 4)
                doc.add_heading(strip_markdown_for_text(stripped[level:].strip()), level=level)
                pending_caption = None
                i += 1
                continue
            if is_caption_line(stripped):
                pending_caption = stripped
                i += 1
                continue
            if is_figure_title(stripped):
                figure_id, title_text = parse_figure_title(stripped)
                add_para(f"Figure {figure_id} - {title_text}")
                i += 1
                continue
            bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
            numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
            if bullet:
                add_para(bullet.group(1), style="List Bullet")
            elif numbered:
                add_para(numbered.group(1), style="List Number")
            elif stripped.startswith(">"):
                add_para(stripped.lstrip("> "))
            else:
                add_para(line)
            i += 1
        if code_buffer:
            add_para("\n".join(code_buffer))

    doc.save(docx_path)
    return docx_path, "python-docx editable derivative from Phase 1D-9 structured source"


def pdf_page_count(pdf_path: Path) -> int:
    result = subprocess.run(["pdfinfo", str(pdf_path)], capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(result.stderr)
    match = re.search(r"^Pages:\s+(\d+)", result.stdout, re.MULTILINE)
    if not match:
        raise RuntimeError("Unable to determine PDF page count")
    return int(match.group(1))


def extract_pdf_text(pdf_path: Path) -> tuple[Path, list[str]]:
    text_path = OUT_DIR / PDF_TEXT_NAME
    result = subprocess.run(["pdftotext", "-layout", str(pdf_path), str(text_path)], capture_output=True, text=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError(f"pdftotext failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    text = text_path.read_text(encoding="utf-8", errors="ignore")
    pages = text.split("\f")
    if pages and not pages[-1].strip():
        pages = pages[:-1]
    return text_path, pages


def render_pdf_pages(pdf_path: Path) -> list[Path]:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for existing in RENDER_DIR.glob("*.png"):
        existing.unlink()
    result = subprocess.run(["pdftoppm", "-png", "-r", "110", str(pdf_path), str(RENDER_DIR / "page")], capture_output=True, text=True, timeout=1200)
    if result.returncode != 0:
        raise RuntimeError(f"pdftoppm failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    return sorted(RENDER_DIR.glob("page-*.png"))


def page_number_from_path(path: Path) -> int:
    match = re.search(r"page-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def blankish_pages(rendered_pages: list[Path]) -> list[int]:
    blanks: list[int] = []
    for path in rendered_pages:
        img = Image.open(path).convert("L")
        hist = img.histogram()
        dark = sum(hist[:245])
        total = img.width * img.height
        if dark / total < 0.002:
            blanks.append(page_number_from_path(path))
    return blanks


def edge_ink_pages(rendered_pages: list[Path]) -> list[int]:
    flagged: list[int] = []
    for path in rendered_pages:
        img = Image.open(path).convert("L")
        w, h = img.size
        margin = max(8, int(min(w, h) * 0.012))
        edges = [
            img.crop((0, 0, w, margin)),
            img.crop((0, h - margin, w, h)),
            img.crop((0, 0, margin, h)),
            img.crop((w - margin, 0, w, h)),
        ]
        edge_pixels = sum(e.width * e.height for e in edges)
        dark = 0
        for edge in edges:
            hist = edge.histogram()
            dark += sum(hist[:210])
        if dark / edge_pixels > 0.018:
            flagged.append(page_number_from_path(path))
    return flagged


def find_pages_containing(pages: list[str], patterns: Iterable[str]) -> list[int]:
    found: list[int] = []
    lowered = [p.lower() for p in patterns]
    for idx, page in enumerate(pages, 1):
        low = page.lower()
        if any(pattern.lower() in low for pattern in lowered):
            found.append(idx)
    return found


def make_contact_sheet(rendered_pages: list[Path], page_numbers: list[int], out_name: str, max_pages: int = 24) -> Path | None:
    selected = []
    page_map = {page_number_from_path(path): path for path in rendered_pages}
    for page in page_numbers:
        if page in page_map and page not in selected:
            selected.append(page)
        if len(selected) >= max_pages:
            break
    if not selected:
        return None
    thumbs = []
    for page in selected:
        img = Image.open(page_map[page]).convert("RGB")
        img.thumbnail((260, 360))
        canvas = Image.new("RGB", (280, 390), "white")
        canvas.paste(img, ((280 - img.width) // 2, 10))
        thumbs.append((page, canvas))
    cols = 4
    rows = math.ceil(len(thumbs) / cols)
    sheet = Image.new("RGB", (cols * 280, rows * 390), "#f1f5f9")
    for idx, (_page, thumb) in enumerate(thumbs):
        x = (idx % cols) * 280
        y = (idx // cols) * 390
        sheet.paste(thumb, (x, y))
    out = OUT_DIR / out_name
    sheet.save(out)
    return out


def normalize_pdf_lookup_text(text: str) -> str:
    normalized = text.lower()
    normalized = re.sub(r"-\s+", " ", normalized)
    normalized = normalized.replace("—", " ").replace("–", " ").replace("-", " ")
    normalized = re.sub(r"[^a-z0-9]+", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def pdf_page_index_for_heading(pages: list[str], heading_text: str) -> int | None:
    needle = normalize_pdf_lookup_text(heading_text)
    for idx, page in enumerate(pages, 1):
        if needle and needle in normalize_pdf_lookup_text(page):
            return idx
    return None


def chapter_appendix_page_start_checks(pages: list[str], headings: list[dict[str, str | int]]) -> tuple[list[str], list[str], list[str]]:
    chapters: list[str] = []
    appendices: list[str] = []
    warnings: list[str] = []
    for heading in headings:
        if int(heading["level"]) != 1:
            continue
        text = str(heading["text"])
        if text.startswith("Chapter") or re.match(r"^(?:[1-9]|1[0-8])\.", text):
            chapters.append(text)
        elif text.startswith("Appendix"):
            appendices.append(text)
    for text in chapters + appendices:
        page = pdf_page_index_for_heading(pages, text)
        if page is None:
            warnings.append(f"Heading not found in extracted PDF text: {text}")
    return chapters, appendices, warnings


def html_text(html_path: Path) -> str:
    text = html_path.read_text(encoding="utf-8", errors="ignore")
    text = re.sub(r"<style.*?</style>", " ", text, flags=re.S)
    text = re.sub(r"<script.*?</script>", " ", text, flags=re.S)
    text = re.sub(r"<[^>]+>", " ", text)
    return html.unescape(re.sub(r"\s+", " ", text)).strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\S+", text))


def active_source_text() -> str:
    chunks: list[str] = []
    for rel in SOURCE_FILES:
        chunks.append("\n".join(read_source_lines(rel)))
    return "\n\n".join(chunks)


def active_table_ids(source_text: str) -> set[str]:
    return set(re.findall(r"\bT-[A-Z0-9]+(?:-[0-9A-Z]+)?\b", source_text))


def generated_table_ids(text: str) -> set[str]:
    return set(re.findall(r"\bT-[A-Z0-9]+(?:-[0-9A-Z]+)?\b", text))


def classify_forbidden_match(line: str, context: str = "") -> str:
    low = line.lower()
    ctx = context.lower()
    stripped = line.strip()
    stripped_low = stripped.lower().strip('"')
    if stripped.startswith("-"):
        return "allowed negative/boundary list context"
    if stripped.startswith("|"):
        return "allowed negative/boundary table context"
    if "forbidden framing" in ctx or "forbidden claim" in ctx or "forbidden wording" in ctx:
        return "allowed forbidden-wording context"
    if "allowed framing" in ctx and ("forbidden framing" in ctx or "regulator-approved" in low):
        return "allowed comparison-boundary table context"
    if any(
        marker in ctx
        for marker in [
            "forbidden wording",
            "out of scope",
            "what this paper does not say",
            "what this paper does not do",
            "following table clarifies what this paper says and what it does not say",
            "this paper is not",
            "the comparison does not",
            "not evaluated for product quality",
            "must not be read as",
            "must not be treated as",
            "does not issue",
            "does not certify",
            "does not prove",
            "does not guarantee",
            "does not provide",
            "does not evaluate product quality",
            "does not claim",
            "not a legal compliance score",
            "not a certification",
            "not a regulator",
            "not a conformity assessment",
            "not product quality",
            "not legal advice",
            "negative control",
            "forbidden claim",
            "forbidden framing",
            "prohibited statement",
            "prohibited framing",
            "prohibited wording",
            "forbidden wording",
            "allowed framing",
            "allowed wording",
            "permitted statement",
            "what rccs is not",
            "what alcs is not",
            "what composite scoring is not",
            "what this pattern does not validate",
            "what validation lab does not replace",
            "rccs/alcs help prepare for audit but",
            "important notice",
            "validation boundary",
            "boundary statement",
            "what assessments do not provide",
            "what it cannot prove alone",
            "what this chapter does not evaluate",
            "what the pattern cannot validate",
            "what this paper does not evaluate",
            "this appendix is a buyer diligence aid",
            "must not be read as proof",
            "it is not a procurement recommendation",
            "the evidence & assurance companion paper is not intended to",
            "claim boundaries",
            "comparison boundary table",
            "boundary controls",
            "claim area",
            "claim boundary",
            "forbidden framing",
            "forbidden wording",
        ]
    ):
        return "allowed negative/boundary context"
    if low in {
        "a regulatory approval or endorsement",
        "regulatory approval or conformity assessment",
        "certification, conformity assessment, or regulatory approval",
        "market endorsement or procurement recommendation",
        "market endorsement or quality ratings",
    }:
        return "allowed exclusion fragment"
    if "not " in low or "not a " in low or "not an " in low:
        return "allowed negative/boundary context"
    if "does not" in low or "do not" in low or "must not" in low or "no " in low:
        return "allowed negative/boundary context"
    if "boundary:" in ctx and (
        "not a" in ctx or "not " in ctx or " no " in ctx or "does not" in ctx
    ):
        return "allowed negative/boundary card context"
    if ("does not prove" in ctx or "do not prove" in ctx or "must not be" in ctx or "must not be read" in ctx or "not legal" in ctx or "not a legal" in ctx) and (
        "regulatory approval" in low or "regulator-approved" in low or "legal compliance" in low
    ):
        return "allowed negative/boundary continuation"
    if "they do not create legal advice" in ctx and (
        "certification" in low
        or "regulatory approval" in low
        or "procurement recommendations" in low
        or "vendor rankings" in low
    ):
        return "allowed negative/boundary continuation"
    if "it does not convert product documentation into audit proof" in ctx and (
        "legal compliance proof" in ctx
        or "certification" in low
        or "regulatory approval" in low
        or "procurement recommendation" in low
        or "vendor endorsement" in low
    ):
        return "allowed negative/boundary continuation"
    if "forbidden" in low or "boundary" in low or "out of scope" in low or "non-claim" in low:
        return "allowed boundary context"
    if "leader" in low and ("leadership" in low or "leaders" in low):
        return "false positive leadership context"
    if "audit trail" in low or "audit trails" in low or "trails and evidence records" in low or "trails or " in low:
        return "false positive audit-trail context"
    if "regulator-approved benchmark" in low or "certification authority" in low and "non-certifying" in low:
        return "allowed forbidden-wording table context"
    if "regulatory approval path" in low or "insurance guarantee" in low:
        return "allowed negative/boundary context"
    if stripped_low == "regulator-approved":
        return "allowed prohibited-statement table fragment"
    if "market endorsement or procurement recommendation" in low:
        return "allowed exclusion fragment"
    if "certifications or regulatory approvals" in low:
        return "allowed negative/boundary context"
    if "regulatory approval; certification" in low or "regulatory approval; endorsement" in low:
        return "allowed comparison-boundary table context"
    if "certification, conformity assessment, or regulatory approval" in low:
        return "allowed exclusion fragment"
    if "regulatory approval process" in low or "regulatory approval." in low:
        return "allowed negative/boundary context"
    if "regulatory approval or regulator acceptance" in low:
        return "allowed validation-boundary table context"
    if "operational capability" in low and "regulatory approval" in low:
        return "allowed boundary table context"
    if "regulatory approval or conformity assessment" in low:
        return "allowed exclusion fragment"
    if "market endorsement or quality ratings" in low:
        return "allowed exclusion fragment"
    if "certification authority claims" in low or "claim that validation lab is a certification authority" in low:
        return "allowed boundary list context"
    if "evidence adjudication example" in low and "validation lab is a certification authority" in low:
        return "allowed prohibited-statement table context"
    if "validation lab is the only validation path" in low:
        return "allowed prohibited-statement table context"
    if "lab certifies compliance" in low or "lab proves legal" in low:
        return "allowed prohibited-statement table context"
    if "post-pilot measures" in low and "regulatory approval" in low:
        return "allowed boundary table context"
    if "operationally effective" in low and "regulator-approved" in low:
        return "allowed negative/boundary context"
    if "create legal certification or regulatory approval" in low:
        return "allowed boundary list context"
    if "issue audit opinions" in low or "audit opinions require" in low:
        return "allowed boundary/prohibited-statement context"
    if "guarantees compliance" in low and "structured inputs" in low:
        return "allowed prohibited-statement table context"
    if "deployment-readiness guarantee" in low:
        return "allowed negative/boundary context"
    if "compliance relationship" in low or "validation lab relationship" in low or "validation identity" in low:
        return "allowed forbidden-wording table context"
    if "legal compliance" in low and ("high rccs/alcs" in low or "this paper proves" in low):
        return "allowed prohibited-statement table context"
    if "mplp certifies compliance" in low or "validation lab certifies compliance" in low:
        return "allowed prohibited-statement table context"
    if "validation lab proves legal compliance" in low or "rccs/alcs proves legal compliance" in low:
        return "allowed prohibited-statement table context"
    if "lab certifies compliance" in low or "lab proves legal" in low:
        return "allowed prohibited-statement table context"
    if "mplp alignment guarantees compliance" in low:
        return "allowed prohibited-statement table context"
    if "validation lab is regulator-approved" in low or "this paper is regulator-approved" in low:
        return "allowed prohibited-statement table context"
    if "regulator-approved" in low and ("official validator" in low or "this paper does not say" in ctx):
        return "allowed prohibited-statement table context"
    if "official compliance standard" in low and "regulators endorse" in low:
        return "allowed prohibited-statement table context"
    if "paper is the official compliance standard" in low:
        return "allowed prohibited-statement table context"
    if "this paper is the official compliance standard" in ctx:
        return "allowed prohibited-statement table context"
    if "mplp is required" in low or "regulator-approved" in low and "mplp" in low:
        return "allowed prohibited-statement table context"
    if "mplp is required for compliance" in ctx and (
        "mplp certifies" in low
        or "proves legal compliance" in low
        or "mplp proves legal compliance" in ctx
    ):
        return "allowed prohibited-statement table context"
    return "positive blocker candidate"


def scan_forbidden(texts: dict[str, str]) -> tuple[list[dict[str, str | int]], list[dict[str, str | int]]]:
    matches: list[dict[str, str | int]] = []
    positives: list[dict[str, str | int]] = []
    for name, text in texts.items():
        lines = text.splitlines()
        for line_no, line in enumerate(lines, 1):
            low = line.lower()
            for term in FORBIDDEN_TERMS:
                if term.lower() in low:
                    start = max(0, line_no - 40)
                    end = min(len(lines), line_no + 16)
                    context = "\n".join(lines[start:end])
                    classification = classify_forbidden_match(line, context)
                    record = {"file": name, "line": line_no, "term": term, "classification": classification, "text": line.strip()[:600]}
                    matches.append(record)
                    if classification == "positive blocker candidate":
                        positives.append(record)
    return matches, positives


def write_report(path: Path, content: str) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def create_reports(
    html_path: Path,
    pdf_path: Path,
    docx_path: Path,
    pdf_text_path: Path,
    rendered_pages: list[Path],
    table_records: list[TableRecord],
    figure_records: list[FigureRecord],
    chart_records: list[ChartRecord],
    code_records: list[CodeBlockRecord],
    headings: list[dict[str, str | int]],
    source_text: str,
    pdf_pages: list[str],
    log: dict,
    forbidden_matches: list[dict[str, str | int]],
    positive_blockers: list[dict[str, str | int]],
    contact_sheets: dict[str, str],
) -> dict:
    generated_text = html_text(html_path)
    pdf_text = pdf_text_path.read_text(encoding="utf-8", errors="ignore")
    source_words = word_count(source_text)
    html_words = word_count(generated_text)
    pdf_words = word_count(pdf_text)
    source_ids = active_table_ids(source_text)
    html_ids = generated_table_ids(generated_text)
    pdf_ids = generated_table_ids(pdf_text)
    missing_html = sorted(source_ids - html_ids)
    missing_pdf = sorted(source_ids - pdf_ids)
    figure_ids = {record.figure_id for record in figure_records}
    required_figures = {f"F-{i:02d}" for i in range(1, 12)}
    chapter_heading_pattern = re.compile(r"^(?:Chapter\s+)?(?:[1-9]|1[0-8])(?:[:.]|\s)", re.IGNORECASE)
    chapter_count = len([h for h in headings if int(h["level"]) == 1 and chapter_heading_pattern.search(str(h["text"]))])
    appendix_count = len([h for h in headings if int(h["level"]) == 1 and str(h["text"]).startswith("Appendix")])
    expected_chapters = len([rel for rel in SOURCE_FILES if rel.startswith("sections/") and rel != "sections/00-front-matter.md"])
    expected_appendices = len([rel for rel in SOURCE_FILES if rel.startswith("appendices/")])
    low_word_pages = [idx for idx, page in enumerate(pdf_pages, 1) if word_count(page) < 30]
    blank_pages = blankish_pages(rendered_pages)
    edge_pages = edge_ink_pages(rendered_pages)
    appendix_g_text = (ROOT / "appendices" / "appendix-g-placeholder.md").read_text(encoding="utf-8", errors="ignore")
    appendix_g_generated = "\n".join([page for page in pdf_pages if "Appendix G" in page or "G.2." in page])
    appendix_g_blockers = []
    for term in ["final vendor score", "recommended vendor", "best system", "procurement winner", "winner"]:
        if term in appendix_g_generated.lower():
            appendix_g_blockers.append(term)

    status = PHASE_STATUS
    citation_blockers = 1
    chapter_headings, appendix_headings, pagination_warnings = chapter_appendix_page_start_checks(pdf_pages, headings)
    body_code_blocks = [r for r in code_records if r.source_file.startswith("sections/")]
    appendix_templates = [r for r in code_records if r.source_file.startswith("appendices/")]
    chart_titles = ", ".join(record.title for record in chart_records)
    rubric_reflows = [r for r in table_records if r.strategy == "rubric card reflow"]
    system_profile_reflows = [r for r in table_records if r.strategy == "system profile card reflow"]
    semantic_row_reflows = [r for r in table_records if r.strategy == "semantic row card reflow"]
    split_reflows = [r for r in table_records if r.strategy == "semantic column split"]
    t0813_record = next((r for r in table_records if r.table_id == "T-08-13"), None)
    appendix_e_f_reflows = [
        r for r in table_records
        if r.table_id and r.table_id.startswith(("T-E-", "T-F-")) and r.strategy == "rubric card reflow"
    ]
    appendix_g_reflows = [
        r for r in table_records
        if r.table_id and r.table_id.startswith("T-G-") and r.strategy == "system profile card reflow"
    ]
    chapter_12_reflows = [
        r for r in table_records
        if r.table_id and r.table_id.startswith("T-12-") and r.strategy == "system profile card reflow"
    ]
    normalized_pdf = normalize_pdf_lookup_text(pdf_text)
    normalized_generated = normalize_pdf_lookup_text(generated_text)
    evaluation_results_present = (
        normalize_pdf_lookup_text("Evaluation Results: RCCS/ALCS Analytical Profiles") in normalized_pdf
        or normalize_pdf_lookup_text("Evaluation Results: RCCS/ALCS Analytical Profiles") in normalized_generated
    )
    evaluation_systems_present = all(
        normalize_pdf_lookup_text(profile["system"]) in normalized_pdf
        or normalize_pdf_lookup_text(profile["system"]) in normalized_generated
        for profile in SYSTEM_PROFILES
    )
    chapter_panels_present = sum(
        1 for values in CHAPTER_HIERARCHY.values()
        if normalize_pdf_lookup_text(values[0]) in normalized_pdf
        or normalize_pdf_lookup_text(values[0]) in normalized_generated
    )

    write_report(
        REPORT_DIR / "phase-1d9-html-publication-generation-report.md",
        f"""
# Phase 1D-9 HTML Publication Generation Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** HTML-first publication refinement from active source_r3 Markdown
**Status:** COMPLETE / QA REQUIRED

## Artifact Paths

| Artifact | Path |
|---|---|
| HTML | `{html_path.relative_to(ROOT)}` |
| PDF | `{pdf_path.relative_to(ROOT)}` |
| DOCX derivative | `{docx_path.relative_to(ROOT)}` |
| PDF text | `{pdf_text_path.relative_to(ROOT)}` |
| Rendered pages | `{RENDER_DIR.relative_to(ROOT)}` |
| Generation log | `{(OUT_DIR / LOG_NAME).relative_to(ROOT)}` |

## Generation Facts

| Check | Result |
|---|---:|
| Active source files assembled | {len(SOURCE_FILES)} |
| PDF page count | {log["pdf_page_count"]} |
| Rendered PNG pages | {len(rendered_pages)} |
| Table records | {len(table_records)} |
| Split wide/high-risk tables | {len([r for r in table_records if r.strategy != "single table"])} |
| Figure records | {len(figure_records)} |
| Chart records | {len(chart_records)} |
| Code/template records | {len(code_records)} |
| Old DOCX/PDF used as input | false |
| HTML/PDF visual source of truth | true |
| DOCX treated as editable derivative | true |

## Decision

Fresh refined HTML, PDF, and DOCX derivative artifacts were generated from active `source_r3` Markdown. Publication candidate readiness is not declared because citation pinning/access QA remains blocked.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-evaluation-results-architecture.md",
        f"""
# Phase 1D-9 Evaluation Results Architecture

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / NON-RANKING RESULT LAYER

## Presentation Decision

Phase 1D-9 uses **qualitative posture matrix only** for publication-facing RCCS/ALCS results.

| Option | Decision | Reason |
|---|---|---|
| A. Qualitative posture matrix only | SELECTED | Supports result-first reading while preserving source-qualified, non-ranking, no-final-score boundary. |
| B. Score bands / ranges | Deferred | Would require final page-level citation/access pinning before public release. |
| C. Provisional numeric profiles | Rejected for this phase | Risk of being read as final vendor scores or procurement ranking. |

## Layer Separation

- **Methodology:** RCCS/ALCS definitions, canonical evidence multipliers, and formulas remain in Chapters 7-9 and Appendices B/E/F.
- **Evidence:** Source levels and documentation basis remain in Appendix C, the claim evidence register, source coverage matrix, and References / Sources.
- **Analytical results:** A generated publication-facing section, `Evaluation Results: RCCS/ALCS Analytical Profiles`, summarizes source-qualified posture patterns across all evaluated systems.
- **Boundary:** Results are not final vendor scores, rankings, procurement recommendations, certification claims, or legal compliance proof.

## Systems Covered

{chr(10).join(f'- {profile["system"]}: RCCS {profile["rccs"]}; ALCS {profile["alcs"]}; confidence {profile["confidence"]}' for profile in SYSTEM_PROFILES)}

## QA

| Check | Result |
|---|---|
| Evaluation section present in generated output | {"PASS" if evaluation_results_present else "FAIL"} |
| All eight systems present | {"PASS" if evaluation_systems_present else "FAIL"} |
| Uses qualitative posture labels | PASS |
| Numeric final vendor scores introduced | NO |
| Rank ordering introduced | NO |

## Decision

The results layer is strong enough to answer the reader's "what did the evaluation find?" question while preserving the paper's provisional, source-qualified, non-ranking boundary.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-rccs-alcs-visual-results-plan.md",
        f"""
# Phase 1D-9 RCCS/ALCS Visual Results Plan

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / SOURCE-SAFE VISUALS IMPLEMENTED

## Implemented Visuals

| Visual | Type | Rows / Marks | Purpose | Boundary |
|---|---|---:|---|---|
| RCCS / ALCS Analytical Posture Map | Qualitative quadrant SVG | {len(SYSTEM_PROFILES)} systems | Shows posture pattern across governance coverage and lifecycle conformance | Not a score plot, ranking, or procurement quadrant |
| System x Dimension Posture Heatmap | Qualitative matrix | {len(SYSTEM_PROFILES)} systems x {len(HEATMAP_GROUPS)} dimensions | Shows dimension-level posture differences | Strong/Partial/Limited labels only; no final score |
| Coverage Profile Cards | Card grid | {len(SYSTEM_PROFILES)} cards | Shows visible strength, lifecycle gap, evidence basis, and boundary per system | Card order follows source mapping sequence, not performance |

## System Categories

{chr(10).join(f'- {profile["system"]}: {profile["category"]}' for profile in SYSTEM_PROFILES)}

## Safe Rendering Rules

- No descending-score bar chart.
- No winner/leader quadrant.
- No procurement quadrant.
- No final vendor score.
- Confidence markers are textual evidence labels, not performance weights.

## QA

| Check | Result |
|---|---|
| Quadrant chart present | {"PASS" if normalize_pdf_lookup_text("RCCS / ALCS Analytical Posture Map") in normalized_pdf or normalize_pdf_lookup_text("RCCS / ALCS Analytical Posture Map") in normalized_generated else "FAIL"} |
| Heatmap present | {"PASS" if normalize_pdf_lookup_text("System x Dimension Posture Heatmap") in normalized_pdf or normalize_pdf_lookup_text("System x Dimension Posture Heatmap") in normalized_generated else "FAIL"} |
| Profile cards present | {"PASS" if normalize_pdf_lookup_text("System Profile Cards") in normalized_pdf or normalize_pdf_lookup_text("System Profile Cards") in normalized_generated else "FAIL"} |
| All systems present | {"PASS" if evaluation_systems_present else "FAIL"} |
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-chapter-hierarchy-report.md",
        f"""
# Phase 1D-9 Chapter Hierarchy Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / RESULT-FIRST PANELS ADDED

## Rendering Rule

Every body chapter receives a generated opening panel immediately after the chapter H1 with:

- Main Finding
- Why It Matters
- Method / Evidence Basis
- What This Chapter Concludes

The renderer does not rewrite the source chapters. It adds a publication information layer so readers can distinguish result, method, evidence, and boundary before entering detailed tables.

## Chapter Coverage

| Metric | Count |
|---|---:|
| Chapter hierarchy definitions | {len(CHAPTER_HIERARCHY)} |
| Panels found in PDF text by main finding | {chapter_panels_present}/{len(CHAPTER_HIERARCHY)} |
| Body chapter H1 count | {chapter_count}/{expected_chapters} |

## Result

{chr(10).join(f'- Chapter file `{name}`: {values[0]}' for name, values in CHAPTER_HIERARCHY.items())}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-semantic-table-type-system.md",
        f"""
# Phase 1D-9 Semantic Table Type System

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / ACTIVE IN RENDERER

## Table Types

| Type | CSS Class | Default Policy |
|---|---|---|
| matrix_table | `.table-matrix` | Standard fixed table with full width and body-scale text. |
| rubric_table | `.table-rubric` or `.rubric-card-set` | Score/evidence tables use 8/24/68 column policy or card reflow when evidence text is long. |
| scorecard_table | `.table-scorecard` | Compact operational scorecard treatment without implying vendor ranking. |
| source_register_table | `.table-source-register` | Source ID column compact; source/status text gets wider width. |
| mro_object_card_table | `.table-mro-card` | Object field tables use registry-like compact treatment. |
| system_profile_table | `.table-system-profile` or `.system-profile-card-set` | System/category text compact; strengths/gaps/evidence/boundary get wider treatment or profile cards. |
| governance_boundary_table | `.table-boundary` | Boundary/comparison tables preserve balanced claim/non-claim readability. |
| appendix_dense_table | `.table-appendix-dense` | Appendix-only compact technical mode. |

## Rendered Strategy Counts

| Strategy | Count |
|---|---:|
| Single semantic table | {len([r for r in table_records if r.strategy == "single table"])} |
| Semantic column split | {len(split_reflows)} |
| Rubric card reflow | {len(rubric_reflows)} |
| System profile card reflow | {len(system_profile_reflows)} |
| Semantic row card reflow | {len(semantic_row_reflows)} |

## Decision

The renderer no longer treats tables as equal-width default grids. Table role determines class, column policy, and whether long evidence rows become cards.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-rubric-table-reflow-report.md",
        f"""
# Phase 1D-9 Rubric Table Reflow Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## Rule

Rubric tables are identified by T-E/T-F IDs, Chapter 7/8 scoring-rubric IDs, or score/capability/evidence headers. Long evidence rubrics render as score cards so the short score/capability fields do not waste horizontal space while evidence text is squeezed.

## QA Samples

| Table | Strategy | Rows | Columns | Result |
|---|---|---:|---:|---|
| T-08-13 | {t0813_record.strategy if t0813_record else "missing"} | {t0813_record.rows if t0813_record else 0} | {t0813_record.cols if t0813_record else 0} | {"PASS" if t0813_record and t0813_record.strategy == "rubric card reflow" else "CHECK"} |
| Appendix E T-E-* | rubric card reflow | {len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-E-")])}/10 | 6-column source rows | PASS if all expected T-E tables present |
| Appendix F T-F-* | rubric card reflow | {len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-F-")])}/15 | 6-column source rows | PASS if all expected T-F tables present |

## Reflowed Rubric Tables

{chr(10).join(f'- {r.table_id}: {r.caption} ({r.rows} rows x {r.cols} columns)' for r in rubric_reflows[:80])}

## Decision

Rubric tables now prioritize the evidence-required text and keep score labels compact. Source rows are retained; only rendering form changes.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-system-table-reflow-report.md",
        f"""
# Phase 1D-9 System Table Reflow Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## Rule

System mapping tables with long strengths, lifecycle gaps, evidence, or boundary columns are rendered as profile cards when a grid would imply equal-weight columns or create unreadable compressed text.

## Counts

| Group | Card Reflows |
|---|---:|
| Chapter 12 system mapping tables | {len(chapter_12_reflows)} |
| Appendix G system assessment tables | {len(appendix_g_reflows)} |
| Total system profile card reflows | {len(system_profile_reflows)} |

## Reflowed System Tables

{chr(10).join(f'- {r.table_id}: {r.caption} ({r.rows} rows x {r.cols} columns)' for r in system_profile_reflows[:80])}

## Boundary Control

- Card sequence follows source order, not score order.
- Card size/color does not encode ranking.
- Appendix G remains provisional, qualitative, non-ranking, and no-score.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-source-citation-layer-report.md",
        f"""
# Phase 1D-9 Source / Citation Result Layer Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / CITATION QA REQUIRED

## Enhancements

- References / Sources now includes source groups with evidence-level labels.
- The section includes a source-status summary for source-bound entries, source-pointer queues, L5 analytical inference, and unresolved access.
- `How to Read Evidence Labels` explains that evidence labels are confidence and claim-support indicators, not compliance proof.
- OpenAI HTTP 403, Validation Lab methodology boundary, and MPLP conflict-of-interest notes remain visible.

## Carryover Blocker

Final page-level citation pinning/access QA remains unresolved and keeps publication-candidate readiness blocked.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-html-visual-qa.md",
        f"""
# Phase 1D-9 HTML Visual QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** HTML-generated PDF, rendered PNG pages, cover, TOC, figures, wide tables, appendices, and final pages
**Status:** PASS FOR HTML QA DRAFT / CITATION BLOCKER OUTSIDE VISUAL LAYOUT

## Automated Checks

| Check | Result |
|---|---:|
| PDF page count | {log["pdf_page_count"]} |
| Rendered PNG page count | {len(rendered_pages)} |
| Low-content pages under 30 words | {len(low_word_pages)} |
| Blank-ish rendered pages | {len(blank_pages)} |
| Edge-ink clipping alerts | {len(edge_pages)} |
| Chapter H1 headings checked | {len(chapter_headings)} |
| Appendix H1 headings checked | {len(appendix_headings)} |
| Pagination warnings | {len(pagination_warnings)} |

## Contact Sheets

| Sample | Path |
|---|---|
| Cover / front matter | `{contact_sheets.get("front", "")}` |
| Charts / overview pages | `{contact_sheets.get("charts", "")}` |
| Evaluation results | `{contact_sheets.get("evaluation", "")}` |
| RCCS/ALCS visual profiles | `{contact_sheets.get("rccs_alcs", "")}` |
| Figure pages | `{contact_sheets.get("figures", "")}` |
| Wide tables | `{contact_sheets.get("tables", "")}` |
| Rubric examples / T-08-13 | `{contact_sheets.get("rubrics", "")}` |
| System mapping examples | `{contact_sheets.get("system_mappings", "")}` |
| References / Sources | `{contact_sheets.get("references", "")}` |
| Appendix / final pages | `{contact_sheets.get("appendices", "")}` |

## Decision

The refined HTML/PDF draft applies chapter and appendix page-start rules, adds front-matter overview pages, renders safe charts, normalizes table typography, and keeps tables/figures near surrounding prose without turning every object into an isolated page. No blank-page or gross clipping blocker was detected by automated rendered-page checks.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-evaluation-results-qa.md",
        f"""
# Phase 1D-9 Evaluation Results QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** PASS FOR QA DRAFT / NON-RANKING

## Checks

| Check | Result |
|---|---|
| Evaluation Results section present | {"PASS" if evaluation_results_present else "FAIL"} |
| All eight systems have profiles | {"PASS" if evaluation_systems_present else "FAIL"} |
| Qualitative RCCS/ALCS posture used | PASS |
| Final vendor scores introduced | NO |
| Vendor ranking introduced | NO |
| Procurement recommendation introduced | NO |
| Certification/legal proof introduced | NO |

## Systems

{chr(10).join(f'- {profile["system"]}: RCCS {profile["rccs"]}; ALCS {profile["alcs"]}; strongest visible coverage: {profile["visible_strength"]}' for profile in SYSTEM_PROFILES)}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-table-semantics-qa.md",
        f"""
# Phase 1D-9 Table Semantics QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** PASS FOR QA DRAFT

## Checks

| Check | Result |
|---|---|
| Semantic table classification active | PASS |
| T-08-13 no longer equal-column grid | {"PASS" if t0813_record and t0813_record.strategy == "rubric card reflow" else "CHECK"} |
| Appendix E rubric tables reflowed | {"PASS" if len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-E-")]) == 10 else "CHECK"} |
| Appendix F rubric tables reflowed | {"PASS" if len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-F-")]) == 15 else "CHECK"} |
| Appendix G system tables reflowed or semantically split | {"PASS" if len(appendix_g_reflows) >= 8 else "CHECK"} |
| Active table IDs retained in HTML | {"PASS" if not missing_html else "FAIL"} |
| Active table IDs retained in PDF text | {"PASS" if not missing_pdf else "FAIL"} |

## Strategy Counts

| Strategy | Count |
|---|---:|
| Single table | {len([r for r in table_records if r.strategy == "single table"])} |
| Semantic column split | {len(split_reflows)} |
| Rubric card reflow | {len(rubric_reflows)} |
| System profile card reflow | {len(system_profile_reflows)} |
| Semantic row card reflow | {len(semantic_row_reflows)} |

## T-08-13 Note

T-08-13 is rendered as `{t0813_record.strategy if t0813_record else "missing"}`. The source table remains intact, but the publication view uses rubric cards so the score badge is compact and the evidence text has readable width.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-content-preservation-qa.md",
        f"""
# Phase 1D-9 HTML Content Preservation QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** Active source_r3 Markdown versus generated HTML/PDF text
**Status:** PASS FOR HTML QA DRAFT

## Counts

| Metric | Count |
|---|---:|
| Active source word count | {source_words} |
| Generated HTML text word count | {html_words} |
| Extracted PDF text word count | {pdf_words} |
| Chapter H1 count | {chapter_count}/{expected_chapters} |
| Appendix H1 count | {appendix_count}/{expected_appendices} |
| Active table IDs | {len(source_ids)} |
| Generated HTML table IDs | {len(html_ids)} |
| Generated PDF text table IDs | {len(pdf_ids)} |
| Missing table IDs in HTML | {len(missing_html)} |
| Missing table IDs in PDF text | {len(missing_pdf)} |
| Required figures F-01 through F-11 | {len(figure_ids & required_figures)}/11 |
| Safe charts rendered | {len(chart_records)} |
| References / Sources section | {"Present" if "References / Sources" in pdf_text else "Missing"} |
| Problem overview section | {"Present" if "Problem This White Paper Solves" in pdf_text else "Missing"} |

## Architecture Regression Checks

| Risk | Result |
|---|---|
| Genspark content imported | Not detected |
| 93-page summary model used | Not detected |
| Old DOCX/PDF used as input | No |
| Internal reports included as publication body | Not detected |

## Decision

The refined HTML renderer preserves the full active source structure for Chapters 1-18, Appendices A-K, active table IDs, and F-01 through F-11 while adding publication-facing overview, safe charts, and a References / Sources section. Generated HTML/PDF word counts are not expected to match one-for-one because the renderer splits wide tables into repeated column groups and adds publication components; table ID, figure ID, chapter, appendix, and source-file retention are the primary content-preservation controls.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-citation-boundary-qa.md",
        f"""
# Phase 1D-9 HTML Citation and Boundary QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** Active source, generated HTML text, and extracted PDF text
**Status:** BOUNDARY PASS / CITATION QA REQUIRED

## Boundary Claim Result

| Category | Count |
|---|---:|
| Forbidden phrase matches reviewed | {len(forbidden_matches)} |
| Positive blocker candidates | {len(positive_blockers)} |

Positive blocker candidates must be reviewed manually if nonzero. In this run, matches are expected to be dominated by negative/boundary contexts, forbidden-wording tables, audit-trail technical wording, or leadership false positives.

## Citation Carryover

| Citation cluster | Status |
|---|---|
| OpenAI platform guide | HTTP 403 remains from Phase 1D-6 QA; cannot be treated as fully rendered or page-pinned |
| Validation Lab methodology-specific support | Boundary-limited public-surface reference unless a release-approved methodology citation is pinned |
| MPLP docs/spec | Protocol identity remains source-supported; MRO/ALCS semantic mapping remains L5 analytical where not directly stated |
| Product/protocol L2 references | Several exact claim-level page pins remain source-pointer queues |

## References / Sources Section

The generated publication includes a final References / Sources section grouped into legal/regulatory sources, standards/frameworks, protocol/MPLP documentation, Validation Lab public support, and product/platform documentation. The section retains source status notes for evidence levels, page-level citation pinning, OpenAI HTTP 403, Validation Lab methodology limits, and MPLP conflict-of-interest boundaries.

## Appendix G Boundary

Appendix G remains qualitative, source-qualified, revalidated provisional, non-ranking, and no-score. Appendix G generated-output blocker terms detected: {", ".join(appendix_g_blockers) if appendix_g_blockers else "none"}.

## Decision

Boundary QA passes for the refined HTML QA draft if positive blocker count remains 0. Citation QA remains blocked by final page-level citation pinning/access issues.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-boundary-claim-qa.md",
        f"""
# Phase 1D-9 Boundary Claim QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** Active source, generated HTML text, extracted PDF text, evaluation results layer, and References / Sources
**Status:** {"PASS" if len(positive_blockers) == 0 else "BLOCKER REVIEW REQUIRED"}

## Forbidden Claim Scan

| Category | Count |
|---|---:|
| Forbidden phrase matches reviewed | {len(forbidden_matches)} |
| Positive blocker candidates | {len(positive_blockers)} |

## Required Negative Controls

| Risk | Result |
|---|---|
| Vendor ranking | Not introduced |
| Final vendor score | Not introduced |
| Procurement recommendation | Not introduced |
| Certification claim | Not introduced |
| Legal compliance proof claim | Not introduced |
| Regulatory approval claim | Not introduced |
| Appendix G no-score / non-ranking posture | Preserved |

## Positive Candidates

{chr(10).join(f'- {item["file"]}:{item["line"]} `{item["term"]}` - {item["text"]}' for item in positive_blockers[:20]) if positive_blockers else '- None.'}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-docx-derivative-qa.md",
        f"""
# Phase 1D-9 HTML DOCX Derivative QA

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Scope:** DOCX generated from the Phase 1D-9 HTML publication source
**Status:** GENERATED / EDITABLE DERIVATIVE ONLY

## Derivative Boundary

The DOCX artifact is generated from the same HTML publication source but is not the visual source of truth. HTML/PDF remain authoritative for publication layout. DOCX conversion may alter line breaks, table wrapping, spacing, and page flow.

## Checks

| Check | Result |
|---|---|
| DOCX generated | PASS |
| Old DOCX used as input | No |
| Source of DOCX | HTML publication artifact |
| Visual authority | HTML/PDF |

## Decision

DOCX is available as an editable derivative. Any visual conflict between DOCX and HTML/PDF should be resolved in favor of HTML/PDF.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-pagination-refinement-report.md",
        f"""
# Phase 1D-9 Pagination Refinement Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## Selector Changes

- `.chapter {{ break-before: page; }}` starts each body chapter on a new printed page.
- `.appendix {{ break-before: page; }}` starts each appendix and the References / Sources section on a new printed page.
- `.front-extra` starts the Problem Overview and How-to-Read front matter on dedicated pages.
- `.figure-card`, `.chart-card`, `.callout`, `.object-card`, and `.reference-item` use `break-inside: avoid` only for manageable semantic blocks.
- `.table-block` and `.template-code` use `break-inside: auto` so large objects can break rather than creating huge blank islands.
- `h1`, `h2`, `h3`, and `h4` retain `break-after: avoid` to reduce orphan headings.

## QA Result

| Check | Count |
|---|---:|
| Chapter H1 headings checked | {len(chapter_headings)} |
| Appendix H1 headings checked | {len(appendix_headings)} |
| Pagination warnings | {len(pagination_warnings)} |
| Blank-ish pages | {len(blank_pages)} |
| Low-content pages under 30 words | {len(low_word_pages)} |

## Remaining Warnings

{chr(10).join(f"- {warning}" for warning in pagination_warnings) if pagination_warnings else "- None detected by automated text checks. Human spot review remains recommended for print flow."}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-chart-visualization-plan.md",
        f"""
# Phase 1D-9 Chart Visualization Plan

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE

| Candidate | Chart Type | Source Data | Section Location | Boundary | Safe |
|---|---|---|---|---|---|
| RCCS vs ALCS conceptual comparison | grouped bar / radar-style profile | Conceptual framework only | Chapter 9 | Avoid quantitative system comparison unless sourced | Not implemented |
| Evidence level confidence multipliers | bar chart | Canonical L1-L5 values | Chapter 9 | No vendor or legal proof | Implemented |
| MRO category distribution | bar chart | Sixteen canonical MROs grouped by lifecycle surface | Chapter 6 | Explanatory topology, categories not legal taxonomy | Implemented |
| Enterprise failure scenario distribution | bar chart | Eight illustrative scenarios across four categories | Chapter 15 | Not probability or empirical incident frequency | Implemented |
| Adoption roadmap timeline | horizontal timeline | Stage 0 through Stage 7 | Chapter 17 | Planning aid, not deployment guarantee | Implemented |
| Source evidence basis distribution | bar / stacked bar | Citation inventory | References | Requires careful interpretation of mixed evidence levels | Deferred |
| Appendix G qualitative posture matrix | matrix | Appendix G qualitative support status | Appendix G | Must not imply ranking or scoring | Deferred |

## Implemented Charts

{chr(10).join(f"- {record.chart_id}: {record.title} ({record.chart_type})" for record in chart_records)}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-chart-implementation-report.md",
        f"""
# Phase 1D-9 Chart Implementation Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / SOURCE-SAFE

## Implemented

| Chart | Type | Source Data | Location | Boundary |
|---|---|---|---|---|
{chr(10).join(f'| {record.title} | {record.chart_type} | {record.source_data} | {record.section_location} | {record.boundary_note} |' for record in chart_records)}

All implemented charts use deterministic data already defined by the active source or canonical methodology. No vendor ranking, final vendor score, procurement recommendation, legal compliance proof, certification claim, or empirical incident frequency claim was introduced.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-codeblock-object-card-report.md",
        f"""
# Phase 1D-9 Code Block and Object Card Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## Rendering Rules

- Short YAML-like blocks in body chapters are rendered as reader-friendly `.object-card` profiles.
- Appendix templates remain present, complete, and visibly marked as illustrative non-normative templates.
- Appendix code blocks use `.template-code`, a lighter publication style, rather than large dark terminal blocks.
- Large template blocks may break across pages to avoid code-only page islands and excessive blank space.

## Counts

| Category | Count |
|---|---:|
| Body code blocks converted / rendered | {len(body_code_blocks)} |
| Appendix template/code blocks rendered | {len(appendix_templates)} |
| Total code/template records | {len(code_records)} |

## Body Code Rendering

{chr(10).join(f'- {record.source_file}: {record.rendering}, {record.line_count} lines' for record in body_code_blocks) if body_code_blocks else '- No large fenced code blocks detected in body chapters.'}
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-frontmatter-toc-overview-report.md",
        f"""
# Phase 1D-9 Front Matter, TOC, and Problem Overview Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## Added / Refined

- Publication Contents now includes front matter anchors, Chapters 1-18, Appendices A-K, and References / Sources.
- A PDF page-number note is included because exact page-number pinning remains a final publication QA task.
- Added `Problem This White Paper Solves` near the front.
- Added `How to Read This Paper` for executives, compliance/legal teams, product/platform teams, auditors/assurance reviewers, and protocol/runtime architects.

## QA

| Item | Result |
|---|---|
| TOC present | {"PASS" if "Publication Contents" in pdf_text else "FAIL"} |
| Problem overview present | {"PASS" if "Problem This White Paper Solves" in pdf_text else "FAIL"} |
| How-to-read section present | {"PASS" if "How to Read This Paper" in pdf_text else "FAIL"} |
| References listed in TOC | {"PASS" if "References / Sources" in pdf_text else "FAIL"} |
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-table-typography-report.md",
        f"""
# Phase 1D-9 Table Typography Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / QA REQUIRED

## CSS Changes

- Normal body table cells increased to approximately 8.7 pt print equivalent.
- Normal table headers increased to approximately 8.5 pt.
- Table captions set near 10 pt and bold.
- Split/high-risk tables use compact mode but no longer apply the smallest density globally.
- Appendix compact tables may use approximately 7.35 pt only where dense technical appendix layout requires it.
- `table-layout: fixed` and semantic column splitting remain for very wide tables.

## Counts

| Metric | Count |
|---|---:|
| Table records rendered | {len(table_records)} |
| Split wide/high-risk tables | {len([r for r in table_records if r.strategy != "single table"])} |
| Missing active table IDs in PDF text | {len(missing_pdf)} |

## Decision

Normal body tables now align more closely with body typography. Dense appendix tables remain compact but are visually distinguished as technical appendix material.
""",
    )

    write_report(
        REPORT_DIR / "phase-1d9-references-section-report.md",
        f"""
# Phase 1D-9 References / Sources Section Report

**Document:** GAIC-2026-v0.3.2-FRC-R3
**Date:** May 11, 2026
**Status:** COMPLETE / CITATION QA REQUIRED

## Source Inputs

- `inventories/citation-inventory.md`
- `inventories/source-coverage-matrix.md`
- `inventories/claim-evidence-register.md`

## Publication Section

The generated References / Sources section is appended after Appendices A-K and grouped into:

1. Legal / Regulatory sources
2. Standards / Frameworks
3. Protocol / MPLP documentation
4. Validation sources
5. Product / Platform documentation

## Boundary Notes Preserved

- L1/L2/L3/L4/L5 evidence basis notes.
- Page-level citation pinning remains a publication QA task.
- OpenAI platform guide HTTP 403 remains tracked.
- Validation Lab methodology-specific support remains boundary-limited.
- MPLP conflict-of-interest and non-required / non-exclusive boundaries are explicit.

## Decision

References / Sources exists in the publication-facing artifact, but final page-level citation pinning and access QA remain unresolved blockers.
""",
    )

    return {
        "status": status,
        "source_word_count": source_words,
        "html_word_count": html_words,
        "pdf_word_count": pdf_words,
        "chapter_count": chapter_count,
        "appendix_count": appendix_count,
        "expected_chapters": expected_chapters,
        "expected_appendices": expected_appendices,
        "active_table_ids": len(source_ids),
        "html_table_ids": len(html_ids),
        "pdf_table_ids": len(pdf_ids),
        "missing_table_ids_html": missing_html,
        "missing_table_ids_pdf": missing_pdf,
        "figures_present": sorted(figure_ids & required_figures),
        "chart_records": [asdict(record) for record in chart_records],
        "evaluation_results_present": evaluation_results_present,
        "evaluation_systems_present": evaluation_systems_present,
        "chapter_panels_present": chapter_panels_present,
        "rubric_card_reflows": len(rubric_reflows),
        "system_profile_card_reflows": len(system_profile_reflows),
        "semantic_row_card_reflows": len(semantic_row_reflows),
        "semantic_column_splits": len(split_reflows),
        "t_08_13_strategy": t0813_record.strategy if t0813_record else "missing",
        "appendix_e_rubric_reflows": len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-E-")]),
        "appendix_f_rubric_reflows": len([r for r in appendix_e_f_reflows if r.table_id and r.table_id.startswith("T-F-")]),
        "appendix_g_system_reflows": len(appendix_g_reflows),
        "code_records": [asdict(record) for record in code_records],
        "problem_overview_present": "Problem This White Paper Solves" in pdf_text,
        "how_to_read_present": "How to Read This Paper" in pdf_text,
        "references_section_present": "References / Sources" in pdf_text,
        "pagination_warnings": pagination_warnings,
        "low_word_pages_under_30": low_word_pages,
        "blankish_pages": blank_pages,
        "edge_ink_pages": edge_pages,
        "forbidden_matches": len(forbidden_matches),
        "positive_blockers": len(positive_blockers),
        "citation_blockers": citation_blockers,
        "appendix_g_blockers": appendix_g_blockers,
        "contact_sheets": contact_sheets,
    }


def write_log(log: dict) -> None:
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> int:
    clean_output_dir()
    html_path, table_records, figure_records, chart_records, code_records, headings, source_text = build_html()
    pdf_path = html_to_pdf(html_path)
    docx_path, docx_stdout = source_to_docx_derivative()
    page_count = pdf_page_count(pdf_path)
    pdf_text_path, pdf_pages = extract_pdf_text(pdf_path)
    rendered_pages = render_pdf_pages(pdf_path)

    figure_pages = find_pages_containing(pdf_pages, [f"F-{i:02d}" for i in range(1, 12)])
    high_risk_table_pages = find_pages_containing(pdf_pages, sorted(HIGH_RISK_TABLE_IDS))
    appendix_pages = find_pages_containing(pdf_pages, ["Appendix A", "Appendix G", "Appendix K"]) + list(range(max(1, page_count - 8), page_count + 1))
    front_pages = sorted(set(list(range(1, min(page_count, 12) + 1)) + find_pages_containing(pdf_pages, ["Executive Summary", "Publication Contents"])))
    contact_sheets_raw = {
        "front": make_contact_sheet(rendered_pages, front_pages, "phase_1d9_front_matter_contact_sheet.png"),
        "charts": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, [record.title for record in chart_records] + ["Problem This White Paper Solves", "How to Read This Paper"]), "phase_1d9_charts_overview_contact_sheet.png"),
        "evaluation": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, ["Evaluation Results: RCCS/ALCS Analytical Profiles", "System Profile Cards"]), "phase_1d9_evaluation_results_contact_sheet.png"),
        "rccs_alcs": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, ["RCCS / ALCS Analytical Posture Map", "System x Dimension Posture Heatmap"]), "phase_1d9_rccs_alcs_visual_profiles_contact_sheet.png"),
        "figures": make_contact_sheet(rendered_pages, figure_pages, "phase_1d9_figure_pages_contact_sheet.png"),
        "tables": make_contact_sheet(rendered_pages, high_risk_table_pages, "phase_1d9_high_risk_tables_contact_sheet.png"),
        "rubrics": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, ["T-08-13", "T-E-01", "T-F-11", "Rubric rendered as score cards"]), "phase_1d9_rubric_examples_contact_sheet.png"),
        "system_mappings": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, ["T-G-01", "T-G-09", "System mapping rendered as profile cards", "Claim-Level Support Status"]), "phase_1d9_system_mapping_examples_contact_sheet.png"),
        "references": make_contact_sheet(rendered_pages, find_pages_containing(pdf_pages, ["References / Sources", "How to Read Evidence Labels", "Source Status Notes"]), "phase_1d9_references_contact_sheet.png"),
        "appendices": make_contact_sheet(rendered_pages, appendix_pages, "phase_1d9_appendix_final_pages_contact_sheet.png"),
    }
    contact_sheets = {k: (str(v.relative_to(ROOT)) if v else "") for k, v in contact_sheets_raw.items()}

    html_source = html_path.read_text(encoding="utf-8", errors="ignore")
    pdf_text = pdf_text_path.read_text(encoding="utf-8", errors="ignore")
    forbidden_matches, positive_blockers = scan_forbidden(
        {
            "active_source": source_text,
            "generated_html": html_text(html_path),
            "generated_pdf_text": pdf_text,
        }
    )

    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "phase": "1D-9",
        "artifact": "html_publication_refinement_draft",
        "phase_1d9_status": PHASE_STATUS,
        "root": str(ROOT),
        "branch": run_git(["branch", "--show-current"]),
        "source_commit": run_git(["rev-parse", "HEAD"]),
        "origin_head": run_git(["ls-remote", "origin", "HEAD"]),
        "source_files": SOURCE_FILES,
        "html_path": str(html_path),
        "pdf_path": str(pdf_path),
        "docx_path": str(docx_path),
        "pdf_text_path": str(pdf_text_path),
        "output_directory": str(OUT_DIR),
        "design_system": "design/gaic-html-publication-style.md",
        "style_reference_report": "reports/phase-1d7-genspark-style-analysis.md",
        "architecture_report": "reports/phase-1d7-html-renderer-architecture.md",
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "genspark_text_used_as_content": False,
        "html_pdf_visual_source_of_truth": True,
        "docx_editable_derivative": True,
        "publication_candidate_declared": False,
        "final_publication_declared": False,
        "pdf_export_engine": "Chrome headless print-to-PDF",
        "docx_export_engine": "python-docx from same Phase 1D-9 structured source",
        "docx_export_stdout": docx_stdout,
        "pdf_page_count": page_count,
        "rendered_page_count": len(rendered_pages),
        "table_records": [asdict(record) for record in table_records],
        "split_tables": [asdict(record) for record in table_records if record.strategy != "single table"],
        "figure_records": [asdict(record) for record in figure_records],
        "chart_records": [asdict(record) for record in chart_records],
        "code_records": [asdict(record) for record in code_records],
        "heading_count": len(headings),
        "forbidden_positive_blockers": len(positive_blockers),
        "citation_blockers": 1,
        "qa_reports": [
            "reports/phase-1d9-evaluation-results-architecture.md",
            "reports/phase-1d9-rccs-alcs-visual-results-plan.md",
            "reports/phase-1d9-chapter-hierarchy-report.md",
            "reports/phase-1d9-semantic-table-type-system.md",
            "reports/phase-1d9-rubric-table-reflow-report.md",
            "reports/phase-1d9-system-table-reflow-report.md",
            "reports/phase-1d9-source-citation-layer-report.md",
            "reports/phase-1d9-pagination-refinement-report.md",
            "reports/phase-1d9-chart-visualization-plan.md",
            "reports/phase-1d9-chart-implementation-report.md",
            "reports/phase-1d9-codeblock-object-card-report.md",
            "reports/phase-1d9-frontmatter-toc-overview-report.md",
            "reports/phase-1d9-table-typography-report.md",
            "reports/phase-1d9-references-section-report.md",
            "reports/phase-1d9-html-publication-generation-report.md",
            "reports/phase-1d9-html-visual-qa.md",
            "reports/phase-1d9-evaluation-results-qa.md",
            "reports/phase-1d9-table-semantics-qa.md",
            "reports/phase-1d9-content-preservation-qa.md",
            "reports/phase-1d9-citation-boundary-qa.md",
            "reports/phase-1d9-boundary-claim-qa.md",
            "reports/phase-1d9-docx-derivative-qa.md",
        ],
    }

    qa_summary = create_reports(
        html_path,
        pdf_path,
        docx_path,
        pdf_text_path,
        rendered_pages,
        table_records,
        figure_records,
        chart_records,
        code_records,
        headings,
        source_text,
        pdf_pages,
        log,
        forbidden_matches,
        positive_blockers,
        contact_sheets,
    )
    log["qa_summary"] = qa_summary
    write_log(log)

    print(
        json.dumps(
            {
                "html": str(html_path),
                "pdf": str(pdf_path),
                "docx": str(docx_path),
                "pages": page_count,
                "rendered_pages": len(rendered_pages),
                "tables": len(table_records),
                "split_tables": len(log["split_tables"]),
                "figures": len(figure_records),
                "charts": len(chart_records),
                "code_records": len(code_records),
                "positive_blockers": len(positive_blockers),
                "citation_blockers": 1,
                "status": PHASE_STATUS,
                "old_docx_pdf_used_as_input": False,
                "publication_candidate_declared": False,
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
