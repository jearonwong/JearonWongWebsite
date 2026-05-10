#!/usr/bin/env python3
"""
Phase 1D-3 public white paper generator for GAIC-2026 v0.3.2 FRC-R3.

This script creates a compressed public-facing white paper from the current
source_r3 corpus. It deliberately does not include full Appendices A-K.
Dense rubrics, inventories, source registers, and full mapping tables belong
in the separate technical evidence pack.
"""

from __future__ import annotations

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import generate_r3_phase1d_docx_pdf as base
import generate_r3_phase1d2_publication_draft as pub


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "phase_1d3" / "public_whitepaper"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
DOCX_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Public-White-Paper-Draft.docx"
PDF_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Public-White-Paper-Draft.pdf"
MD_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Public-White-Paper-Draft.md"
LOG_NAME = "public_whitepaper_generation_log.json"


def configure_base() -> None:
    base.OUT_DIR = OUT_DIR
    base.RENDER_DIR = RENDER_DIR
    base.DOCX_NAME = DOCX_NAME
    base.PDF_NAME = PDF_NAME
    base.ASSEMBLED_MD_NAME = MD_NAME
    base.LOG_NAME = LOG_NAME


def clean_text(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


PUBLIC_BRIEF_FILES = [
    ("sections/01-scope-methodology.md", "Scope, Methodology, and Boundary", 6200),
    ("sections/02-missing-layer.md", "The Missing Layer", 4200),
    ("sections/03-agentic-ai-breaks-model-compliance.md", "Why Agentic AI Breaks Model-Centric Compliance", 4200),
    ("sections/04-regulatory-standards-baseline.md", "Regulatory and Standards Baseline", 7600),
    ("sections/05-regulatory-language-to-engineering-objects.md", "Regulatory Language to Engineering Objects", 7200),
    ("sections/06-missing-regulatory-objects.md", "Missing Regulatory Objects", 13000),
    ("sections/07-rccs-regulatory-compliance-coverage.md", "RCCS Public Method Brief", 11000),
    ("sections/08-alcs-agentic-lifecycle-conformance.md", "ALCS Public Method Brief", 11500),
    ("sections/09-composite-scoring-method.md", "Composite Method Boundary", 8200),
    ("sections/10-enterprise-control-crosswalk.md", "Enterprise Control Translation", 9000),
    ("sections/11-comparative-field.md", "Comparative Field Boundary", 6800),
    ("sections/12-detailed-system-mappings.md", "System Mapping Synthesis", 9000),
    ("sections/13-mplp-deep-mapping.md", "MPLP Protocol Boundary", 7200),
    ("sections/14-evidence-validation-placeholder.md", "Evidence-Based Validation Pattern", 7600),
    ("sections/15-failure-scenarios-placeholder.md", "Failure Scenario Reasoning", 11000),
    ("sections/16-companion-paper-boundary.md", "Companion Paper Boundary", 6500),
    ("sections/17-adoption-roadmap-placeholder.md", "Adoption Roadmap Detail", 8200),
    ("sections/18-conclusion-placeholder.md", "Conclusion and Contribution", 6200),
]


def strip_tables_code_and_internal_noise(text: str) -> str:
    cleaned: list[str] = []
    in_code = False
    in_table = False
    skip_citation_block = False

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue

        is_table_line = stripped.startswith("|") and stripped.endswith("|")
        if is_table_line:
            in_table = True
            continue
        if in_table and not is_table_line:
            in_table = False

        lower = stripped.lower()
        if lower.startswith("citation placeholders"):
            skip_citation_block = True
            continue
        if skip_citation_block:
            if not stripped:
                skip_citation_block = False
            continue

        if not stripped:
            cleaned.append("")
            continue

        if stripped.startswith("<!--"):
            continue
        if "figure placeholder" in lower or stripped.startswith("[FIGURE"):
            continue
        if "Table T-" in stripped or "Table T-" in raw:
            if stripped.lower().startswith(("table ", "**table", "### table")):
                continue
            line = re.sub(r"\(?Table T-[A-Z0-9-]+[A-Z]?\)?", "the corresponding evidence-pack table", line)
            stripped = line.strip()
            lower = stripped.lower()
        if re.match(r"^\*?\*?Table\s+(T-)?[A-Z0-9-]+:", stripped, re.IGNORECASE):
            continue
        if re.match(r"^Table\s+(T-)?[A-Z0-9-]+:", stripped, re.IGNORECASE):
            continue
        if re.match(r"^\*?\*?Figure\s+(F-)?\d+", stripped, re.IGNORECASE):
            continue
        if lower.startswith("figure description") or "figure description:" in lower:
            continue
        if lower.startswith("model risk management table"):
            continue
        if lower.startswith("**status:**"):
            continue
        if "phase 1a-2 to be filled" in lower:
            continue
        if "phase 1a-2 to be generated" in lower:
            continue
        if lower.startswith("uses citation placeholders"):
            continue
        if stripped.startswith("**Appendix") and "Status:" in stripped:
            continue

        if stripped.startswith("#"):
            title = base.clean_inline_markdown(stripped.lstrip("#").strip())
            if "Table T-" in title or title.lower().startswith("table "):
                continue
            if title.lower().startswith("figure description"):
                continue
            if title.lower().startswith("model risk management table"):
                continue
            if title:
                cleaned.append(f"### {title}")
            continue

        cleaned.append(line)

    text = "\n".join(cleaned)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def limit_on_paragraph_boundary(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    paragraphs = re.split(r"\n\s*\n", text)
    out: list[str] = []
    total = 0
    for paragraph in paragraphs:
        proposed = total + len(paragraph) + 2
        if out and proposed > max_chars:
            break
        out.append(paragraph)
        total = proposed
    excerpt = "\n\n".join(out).strip()
    if len(excerpt) < max_chars * 0.55:
        excerpt = text[:max_chars].rsplit("\n", 1)[0].strip()
    return excerpt


def build_public_chapter_briefs() -> str:
    chunks = [
        "# 13. Public Technical Briefs",
        "",
        "This section expands the public argument without reintroducing the full technical appendix corpus. It draws from the active body chapters, removes dense tables, omits internal placeholder language, and preserves the paper's boundary discipline. The complete tables, rubrics, mappings, source ledgers, and QA reports remain in the separate technical evidence pack.",
    ]
    for rel, title, max_chars in PUBLIC_BRIEF_FILES:
        path = ROOT / rel
        if not path.exists():
            continue
        excerpt = limit_on_paragraph_boundary(strip_tables_code_and_internal_noise(path.read_text(encoding="utf-8")), max_chars)
        if not excerpt:
            continue
        chunks.extend(
            [
                "",
                f"## {title}",
                "",
                excerpt,
            ]
        )
    return "\n".join(chunks)


def write_public_markdown() -> Path:
    public_chapter_briefs = build_public_chapter_briefs()
    md = f"""# Global AI Compliance White Paper 2026

**From Model Governance to Agentic Lifecycle Conformance**

**Document ID:** GACWP-2026-v0.3.2-FRC-R3  
**Artifact:** Public White Paper Draft  
**Document Type:** Public White Paper Draft  
**Compiled:** {datetime.now().strftime('%B %d, %Y')}

---

## Document Control and Boundary Statement

This public white paper is a non-legal technical governance analysis. It is not legal advice, a certification, a regulatory approval, a procurement recommendation, a vendor ranking, or a final vendor assessment.

MPLP is discussed as one lifecycle protocol path. It is not required, exclusive, certified, regulator-approved, or legally sufficient by itself. Validation Lab is referenced only as a non-certifying evidence adjudication example.

RCCS and ALCS are analytical tools. They do not prove legal compliance.

## How to Read This Edition

This publication draft separates the public argument from the technical evidence substrate. The public white paper carries the narrative, core concepts, selected figures, methodology summary, and executive-level adoption guidance. The separate technical evidence pack preserves full appendices, rubrics, mapping tables, source registers, claim evidence register, inventories, and QA reports.

The public white paper intentionally does not include full Appendices A-K.

## Curated Table of Contents

1. Executive Summary
2. Core Thesis
3. Why Model Governance Is Not Enough
4. The Missing Layer
5. Missing Regulatory Objects
6. RCCS, ALCS, and Composite Scoring
7. Enterprise Control Implications
8. Comparative System Evidence
9. MPLP Boundary and Protocol Role
10. Evidence-Based Validation Pattern
11. Enterprise Failure Scenario Chain
12. Adoption Roadmap
13. Public Technical Briefs
14. Conclusion
15. Public Appendices

## Curated Figures

- Figure F-01: AI Agent Lifecycle Governance Stack
- Figure F-03: MRO Topology
- Figure F-04: RCCS/ALCS Dual Scoring Model
- Figure F-08: Evidence-Based Validation Pattern Flow
- Figure F-09: Enterprise Failure Scenario Chain
- Figure F-11: Stage 0-to-7 Adoption Roadmap

## Curated Major Tables

- Public Table 1: Evidence Levels
- Public Table 2: Missing Regulatory Object Clusters
- Public Table 3: RCCS/ALCS Method Summary
- Public Table 4: Representative System Evidence Posture
- Public Table 5: Adoption Roadmap Summary
- Public Appendix Tables A-D

---

# 1. Executive Summary

AI compliance is entering a structural transition. The compliance object is no longer limited to a model, a dataset, a prompt, or a single application output. In enterprise environments, AI systems increasingly plan work, call tools, retrieve memory, delegate tasks, collaborate with other agents, and produce operational outcomes.

This changes the compliance question. The question is no longer only whether a model produced a safe output. The question is whether an organized unit of agentic work was authorized, executed, evidenced, reviewed, accepted, disputed, remediated, and improved under clear responsibility boundaries.

This paper names the missing layer **AI Agent Lifecycle Governance**. It is the governance layer that defines how agentic work is authorized, executed, evidenced, reviewed, accepted, disputed, remediated, and reused across agents, humans, tools, projects, vendors, and organizational boundaries.

**Figure F-01: AI Agent Lifecycle Governance Stack**

`Intent -> Context -> Plan -> Authority -> Agent Work -> Tool Actions -> Evidence -> Review -> Accepted Outcome -> Dispute / Remediate`

The framework introduced here has four parts. First, it defines Missing Regulatory Objects, or MROs, that make regulatory abstractions testable in agentic workflows. Second, it introduces RCCS, an analytical score for regulatory compliance coverage. Third, it introduces ALCS, an analytical score for lifecycle conformance. Fourth, it separates public white paper findings from the technical evidence pack that supports detailed review.

## Key Takeaways

| Takeaway | Meaning |
|----------|---------|
| Model governance remains necessary | It still covers model risk, documentation, transparency, robustness, monitoring, and governance processes |
| Agentic workflows add lifecycle responsibility | Tool actions, delegation, handoffs, evidence chains, and accepted outcomes require additional objects |
| Logs are not enough | Logs may show execution, but they do not automatically prove authority, acceptance, privacy treatment, or remediation closure |
| RCCS and ALCS are analytical | They help compare governance capability, but they are not legal compliance scores |
| Public evidence is uneven | Official product documentation can support visible capability claims, but does not prove deployment practice |
| Technical details belong in the evidence pack | Full rubrics, source registers, mappings, and QA ledgers are preserved separately |

# 2. Core Thesis

**AI Agent / MAS compliance is not only model compliance. It is lifecycle responsibility compliance.**

That thesis does not claim that current AI regulation is absent. The claim is narrower and more technical. Existing regulation and governance frameworks increasingly require oversight, documentation, monitoring, accountability, transparency, privacy controls, and risk management. What they often do not specify is the machine-readable lifecycle object model needed to prove those properties inside dynamic agentic and multi-agent execution.

The lifecycle responsibility problem appears whenever an AI system moves from producing content to performing work. The moment an agent sends an email, approves a refund, calls an API, modifies a database, routes a case, changes access, triggers another agent, or produces an operational record, the governance object is no longer only a model output. It is the full lifecycle of delegated work.

## The Governance Shift

| Earlier Governance Focus | Agentic Governance Question |
|--------------------------|-----------------------------|
| Is the model documented? | Was the work unit authorized and evidenced? |
| Is the output safe? | Was the action within delegated authority? |
| Is there human review? | Which human role held responsibility at each lifecycle state? |
| Are logs available? | Can evidence reconstruct authority, action, review, acceptance, and remediation? |
| Is a vendor documented? | Does conformance survive model, runtime, or tool substitution? |

This shift is why the public white paper is shorter than the technical evidence pack. The public paper explains the governance architecture. The evidence pack preserves the full proof substrate.

# 3. Why Model Governance Is Not Enough

Model governance answers important questions about the model. It does not, by itself, answer the lifecycle responsibility question. Agentic systems are not merely models; they are work systems. They combine model behavior with orchestration, tools, memory, roles, approvals, logs, evidence policies, privacy constraints, and remediation workflows.

The problem is not that model governance is wrong. The problem is that model governance is incomplete once the system delegates consequential work.

## The Layer Gap

| Layer | What It Governs | Why It Is Insufficient Alone |
|-------|-----------------|------------------------------|
| Model governance | Model quality, documentation, evaluation, monitoring | Does not prove responsibility for delegated work |
| Agent orchestration | Task flow, tools, graph state, handoffs | Does not assign legal or organizational responsibility |
| Observability | Traces, logs, metrics, debugging | Does not prove authorization, acceptance, privacy, or remediation |
| Human review | Review or approval moment | Does not map human role to MAS responsibility |
| Lifecycle governance | Authority, responsibility, evidence, acceptance, dispute, remediation | Requires explicit lifecycle objects and evidence |

Agentic compliance becomes fragile when enterprises assume that orchestration traces, model cards, or human-in-the-loop labels are enough. They are useful artifacts, but they are not a complete responsibility model.

# 4. The Missing Layer

AI Agent Lifecycle Governance is the missing layer between model compliance and enterprise agent deployment. It defines the lifecycle control surface for agentic work.

This layer translates regulatory abstractions into engineering objects:

- Human oversight becomes human-role-to-MAS responsibility mapping.
- Logging becomes partitioned evidence chain.
- Accountability becomes tool-action liability and responsibility transfer.
- Transparency becomes evidence-linked review and accepted outcome.
- Monitoring becomes authority drift, incident closure, and continuous improvement.
- Privacy becomes lifecycle data-flow mapping, minimization, and selective disclosure.

This is not bureaucracy for its own sake. It is the recognition that the system boundary has changed. The work unit is now distributed across agents, tools, memory, humans, projects, vendors, and evidence boundaries.

# 5. Missing Regulatory Objects

The paper defines sixteen Missing Regulatory Objects. They are not proposed statutes. They are engineering and assurance objects that allow regulatory concepts to become testable in agentic workflows.

**Figure F-03: MRO Topology**

`Responsibility -> Authority -> Evidence -> Privacy -> Substitution -> Closure`

## Public MRO Summary

| Cluster | Objects | Public Meaning |
|---------|---------|----------------|
| Responsibility | MRO-01, MRO-03, MRO-04, MRO-06 | Who owns intent, agent role, handoff, and accepted outcome |
| Authority | MRO-02, MRO-05, MRO-07 | What the agent may do, when authority drifts, and who owns tool consequences |
| Evidence | MRO-08, MRO-11, MRO-12 | How lifecycle evidence is partitioned, validated, minimized, and selectively disclosed |
| Privacy | MRO-10, MRO-13, MRO-14 | How personal data, data subject rights, and processor chains are preserved across lifecycle states |
| Reuse and substitution | MRO-09, MRO-15 | Whether conformance survives reuse, runtime change, vendor change, or model substitution |
| Closure | MRO-16 | Whether incidents, disputes, remediation, and accepted closure are evidenced |

The full object cards, fields, controls, audit questions, and implementation implications are preserved in the technical evidence pack.

# 6. RCCS, ALCS, and Composite Scoring

The paper introduces two analytical scoring layers:

1. **RCCS - Regulatory Compliance Coverage Score** measures whether a system provides primitives aligned with regulatory and governance requirements such as risk management, data governance, documentation, record-keeping, transparency, human oversight, security, accountability, contestability, and monitoring.
2. **ALCS - Agentic Lifecycle Conformance Score** measures whether a system defines lifecycle responsibility objects for agentic and multi-agent execution.

**Figure F-04: RCCS/ALCS Dual Scoring Model**

`RCCS regulatory coverage + ALCS lifecycle conformance -> evidence-adjusted analytical composite`

## Canonical Method

| Method Element | Public Formula / Rule |
|----------------|-----------------------|
| Evidence levels | L1 legal/standards; L2 official docs; L3 independent; L4 vendor statement; L5 analytical inference |
| Evidence multipliers | L1=1.00, L2=0.85, L3=0.75, L4=0.55, L5=0.35 |
| Dimension score | Dimension Score = (Raw Score / 5) x 100 x Evidence Confidence Multiplier |
| Composite score | Composite Score = 0.40 x RCCS + 0.60 x ALCS |
| Boundary | Analytical only; not legal compliance proof, certification, procurement recommendation, or regulatory approval |

The technical evidence pack preserves full RCCS and ALCS rubrics.

# 7. Enterprise Control Implications

Enterprise teams evaluate controls, procurement risk, board reporting, audit readiness, vendor substitution, cybersecurity, incident response, and monitoring. Lifecycle governance must therefore translate into enterprise control language.

The core enterprise implication is simple: agentic AI governance needs evidence that follows the work. If an agentic workflow cannot prove who authorized work, who executed it, what tool action occurred, what evidence was retained, who reviewed the result, who accepted the outcome, and how remediation closed, the enterprise has an accountability gap.

## Control Translation

| Enterprise Function | Lifecycle Question |
|---------------------|-------------------|
| Internal audit | Can the organization reconstruct authority, action, review, and accepted outcome? |
| Privacy | Can evidence be retained without over-disclosing personal data? |
| Security | Can delegated authority and tool-action boundaries survive attack or misuse? |
| Procurement | Can vendors expose the evidence and lifecycle primitives buyers need? |
| Board reporting | Can agentic risk be explained in terms of authority, evidence, accepted outcomes, and open remediation? |
| Incident response | Can disputes close with evidence, responsibility, correction, and accepted closure? |

# 8. Comparative System Evidence

The paper evaluates representative systems and frameworks as responsibility-semantics surfaces, not as product rankings. The public white paper summarizes the evidence posture. Full matrices and system-by-system assessments are preserved in the evidence pack.

## Representative Systems

| System / Category | Public Role in the Analysis | Evidence Posture |
|-------------------|-----------------------------|------------------|
| MPLP | Lifecycle protocol path | Identity and protocol entry points L2-supported; MRO/ALCS mapping remains L2 + L5 |
| IBM watsonx.governance | Enterprise AI governance platform | Official docs support governance primitives; lifecycle-object depth remains claim-bound |
| Microsoft Azure AI Foundry | Enterprise AI platform and control plane | Official docs support platform governance primitives; lifecycle responsibility objects require integration |
| AWS Bedrock / Guardrails / AgentCore | Managed agent runtime and guardrail infrastructure | Official docs support guardrail and runtime primitives; lifecycle object mapping remains bounded |
| Google Vertex AI / ADK / Model Armor | Managed AI and agent safety infrastructure | Official docs support platform and safety primitives; lifecycle conformance requires implementer design |
| LangGraph / LangSmith | Stateful orchestration and observability tooling | Official docs support orchestration and traceability primitives; governance semantics remain implementer-dependent |
| OpenAI Agents SDK | Developer agent SDK | Accessible Agents SDK docs support SDK-surface claims; platform guide 403 remains tracked |
| CrewAI | Multi-agent orchestration framework | Official repo/docs support orchestration primitives; lifecycle responsibility semantics require integration |

This public summary is qualitative and provisional. It is not a final vendor assessment and must not be used as a procurement recommendation.

# 9. MPLP Boundary and Protocol Role

MPLP is discussed because it is a lifecycle protocol path for expressing agentic responsibility objects. The author created MPLP, so the conflict of interest is disclosed. The analysis therefore uses boundary controls.

MPLP is not required for compliance. It is not the only validation path. It is not a certification program, regulatory authority, conformity assessment body, legal compliance proof, or procurement recommendation.

The public claim is narrower: agentic AI compliance requires lifecycle responsibility objects, and MPLP is one coherent way to express such objects. Other systems may implement comparable objects differently.

# 10. Evidence-Based Validation Pattern

Evidence-based validation separates evidence generation from evidence adjudication. Organizations generate lifecycle responsibility evidence. Validation providers or internal reviewers adjudicate that evidence against predefined rulesets without requiring raw operational exposure.

**Figure F-08: Evidence-Based Validation Pattern Flow**

`Evidence Pack -> Hash Manifest -> Ruleset -> Evidence Pointers -> Verdict Hash -> Replay / Recheck`

Validation Lab is referenced only as a non-certifying evidence adjudication example. It does not certify compliance, prove legal compliance, guarantee regulatory approval, or act as a certification authority.

## Validation Boundary

| Pattern Can Support | Pattern Does Not Prove |
|---------------------|------------------------|
| Lifecycle object existence | Legal compliance |
| Evidence integrity | Regulatory approval |
| Selective disclosure | Business logic correctness |
| Replay and recheck | Model accuracy or operational effectiveness |
| Privacy-preserving review | Certification or conformity assessment |

# 11. Enterprise Failure Scenario Chain

Enterprise failures often occur when a governance gap moves through a lifecycle chain: a trigger happens, an object is missing, evidence is incomplete, responsibility becomes unclear, and remediation cannot close cleanly.

**Figure F-09: Enterprise Failure Scenario Chain**

`Trigger -> Missing Object -> Evidence Gap -> Enterprise Consequence -> Remediation Closure`

## Failure Categories

| Failure Category | Typical Missing Object | Enterprise Risk |
|------------------|------------------------|-----------------|
| Authority and boundary failure | Delegated Authority Boundary; Authority Drift | Unauthorized action and liability dispute |
| Evidence and traceability failure | MAS Evidence Partitioning; Evidence Minimization | Audit failure and dispute resolution failure |
| Responsibility and acceptance failure | Human Role Mapping; Accepted Outcome | No accountable owner for outcome |
| Integration and substitution failure | Cross-Project Reuse; Vendor/Runtime Substitution | Compliance drift across reuse or vendor change |

The full scenario set is preserved in the technical evidence pack.

# 12. Adoption Roadmap

Enterprises should treat lifecycle governance as an adoption path, not as a one-time checklist.

**Figure F-11: Stage 0-to-7 Adoption Roadmap**

`0 Inventory -> 1 Authority -> 2 Responsibility -> 3 Evidence -> 4 Privacy -> 5 Validation -> 6 Remediation -> 7 Scale`

## Roadmap Summary

| Stage | Goal | Output |
|-------|------|--------|
| 0 Inventory | Identify agentic behavior | Agentic behavior inventory |
| 1 Authority | Define delegated authority boundaries | Authority records and drift monitoring |
| 2 Responsibility | Map human roles to agentic responsibility | Role-responsibility map |
| 3 Evidence | Convert logs into evidence chains | Partitioned evidence chain |
| 4 Privacy | Map privacy and subject rights across lifecycle | Lifecycle privacy map |
| 5 Validation | Prepare evidence-based validation | Validation packs and rulesets |
| 6 Remediation | Close incidents and disputes | Remediation closure records |
| 7 Scale | Integrate with enterprise controls | Lifecycle governance operating model |

{public_chapter_briefs}

# 14. Conclusion

The central contribution of this paper is to name and structure the missing layer between model governance and enterprise agent deployment. That layer is AI Agent Lifecycle Governance.

The public white paper makes the argument. The technical evidence pack preserves the detail. Together, they support a publication model that is readable for enterprise leaders while retaining audit-grade traceability for technical reviewers.

The conclusion is deliberately bounded:

- This paper does not provide legal advice.
- It does not certify systems.
- It does not recommend vendors.
- It does not rank products.
- It does not claim that MPLP is required.
- It does not claim that Validation Lab certifies compliance.
- It does claim that agentic AI governance requires lifecycle responsibility objects.

# Public Appendix A: Methodology Summary

The methodology translates regulatory language into engineering objects, maps those objects to system primitives, and applies evidence-level discipline to prevent overclaiming.

The public white paper includes methodology summaries only. Full rubrics and scoring worksheets are in the technical evidence pack.

# Public Appendix B: Source and Evidence Note

Evidence levels are preserved as L1 through L5. L1 legal and standards sources establish baseline obligations but do not prove implementation. L2 official documentation supports public capability mapping but does not prove deployment practice. L5 author inference is marked as analytical interpretation.

The OpenAI platform guide HTTP 403 access note remains tracked. OpenAI SDK-surface claims are narrowed to accessible Agents SDK documentation unless final access succeeds.

# Public Appendix C: Non-Claim Boundary

The public white paper must not be read as legal advice, certification, regulatory approval, conformity assessment, procurement recommendation, market endorsement, vendor ranking, or final vendor score.

Appendix G remains revalidated provisional. It is qualitative, non-ranking, non-score-based, and non-procurement-oriented.

# Public Appendix D: Technical Evidence Pack Access Guide

The technical evidence pack contains the full materials omitted from this public edition:

- Appendices A-K
- full RCCS and ALCS rubrics
- full system mappings and Appendix G assessments
- source audit register
- claim evidence register
- table, figure, citation, and source coverage inventories
- QA and phase reports

This split is intentional. The public white paper is for narrative understanding. The technical evidence pack is for audit traceability and detailed review.
"""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / MD_NAME
    path.write_text(clean_text(md) + "\n", encoding="utf-8")
    return path


def set_public_defaults(doc: Document) -> None:
    pub.set_publication_defaults(doc)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("GAIC-2026 v0.3.2-FRC-R3 | Public White Paper Draft")
        run.font.name = "Arial"
        run.font.size = Pt(7.2)
        run.font.color.rgb = pub.ANODIZED_SLATE


def add_public_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(92)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Global AI Compliance White Paper 2026")
    run.font.name = "Arial"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = pub.MONOLITH_CHARCOAL

    subtitle = doc.add_paragraph()
    srun = subtitle.add_run("Public White Paper Draft")
    srun.font.name = "Arial"
    srun.font.size = Pt(15)
    srun.font.color.rgb = pub.ANODIZED_SLATE

    deck = doc.add_paragraph()
    deck.paragraph_format.space_before = Pt(10)
    drun = deck.add_run("From Model Governance to Agentic Lifecycle Conformance")
    drun.font.name = "Arial"
    drun.font.size = Pt(12)
    drun.font.color.rgb = pub.PROTOCOL_BLUE

    pub.add_publication_figure_panel(doc, "Figure F-01: AI Agent Lifecycle Governance Stack")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(22)
    mrun = meta.add_run("GACWP-2026-v0.3.2-FRC-R3 | Public white paper plus separate technical evidence pack")
    mrun.font.name = "Arial"
    mrun.font.size = Pt(8.8)
    mrun.font.color.rgb = pub.ANODIZED_SLATE

    boundary = doc.add_paragraph()
    brun = boundary.add_run("Non-legal technical governance analysis. Not legal advice, certification, regulatory approval, procurement recommendation, or final vendor assessment.")
    brun.font.name = "Arial"
    brun.font.size = Pt(8.4)
    brun.font.color.rgb = pub.BOUNDARY_AMBER
    doc.add_page_break()


def build_docx(md_path: Path) -> tuple[Path, list[base.TableRecord], str]:
    doc = Document()
    set_public_defaults(doc)
    add_public_cover(doc)
    records: list[base.TableRecord] = []
    text = pub.add_file_to_publication_doc(doc, md_path, records)
    docx_path = OUT_DIR / DOCX_NAME
    doc.save(docx_path)
    return docx_path, records, text


def write_log(docx_path: Path, pdf_path: Path, rendered_pages: list[Path], records: list[base.TableRecord], source_text: str) -> dict:
    page_count = base.pdf_page_count(pdf_path)
    if page_count is None:
        length_status = "UNKNOWN"
    elif page_count > 160:
        length_status = "QA REQUIRED / TOO LONG"
    elif 80 <= page_count <= 140:
        length_status = "PASS TARGET"
    elif page_count < 80:
        length_status = "QA REQUIRED / BELOW TARGET"
    else:
        length_status = "PASS HARD MAX / ABOVE TARGET BAND"

    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "phase": "1D-3",
        "artifact": "public_whitepaper",
        "root": str(ROOT),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "assembled_markdown_path": str(OUT_DIR / MD_NAME),
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "full_appendices_a_k_included": False,
        "technical_evidence_pack_required": True,
        "pdf_page_count": page_count,
        "rendered_page_count": len(rendered_pages),
        "table_count": len(records),
        "table_records": [base.asdict(record) for record in records],
        "length_status": length_status,
        "publication_candidate_declared": False,
        "source_text_character_count": len(source_text),
    }
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    return log


def main() -> int:
    configure_base()
    md_path = write_public_markdown()
    docx_path, records, source_text = build_docx(md_path)
    pdf_path = base.convert_docx_to_pdf(docx_path)
    rendered_pages = base.render_pdf_pages(pdf_path)
    log = write_log(docx_path, pdf_path, rendered_pages, records, source_text)
    print(json.dumps({
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "pages": log["pdf_page_count"],
        "length_status": log["length_status"],
        "tables": len(records),
        "rendered_pages": len(rendered_pages),
        "log": str(OUT_DIR / LOG_NAME),
    }, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
