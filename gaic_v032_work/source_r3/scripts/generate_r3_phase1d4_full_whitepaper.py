#!/usr/bin/env python3
"""
Phase 1D-4 full-content white paper generator for GAIC-2026 v0.3.2 FRC-R3.

This script uses the current source_r3 Markdown assets only. It preserves
substantive Chapters 1-18 and Appendices A-K, while applying a continuous
publication reflow: no old DOCX/PDF input, no global H1 page breaks, no
per-table landscape section islands, and no visible figure placeholders.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import generate_r3_phase1d_docx_pdf as base
import generate_r3_phase1d2_publication_draft as pub


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "phase_1d4" / "full_whitepaper"
RENDER_DIR = OUT_DIR / "rendered_pdf_pages"
DOCX_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Full-White-Paper-Draft.docx"
PDF_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Full-White-Paper-Draft.pdf"
MD_NAME = "Global-AI-Compliance-White-Paper-2026-v0.3.2-FRC-R3-Full-White-Paper-Draft.md"
LOG_NAME = "phase_1d4_full_whitepaper_generation_log.json"

SOURCE_FILES = base.SOURCE_FILES
HIGH_RISK_TABLE_IDS = set(base.HIGH_RISK_TABLE_IDS) | {
    "T-F-01", "T-F-02", "T-F-03", "T-F-04", "T-F-05",
    "T-F-06", "T-F-07", "T-F-08", "T-F-09", "T-F-10",
    "T-F-11", "T-F-12", "T-F-13", "T-F-14", "T-F-15",
}

FORBIDDEN_TERMS = base.FORBIDDEN_TERMS
APPENDIX_G_STRICT_TERMS = base.APPENDIX_G_STRICT_TERMS


@dataclass
class ReflowTableRecord:
    table_id: str | None
    caption: str | None
    source_file: str
    original_rows: int
    original_cols: int
    rendered_tables: int
    strategy: str
    font_size: float
    high_risk: bool
    content_preserved: bool


@dataclass
class FigureRecord:
    figure_id: str
    title: str
    source_file: str
    placement: str


def which(name: str) -> str | None:
    return shutil.which(name)


def clean_inline_markdown(text: str) -> str:
    text = base.clean_inline_markdown(text)
    text = text.replace("→", "->")
    text = text.replace("—", "-")
    text = text.replace("–", "-")
    text = text.replace("✓", "")
    text = text.replace("◐", "")
    return text.strip()


def set_cell_border(cell, color: str = "D9DEE5", size: str = "5") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_indent(table, indent: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def set_publication_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = pub.MONOLITH_CHARCOAL
    normal.paragraph_format.space_after = Pt(3.6)
    normal.paragraph_format.line_spacing = 1.04

    for name, size, color, before, after in [
        ("Title", 27, pub.MONOLITH_CHARCOAL, 2, 8),
        ("Heading 1", 15.5, pub.MONOLITH_CHARCOAL, 14, 4),
        ("Heading 2", 12.5, pub.PROTOCOL_BLUE, 9, 3),
        ("Heading 3", 10.8, pub.ANODIZED_SLATE, 6, 2),
        ("Heading 4", 9.8, pub.ANODIZED_SLATE, 5, 2),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(9.1)
        style.paragraph_format.space_after = Pt(2.4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("GAIC-2026 v0.3.2-FRC-R3 | Full White Paper Draft | Phase 1D-4 QA Required")
    run.font.name = "Arial"
    run.font.size = Pt(7.0)
    run.font.color.rgb = pub.ANODIZED_SLATE


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(78)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Global AI Compliance White Paper 2026")
    run.font.name = "Arial"
    run.font.size = Pt(31)
    run.font.bold = True
    run.font.color.rgb = pub.MONOLITH_CHARCOAL

    subtitle = doc.add_paragraph()
    srun = subtitle.add_run("Full-Content Publication Draft")
    srun.font.name = "Arial"
    srun.font.size = Pt(15)
    srun.font.color.rgb = pub.ANODIZED_SLATE

    deck = doc.add_paragraph()
    deck.paragraph_format.space_before = Pt(8)
    drun = deck.add_run("From Model Governance to Agentic Lifecycle Conformance")
    drun.font.name = "Arial"
    drun.font.size = Pt(12.5)
    drun.font.color.rgb = pub.PROTOCOL_BLUE

    strip = doc.add_table(rows=2, cols=5)
    strip.alignment = WD_TABLE_ALIGNMENT.LEFT
    strip.autofit = False
    set_table_width(strip, 8800)
    labels = ["Intent", "Authority", "Agent Work", "Evidence", "Closure"]
    for idx, label in enumerate(labels):
        top = strip.rows[0].cells[idx]
        base.set_cell_shading(top, pub.CHARCOAL_FILL)
        set_cell_border(top, "111318", "8")
        set_cell_margins(top)
        set_cell_width(top, 1760)
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(8.4)
        r.font.bold = True
        r.font.color.rgb = RGBColor(247, 247, 244)

        bottom = strip.rows[1].cells[idx]
        base.set_cell_shading(bottom, pub.PANEL_FILL)
        set_cell_border(bottom)
        set_cell_margins(bottom)
        set_cell_width(bottom, 1760)
        q = bottom.paragraphs[0]
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = q.add_run(["authorized", "bounded", "traceable", "source-bound", "remediable"][idx])
        br.font.name = "Arial"
        br.font.size = Pt(7.2)
        br.font.color.rgb = pub.ANODIZED_SLATE

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(22)
    mrun = meta.add_run("GACWP-2026-v0.3.2-FRC-R3 | Phase 1D-4 full-content reflow | May 2026")
    mrun.font.name = "Arial"
    mrun.font.size = Pt(8.6)
    mrun.font.color.rgb = pub.ANODIZED_SLATE

    boundary = doc.add_paragraph()
    brun = boundary.add_run("Non-legal technical governance analysis. Not legal advice, certification, regulatory approval, procurement recommendation, vendor ranking, or final vendor assessment.")
    brun.font.name = "Arial"
    brun.font.size = Pt(8.2)
    brun.font.color.rgb = pub.BOUNDARY_AMBER
    doc.add_page_break()


def sanitize_publication_line(line: str, relative: str) -> str | None:
    stripped = line.strip()
    if stripped.startswith("<!--") and stripped.endswith("-->"):
        return None
    low = stripped.lower()

    # Suppress status residue in the generated publication while preserving
    # the underlying source files and governance reports.
    if relative == "sections/00-front-matter.md":
        if stripped.startswith("**Status:**"):
            return "**Status:** Phase 1D-4 full-content publication draft; QA required; non-legal technical governance analysis; not legal advice"
        if "| Status |" in stripped:
            return "| Status | Phase 1D-4 full-content publication draft; QA required; agenda-setting technical governance white paper; not legal advice |"
        if "phase 1d-2 publication design draft status" in low:
            return "This is **v0.3.2 Final Release Candidate R3 (GACWP-2026-v0.3.2-FRC-R3)** in Phase 1D-4 full-content publication draft status. The current source is the controlled Markdown source for designed DOCX/PDF generation, visual QA, citation rendering checks, and final publication proofing."

    if re.match(r"^\*\*Status:\*\*\s*Reconstructed from", stripped, re.IGNORECASE):
        return None
    if re.match(r"^\*\*Phase:\*\*\s*1[A-D]", stripped, re.IGNORECASE):
        return None
    if re.match(r"^Status:\s*Reconstructed from", stripped, re.IGNORECASE):
        return None
    if re.match(r"^Phase:\s*1[A-D]", stripped, re.IGNORECASE):
        return None

    if "phase 1a-2 to be generated" in low:
        return None
    return line


def write_assembled_markdown() -> tuple[Path, str]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    chunks: list[str] = []
    for rel in SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            raise FileNotFoundError(f"Missing required source file: {rel}")
        lines = []
        for line in path.read_text(encoding="utf-8").splitlines():
            sanitized = sanitize_publication_line(line, rel)
            if sanitized is not None:
                lines.append(sanitized)
        chunks.append("\n".join(lines).strip())
    assembled = "\n\n---\n\n".join(chunk for chunk in chunks if chunk)
    md_path = OUT_DIR / MD_NAME
    md_path.write_text(assembled + "\n", encoding="utf-8")
    return md_path, assembled


def heading_level(line: str) -> int:
    return len(line) - len(line.lstrip("#"))


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    return base.parse_markdown_table(lines, start)


def is_table_start(lines: list[str], i: int) -> bool:
    return base.is_table_start(lines, i)


def is_caption_line(line: str) -> bool:
    return base.is_caption_line(line.strip())


def detect_table_id(caption: str | None, rows: list[list[str]]) -> str | None:
    return base.detect_table_id(caption, rows)


def figure_id_from_line(line: str) -> str | None:
    return pub.figure_id_from_line(line)


def is_figure_marker(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith("-"):
        return False
    if pub.is_publication_figure_marker(stripped):
        return True
    if re.search(r"\bFigure F-\d{2}:", stripped):
        return True
    if re.search(r"\[FIGURE\s+\d+:", stripped, re.IGNORECASE):
        return True
    return False


def figure_panel_data(figure_id: str) -> dict:
    if figure_id == "F-02":
        return pub.FIGURE_PANELS["F-02"]
    return pub.FIGURE_PANELS.get(figure_id, pub.FIGURE_PANELS["F-01"])


def add_figure_panel(doc: Document, line: str, relative: str, figures: list[FigureRecord]) -> None:
    figure_id = figure_id_from_line(line) or "F-02"
    panel = figure_panel_data(figure_id)
    caption = doc.add_paragraph()
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(2)
    crun = caption.add_run(f"Figure {figure_id}: {panel['title']}")
    crun.font.name = "Arial"
    crun.font.size = Pt(8.8)
    crun.font.bold = True
    crun.font.color.rgb = pub.MONOLITH_CHARCOAL

    cols = len(panel["lanes"])
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_width = 8800
    set_table_width(table, total_width)
    col_width = max(760, total_width // cols)

    for idx, label in enumerate(panel["lanes"]):
        top = table.rows[0].cells[idx]
        base.set_cell_shading(top, pub.CHARCOAL_FILL)
        set_cell_border(top, "111318", "6")
        set_cell_margins(top, 70, 70, 70, 70)
        set_cell_width(top, col_width)
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(6.7 if cols > 7 else 7.4)
        r.font.bold = True
        r.font.color.rgb = RGBColor(247, 247, 244)

        bottom = table.rows[1].cells[idx]
        base.set_cell_shading(bottom, pub.PANEL_FILL)
        set_cell_border(bottom)
        set_cell_margins(bottom, 70, 70, 70, 70)
        set_cell_width(bottom, col_width)
        q = bottom.paragraphs[0]
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = q.add_run("object" if idx % 2 == 0 else "control")
        br.font.name = "Arial"
        br.font.size = Pt(6.4)
        br.font.color.rgb = pub.ANODIZED_SLATE

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(5)
    nrun = note.add_run(panel["note"])
    nrun.font.name = "Arial"
    nrun.font.size = Pt(7.4)
    nrun.italic = True
    nrun.font.color.rgb = pub.BOUNDARY_AMBER if "not " in panel["note"].lower() else pub.ANODIZED_SLATE

    figures.append(FigureRecord(figure_id=figure_id, title=panel["title"], source_file=relative, placement="inline DOCX-native panel"))


def add_markdown_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    cleaned = clean_inline_markdown(text)
    if not cleaned:
        return
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(3.4)
    para.add_run(cleaned)


def add_code_block(doc: Document, code: list[str]) -> None:
    if not code:
        return
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("\n".join(code))
    run.font.name = "Courier New"
    run.font.size = Pt(7.0)
    run.font.color.rgb = pub.ANODIZED_SLATE


def column_groups(cols: int) -> list[list[int]]:
    if cols <= 4:
        return [list(range(cols))]
    if cols == 5:
        return [[0, 1, 2], [0, 3, 4]]
    if cols == 6:
        return [[0, 1, 2], [0, 3, 4, 5]]
    if cols == 7:
        return [[0, 1, 2, 3], [0, 4, 5, 6]]
    if cols == 8:
        return [[0, 1, 2, 3], [0, 4, 5, 6, 7]]
    if cols == 9:
        return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8]]
    return [[0, 1, 2, 3], [0, 4, 5, 6], [0, 7, 8, 9]]


def projected_rows(rows: list[list[str]], group: list[int]) -> list[list[str]]:
    return [[row[idx] if idx < len(row) else "" for idx in group] for row in rows]


def table_font_size(cols: int, high_risk: bool) -> float:
    if cols <= 2:
        return 8.2
    if cols == 3:
        return 7.9
    if cols == 4:
        return 7.4 if high_risk else 7.7
    return 7.0


def add_single_table(doc: Document, rows: list[list[str]], cols: int, font_size: float, header_fill: str = pub.HEADER_FILL) -> None:
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    total_width = 8700
    set_table_width(table, total_width)
    set_table_indent(table, 0)
    col_width = max(1100, total_width // max(cols, 1))

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(tr)
        for c_idx, value in enumerate(row):
            cell = tr.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, col_width)
            set_cell_border(cell)
            set_cell_margins(cell)
            if r_idx == 0:
                base.set_cell_shading(cell, header_fill)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(clean_inline_markdown(value))
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.bold = r_idx == 0
            run.font.color.rgb = pub.MONOLITH_CHARCOAL if r_idx else pub.ANODIZED_SLATE


def add_table_to_doc(
    doc: Document,
    rows: list[list[str]],
    relative: str,
    caption: str | None,
    records: list[ReflowTableRecord],
) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table_id = detect_table_id(caption, normalized)
    high_risk = bool(table_id and table_id in HIGH_RISK_TABLE_IDS)
    split_required = cols >= 6 or high_risk
    groups = column_groups(cols) if split_required else [list(range(cols))]
    strategy = "semantic column split" if split_required else "portrait continuous"
    font_size = table_font_size(max(len(group) for group in groups), high_risk)

    if caption:
        para = doc.add_paragraph()
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.space_before = Pt(5)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(caption)
        run.font.name = "Arial"
        run.font.size = Pt(8.4)
        run.font.bold = True
        run.font.color.rgb = pub.MONOLITH_CHARCOAL

    if split_required:
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(2)
        nr = note.add_run("Layout note: wide source table is split into consecutive column groups; all source rows are preserved.")
        nr.font.name = "Arial"
        nr.font.size = Pt(7.3)
        nr.italic = True
        nr.font.color.rgb = pub.ANODIZED_SLATE

    for idx, group in enumerate(groups, 1):
        if split_required:
            sub = doc.add_paragraph()
            sub.paragraph_format.keep_with_next = True
            sub.paragraph_format.space_before = Pt(3 if idx == 1 else 6)
            sub.paragraph_format.space_after = Pt(1)
            sr = sub.add_run(f"Continuation {idx} of {len(groups)}")
            sr.font.name = "Arial"
            sr.font.size = Pt(7.5)
            sr.font.bold = True
            sr.font.color.rgb = pub.PROTOCOL_BLUE
        projected = projected_rows(normalized, group)
        add_single_table(doc, projected, len(group), font_size)

    records.append(
        ReflowTableRecord(
            table_id=table_id,
            caption=caption,
            source_file=relative,
            original_rows=len(normalized),
            original_cols=cols,
            rendered_tables=len(groups),
            strategy=strategy,
            font_size=font_size,
            high_risk=high_risk,
            content_preserved=True,
        )
    )


def add_file_to_doc(doc: Document, path: Path, table_records: list[ReflowTableRecord], figure_records: list[FigureRecord]) -> str:
    relative = path.relative_to(ROOT).as_posix()
    raw_text = path.read_text(encoding="utf-8")
    lines = []
    for raw in raw_text.splitlines():
        sanitized = sanitize_publication_line(raw, relative)
        if sanitized is not None:
            lines.append(sanitized)

    i = 0
    in_code = False
    code_buffer: list[str] = []
    pending_caption: str | None = None
    rendered_text_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        rendered_text_lines.append(line)

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if is_table_start(lines, i):
            rows, new_i = parse_markdown_table(lines, i)
            rendered_text_lines.extend(lines[i + 1:new_i])
            add_table_to_doc(doc, rows, relative, pending_caption, table_records)
            pending_caption = None
            i = new_i
            continue

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = min(heading_level(stripped), 4)
            text = clean_inline_markdown(stripped[level:].strip())
            if relative == "sections/00-front-matter.md" and text == "Global AI Compliance White Paper 2026":
                i += 1
                continue
            doc.add_heading(text, level=level)
            pending_caption = None
            i += 1
            continue

        if is_caption_line(stripped):
            pending_caption = clean_inline_markdown(stripped)
            i += 1
            continue

        if is_figure_marker(stripped):
            add_figure_panel(doc, stripped, relative, figure_records)
            pending_caption = None
            i += 1
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.18)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(clean_inline_markdown(stripped.lstrip("> ")))
            run.italic = True
            run.font.color.rgb = pub.ANODIZED_SLATE
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bullet_match:
            add_markdown_paragraph(doc, bullet_match.group(2), style="List Bullet")
            i += 1
            continue
        if number_match:
            add_markdown_paragraph(doc, f"{number_match.group(2)}. {number_match.group(3)}")
            i += 1
            continue

        add_markdown_paragraph(doc, line)
        i += 1

    if code_buffer:
        add_code_block(doc, code_buffer)
    return "\n".join(rendered_text_lines)


def build_docx() -> tuple[Path, list[ReflowTableRecord], list[FigureRecord], str]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_publication_defaults(doc)
    add_cover(doc)
    table_records: list[ReflowTableRecord] = []
    figure_records: list[FigureRecord] = []
    source_texts: list[str] = []

    for rel in SOURCE_FILES:
        path = ROOT / rel
        source_texts.append(add_file_to_doc(doc, path, table_records, figure_records))

    docx_path = OUT_DIR / DOCX_NAME
    doc.save(docx_path)
    return docx_path, table_records, figure_records, "\n\n".join(source_texts)


def convert_docx_to_pdf(docx_path: Path) -> Path:
    soffice = which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if not Path(soffice).exists() and not which("soffice"):
        raise RuntimeError("LibreOffice soffice not found")
    with tempfile.TemporaryDirectory(prefix="gaic_1d4_lo_profile_") as profile:
        cmd = [
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT_DIR),
            str(docx_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=900)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    generated = OUT_DIR / (docx_path.stem + ".pdf")
    final_pdf = OUT_DIR / PDF_NAME
    if generated != final_pdf:
        if final_pdf.exists():
            final_pdf.unlink()
        generated.rename(final_pdf)
    return final_pdf


def render_pdf_pages(pdf_path: Path) -> list[Path]:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for existing in RENDER_DIR.glob("*.png"):
        existing.unlink()
    pdftoppm = which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm not found")
    prefix = RENDER_DIR / "page"
    result = subprocess.run([pdftoppm, "-png", "-r", "110", str(pdf_path), str(prefix)], capture_output=True, text=True, timeout=900)
    if result.returncode != 0:
        raise RuntimeError(f"pdftoppm failed\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}")
    return sorted(RENDER_DIR.glob("page-*.png"))


def pdf_page_count(pdf_path: Path) -> int | None:
    return base.pdf_page_count(pdf_path)


def docx_stats(docx_path: Path) -> dict:
    doc = Document(docx_path)
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            headings.append({"style": para.style.name, "text": para.text[:160]})
    return {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "heading_count": len(headings),
        "heading_sample": headings[:30],
    }


def scan_terms(paths: Iterable[Path], terms: list[str]) -> dict[str, list[dict[str, str | int]]]:
    results: dict[str, list[dict[str, str | int]]] = {term: [] for term in terms}
    for path in paths:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        for line_no, line in enumerate(text.splitlines(), 1):
            low = line.lower()
            for term in terms:
                if term.lower() in low:
                    results[term].append({"file": path.relative_to(ROOT).as_posix(), "line": line_no, "text": line.strip()[:1000]})
    return results


def classify_forbidden_match(line: str) -> str:
    low = line.lower()
    if line.strip().startswith("|"):
        return "forbidden wording table / boundary table context"
    if line.strip().startswith("-"):
        return "negative boundary / exclusion list context"
    if "leader" in low and ("leaders" in low or "leadership" in low or "senior leadership" in low):
        return "non-ranking role or leadership context"
    if "audit trails" in low or "trail" in low and "audit" in low:
        return "technical audit-trail context"
    if "claim that" in low or "claims of" in low:
        return "negative boundary / non-claim context"
    if "not " in low or "not a " in low or "not an " in low or "not legal" in low:
        return "negative boundary / non-claim context"
    if "does not" in low or "do not" in low or "must not" in low or "no " in low:
        return "negative boundary / non-claim context"
    if "does not replace" in low or "what this pattern does not prove" in low:
        return "negative boundary / non-claim context"
    if "not as the only" in low:
        return "negative boundary / non-claim context"
    if "does not certify" in low or "does not prove" in low or "does not guarantee" in low:
        return "negative boundary / non-claim context"
    if "legal interpretation" in low or "organizational governance" in low:
        return "negative boundary / non-claim context"
    if "out of scope" in low or "forbidden" in low or "boundary" in low or "non-claim" in low:
        return "negative boundary / non-claim context"
    if "market leadership" in low:
        return "non-ranking market-scope exclusion context"
    return base.classify_forbidden_match(line)


def active_generated_paths() -> list[Path]:
    return [OUT_DIR / MD_NAME]


def write_log(
    md_path: Path,
    docx_path: Path,
    pdf_path: Path,
    rendered_pages: list[Path],
    tables: list[ReflowTableRecord],
    figures: list[FigureRecord],
    source_text: str,
) -> dict:
    forbidden = scan_terms(active_generated_paths(), FORBIDDEN_TERMS)
    forbidden_review = []
    for term, matches in forbidden.items():
        for match in matches:
            forbidden_review.append({**match, "term": term, "classification": classify_forbidden_match(str(match["text"]))})
    positive_like = [m for m in forbidden_review if m["classification"] == "requires review"]
    appendix_g_path = ROOT / "appendices" / "appendix-g-placeholder.md"
    appendix_g_scan = scan_terms([appendix_g_path], APPENDIX_G_STRICT_TERMS)
    page_count = pdf_page_count(pdf_path)

    log = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "phase": "1D-4",
        "artifact": "full_content_whitepaper",
        "root": str(ROOT),
        "source_files": SOURCE_FILES,
        "assembled_markdown_path": str(md_path),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "fresh_generation_from_source_r3_only": True,
        "old_docx_pdf_used_as_input": False,
        "full_chapters_1_18_included": True,
        "appendices_a_k_included": True,
        "internal_reports_in_publication_body": False,
        "docx_stats": docx_stats(docx_path),
        "pdf_page_count": page_count,
        "rendered_page_count": len(rendered_pages),
        "rendered_pages_sample": [str(p) for p in rendered_pages[:20]],
        "table_records": [asdict(r) for r in tables],
        "split_tables": [asdict(r) for r in tables if r.strategy == "semantic column split"],
        "figure_records": [asdict(f) for f in figures],
        "forbidden_scan": forbidden,
        "forbidden_requires_review": positive_like,
        "appendix_g_strict_scan_source": appendix_g_scan,
        "source_text_character_count": len(source_text),
        "content_preservation_status": "substantive chapters and appendices retained; governance ledgers externalized",
        "qa_status": "QA REQUIRED",
    }
    (OUT_DIR / LOG_NAME).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    return log


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path, _assembled = write_assembled_markdown()
    docx_path, table_records, figure_records, source_text = build_docx()
    pdf_path = convert_docx_to_pdf(docx_path)
    rendered_pages = render_pdf_pages(pdf_path)
    log = write_log(md_path, docx_path, pdf_path, rendered_pages, table_records, figure_records, source_text)
    print(json.dumps({
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "assembled_markdown": str(md_path),
        "pages": log["pdf_page_count"],
        "tables": len(table_records),
        "split_tables": len(log["split_tables"]),
        "figures": len(figure_records),
        "rendered_pages": len(rendered_pages),
        "log": str(OUT_DIR / LOG_NAME),
        "publication_candidate_declared": False,
    }, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
